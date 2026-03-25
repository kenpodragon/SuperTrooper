"""Tests for template_builder — placeholder .docx template generation."""

import json
import os
import tempfile

import pytest
from docx import Document

from utils.template_builder import build_template

# Test resume path
TEST_RESUME = r"c:\Users\ssala\OneDrive\Desktop\Resumes\Archived\Originals\Stephen_Salaka_Resume_v32.docx"


@pytest.fixture
def output_dir():
    """Create a temporary directory for test outputs."""
    with tempfile.TemporaryDirectory() as tmpdir:
        yield tmpdir


@pytest.fixture
def output_paths(output_dir):
    """Return output file paths for template and map."""
    return {
        "docx": os.path.join(output_dir, "template.docx"),
        "map": os.path.join(output_dir, "template_map.json"),
    }


class TestBuildTemplateOutputFiles:
    """Verify that build_template creates the expected output files."""

    def test_creates_output_docx(self, output_paths):
        build_template(TEST_RESUME, output_paths["docx"], output_paths["map"])
        assert os.path.exists(output_paths["docx"]), "Output .docx was not created"

    def test_creates_output_map_json(self, output_paths):
        build_template(TEST_RESUME, output_paths["docx"], output_paths["map"])
        assert os.path.exists(output_paths["map"]), "Output template_map.json was not created"


class TestTemplateMapContent:
    """Verify template_map has correct keys and structure."""

    def test_map_has_header_keys(self, output_paths):
        build_template(TEST_RESUME, output_paths["docx"], output_paths["map"])
        with open(output_paths["map"], "r") as f:
            tmap = json.load(f)
        header_keys = [k for k in tmap if k.startswith("HEADER")]
        assert len(header_keys) > 0, "Template map should contain HEADER keys"

    def test_map_has_job_keys(self, output_paths):
        build_template(TEST_RESUME, output_paths["docx"], output_paths["map"])
        with open(output_paths["map"], "r") as f:
            tmap = json.load(f)
        job_keys = [k for k in tmap if k.startswith("JOB")]
        assert len(job_keys) > 0, "Template map should contain JOB keys"

    def test_map_has_bullet_keys(self, output_paths):
        build_template(TEST_RESUME, output_paths["docx"], output_paths["map"])
        with open(output_paths["map"], "r") as f:
            tmap = json.load(f)
        bullet_keys = [k for k in tmap if "BULLET" in k]
        assert len(bullet_keys) > 0, "Template map should contain BULLET keys"

    def test_each_slot_has_required_fields(self, output_paths):
        build_template(TEST_RESUME, output_paths["docx"], output_paths["map"])
        with open(output_paths["map"], "r") as f:
            tmap = json.load(f)
        for slot_name, slot_data in tmap.items():
            assert "type" in slot_data, f"Slot {slot_name} missing 'type'"
            assert "original_text" in slot_data, f"Slot {slot_name} missing 'original_text'"


class TestOutputDocxContent:
    """Verify the output .docx contains placeholders and preserves formatting."""

    def test_docx_contains_placeholder_markers(self, output_paths):
        build_template(TEST_RESUME, output_paths["docx"], output_paths["map"])
        doc = Document(output_paths["docx"])
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "{{" in all_text, "Output .docx should contain {{ placeholder markers"
        assert "}}" in all_text, "Output .docx should contain }} placeholder markers"

    def test_docx_preserves_bold_formatting(self, output_paths):
        build_template(TEST_RESUME, output_paths["docx"], output_paths["map"])
        doc = Document(output_paths["docx"])
        bold_found = False
        for p in doc.paragraphs:
            if "{{" in p.text and p.runs:
                if p.runs[0].bold:
                    bold_found = True
                    break
        assert bold_found, "Output .docx should preserve bold formatting on at least one placeholder"


class TestReturnValue:
    """Verify build_template returns the expected summary dict."""

    def test_return_has_slot_count(self, output_paths):
        result = build_template(TEST_RESUME, output_paths["docx"], output_paths["map"])
        assert "slot_count" in result
        assert isinstance(result["slot_count"], int)
        assert result["slot_count"] > 0

    def test_return_has_sections_detected(self, output_paths):
        result = build_template(TEST_RESUME, output_paths["docx"], output_paths["map"])
        assert "sections_detected" in result
        assert isinstance(result["sections_detected"], list)

    def test_return_has_layout(self, output_paths):
        result = build_template(TEST_RESUME, output_paths["docx"], output_paths["map"])
        assert "layout" in result
