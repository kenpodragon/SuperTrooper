"""Tests for section_generator — .docx paragraph cloning engine."""

import io

import pytest
from docx import Document

from utils.section_generator import (
    clone_paragraph,
    fill_paragraph,
    generate_from_sections,
    remove_paragraph,
)


# ---------------------------------------------------------------------------
# Helpers to build in-memory test fixtures
# ---------------------------------------------------------------------------

def _make_template_doc() -> bytes:
    """Create minimal template .docx with section markers.

    Para layout:
      0: {{HEADER}}
      1: {{HEADLINE}}
      2: {{SUMMARY}}
      3: Certifications          (section header)
      4: {{CERTIFICATIONS}}      (prototype, repeating)
      5: Education               (section header)
      6: {{EDUCATION}}           (prototype, repeating)
    """
    doc = Document()
    doc.add_paragraph("{{HEADER}}")           # para 0
    doc.add_paragraph("{{HEADLINE}}")         # para 1
    doc.add_paragraph("{{SUMMARY}}")          # para 2
    doc.add_paragraph("Certifications")       # para 3
    doc.add_paragraph("{{CERTIFICATIONS}}")   # para 4
    doc.add_paragraph("Education")            # para 5
    doc.add_paragraph("{{EDUCATION}}")        # para 6
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_template_map() -> dict:
    return {
        "HEADER": {
            "repeating": False,
            "para_index": 0,
            "format": "simple",
        },
        "HEADLINE": {
            "repeating": False,
            "para_index": 1,
            "format": "simple",
        },
        "SUMMARY": {
            "repeating": False,
            "para_index": 2,
            "format": "simple",
        },
        "CERTIFICATIONS": {
            "repeating": True,
            "para_index": 4,
            "format": "bold_label",
            "separator": " | ",
            "section_header_para": 3,
        },
        "EDUCATION": {
            "repeating": True,
            "para_index": 6,
            "format": "bold_label",
            "separator": " | ",
            "section_header_para": 5,
        },
    }


def _doc_texts(blob: bytes) -> list[str]:
    """Extract all paragraph texts from a .docx bytes blob."""
    doc = Document(io.BytesIO(blob))
    return [p.text for p in doc.paragraphs]


# ---------------------------------------------------------------------------
# Tests: clone_paragraph
# ---------------------------------------------------------------------------

class TestCloneParagraph:
    def test_clone_paragraph_preserves_text(self):
        doc = Document()
        para = doc.add_paragraph("original text")
        anchor = doc.add_paragraph("anchor")

        cloned = clone_paragraph(doc, para, anchor)

        assert cloned.text == "original text"

    def test_clone_paragraph_inserts_after_anchor(self):
        doc = Document()
        para = doc.add_paragraph("source")
        anchor = doc.add_paragraph("anchor")
        doc.add_paragraph("after anchor")

        clone_paragraph(doc, para, anchor)

        texts = [p.text for p in doc.paragraphs]
        anchor_idx = texts.index("anchor")
        assert texts[anchor_idx + 1] == "source"

    def test_clone_paragraph_is_independent(self):
        doc = Document()
        para = doc.add_paragraph("original")
        anchor = doc.add_paragraph("anchor")

        cloned = clone_paragraph(doc, para, anchor)
        cloned.runs[0].text = "modified"

        # Original should be unchanged
        assert para.text == "original"


# ---------------------------------------------------------------------------
# Tests: fill_paragraph
# ---------------------------------------------------------------------------

class TestFillParagraph:
    def test_fill_paragraph_simple(self):
        doc = Document()
        para = doc.add_paragraph("{{PLACEHOLDER}}")

        fill_paragraph(para, "Hello World", "simple")

        assert para.text == "Hello World"

    def test_fill_paragraph_simple_preserves_run_count(self):
        doc = Document()
        para = doc.add_paragraph("First run")
        para.runs[0].text = "First run"

        fill_paragraph(para, "Replaced text", "simple")

        # Should fill first run, clear others
        assert "Replaced text" in para.text

    def test_fill_paragraph_bold_label_splits_at_separator(self):
        doc = Document()
        para = doc.add_paragraph("CSM | Scrum Alliance")

        fill_paragraph(para, "CSM | Scrum Alliance", "bold_label", " | ")

        # First run should be bold
        texts = [r.text for r in para.runs if r.text]
        assert any("CSM" in t for t in texts)
        # At least one run should be bold
        bold_runs = [r for r in para.runs if r.bold and r.text.strip()]
        assert len(bold_runs) >= 1
        assert "CSM" in bold_runs[0].text

    def test_fill_paragraph_bold_label_no_separator_falls_back_to_simple(self):
        doc = Document()
        para = doc.add_paragraph("Plain text no separator")

        fill_paragraph(para, "Plain text no separator", "bold_label", " | ")

        assert para.text == "Plain text no separator"

    def test_fill_paragraph_bold_label_with_colon_separator(self):
        doc = Document()
        para = doc.add_paragraph("Action: result follows")

        fill_paragraph(para, "Action: result follows", "bold_label", ": ")

        bold_runs = [r for r in para.runs if r.bold and r.text.strip()]
        assert len(bold_runs) >= 1
        assert "Action" in bold_runs[0].text


# ---------------------------------------------------------------------------
# Tests: remove_paragraph
# ---------------------------------------------------------------------------

class TestRemoveParagraph:
    def test_remove_paragraph_removes_from_doc(self):
        doc = Document()
        doc.add_paragraph("keep this")
        to_remove = doc.add_paragraph("remove this")
        doc.add_paragraph("keep this too")

        remove_paragraph(to_remove)

        texts = [p.text for p in doc.paragraphs]
        assert "remove this" not in texts
        assert "keep this" in texts
        assert "keep this too" in texts


# ---------------------------------------------------------------------------
# Tests: generate_from_sections
# ---------------------------------------------------------------------------

class TestGenerateFromSections:
    def test_generate_singular_sections_fill_in_place(self):
        blob = _make_template_doc()
        tmap = _make_template_map()

        resolved = {
            "HEADER": {"full_name": "Stephen Salaka", "credentials": "PMP"},
            "HEADLINE": "Senior Program Manager",
            "SUMMARY": "Experienced leader driving complex programs.",
            "CERTIFICATIONS": [],
            "EDUCATION": [],
        }

        result = generate_from_sections(blob, tmap, resolved)
        texts = _doc_texts(result)

        assert "Stephen Salaka, PMP" in texts
        assert "Senior Program Manager" in texts
        assert "Experienced leader driving complex programs." in texts

    def test_generate_repeating_section_clones_n_times(self):
        blob = _make_template_doc()
        tmap = _make_template_map()

        certs = [
            {"name": "CSM", "issuer": "Scrum Alliance"},
            {"name": "PMP", "issuer": "PMI"},
            {"name": "CSPO", "issuer": "Scrum Alliance"},
        ]

        resolved = {
            "HEADER": {"full_name": "Test User", "credentials": ""},
            "HEADLINE": "Test Headline",
            "SUMMARY": "Test Summary",
            "CERTIFICATIONS": certs,
            "EDUCATION": [],
        }

        result = generate_from_sections(blob, tmap, resolved)
        texts = _doc_texts(result)

        assert "CSM | Scrum Alliance" in texts
        assert "PMP | PMI" in texts
        assert "CSPO | Scrum Alliance" in texts
        # Original prototype should be gone
        assert "{{CERTIFICATIONS}}" not in texts

    def test_generate_repeating_1_item_fills_in_place(self):
        blob = _make_template_doc()
        tmap = _make_template_map()

        resolved = {
            "HEADER": {"full_name": "Test", "credentials": ""},
            "HEADLINE": "Headline",
            "SUMMARY": "Summary",
            "CERTIFICATIONS": [{"name": "CSM", "issuer": "Scrum Alliance"}],
            "EDUCATION": [],
        }

        result = generate_from_sections(blob, tmap, resolved)
        texts = _doc_texts(result)

        assert "CSM | Scrum Alliance" in texts
        assert "{{CERTIFICATIONS}}" not in texts

    def test_generate_empty_section_removes_header_and_prototype(self):
        blob = _make_template_doc()
        tmap = _make_template_map()

        resolved = {
            "HEADER": {"full_name": "Test", "credentials": ""},
            "HEADLINE": "Headline",
            "SUMMARY": "Summary",
            "CERTIFICATIONS": [],  # Empty — should remove header + prototype
            "EDUCATION": [{"degree": "BS", "field": "CS", "institution": "MIT", "location": "Cambridge MA"}],
        }

        result = generate_from_sections(blob, tmap, resolved)
        texts = _doc_texts(result)

        # Section header "Certifications" and prototype should be removed
        assert "Certifications" not in texts
        assert "{{CERTIFICATIONS}}" not in texts
        # Education should still be present
        assert "BS | CS | MIT | Cambridge MA" in texts

    def test_generate_placeholder_not_in_resolved_leaves_marker(self):
        """Sections not in resolved dict are left untouched (no crash)."""
        blob = _make_template_doc()
        tmap = _make_template_map()

        # Only resolve HEADER — everything else omitted
        resolved = {
            "HEADER": {"full_name": "Stephen", "credentials": "PMP"},
        }

        # Should not raise
        result = generate_from_sections(blob, tmap, resolved)
        texts = _doc_texts(result)
        assert "Stephen, PMP" in texts

    def test_generate_education_multiple_items(self):
        blob = _make_template_doc()
        tmap = _make_template_map()

        education = [
            {"degree": "MBA", "field": "Finance", "institution": "Wharton", "location": "Philadelphia PA"},
            {"degree": "BS", "field": "Engineering", "institution": "Penn State", "location": "University Park PA"},
        ]

        resolved = {
            "HEADER": {"full_name": "Test", "credentials": ""},
            "HEADLINE": "Headline",
            "SUMMARY": "Summary",
            "CERTIFICATIONS": [],
            "EDUCATION": education,
        }

        result = generate_from_sections(blob, tmap, resolved)
        texts = _doc_texts(result)

        assert "MBA | Finance | Wharton | Philadelphia PA" in texts
        assert "BS | Engineering | Penn State | University Park PA" in texts
        assert "{{EDUCATION}}" not in texts

    def test_generate_returns_bytes(self):
        blob = _make_template_doc()
        tmap = _make_template_map()
        resolved = {}

        result = generate_from_sections(blob, tmap, resolved)
        assert isinstance(result, bytes)
        assert len(result) > 0


# ---------------------------------------------------------------------------
# Tests: EXPERIENCE compound section
# ---------------------------------------------------------------------------

class TestGenerateExperienceSection:
    def _make_experience_doc(self) -> bytes:
        doc = Document()
        doc.add_paragraph("{{HEADER}}")        # 0
        doc.add_paragraph("Experience")        # 1 (section header)
        doc.add_paragraph("{{EXPERIENCE}}")    # 2 (prototype)
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def _make_experience_tmap(self) -> dict:
        return {
            "HEADER": {"repeating": False, "para_index": 0, "format": "simple"},
            "EXPERIENCE": {
                "repeating": True,
                "para_index": 2,
                "format": "bold_label",
                "separator": ", ",
                "section_header_para": 1,
            },
        }

    def test_experience_generates_company_and_job_lines(self):
        blob = self._make_experience_doc()
        tmap = self._make_experience_tmap()

        resolved = {
            "HEADER": {"full_name": "Test", "credentials": ""},
            "EXPERIENCE": [
                {
                    "company": {
                        "employer": "Acme Corp",
                        "location": "New York NY",
                        "industry": "Tech",
                    },
                    "jobs": [
                        {
                            "job": {
                                "title": "Senior Manager",
                                "start_date": "2020-01-01",
                                "end_date": None,
                                "is_current": True,
                            },
                            "bullets": [
                                {"text": "Led team: delivered $2M savings"},
                                {"text": "Drove process: reduced cycle time 40%"},
                            ],
                        }
                    ],
                }
            ],
        }

        result = generate_from_sections(blob, tmap, resolved)
        texts = _doc_texts(result)

        # Company header present
        assert any("Acme Corp" in t for t in texts)
        # Job title present
        assert any("Senior Manager" in t for t in texts)
        # Bullets present
        assert any("Led team" in t for t in texts)
        assert any("Drove process" in t for t in texts)
        # Prototype removed
        assert "{{EXPERIENCE}}" not in texts

    def test_experience_empty_removes_section_header(self):
        blob = self._make_experience_doc()
        tmap = self._make_experience_tmap()

        resolved = {
            "HEADER": {"full_name": "Test", "credentials": ""},
            "EXPERIENCE": [],
        }

        result = generate_from_sections(blob, tmap, resolved)
        texts = _doc_texts(result)

        assert "Experience" not in texts
        assert "{{EXPERIENCE}}" not in texts
