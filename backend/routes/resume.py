"""Routes for resume generation, recipes, templates, header, education, certifications."""

from flask import Blueprint, request, jsonify, send_file, Response
import io
import json
import logging
import re
import db
from docx import Document
from ai_providers.router import route_inference
from PIL import Image, ImageDraw, ImageFont
from template_dedup import check_duplicates, compute_hash

logger = logging.getLogger(__name__)

bp = Blueprint("resume", __name__)


# ---------------------------------------------------------------------------
# V1 → V2 Resolved Adapter
# ---------------------------------------------------------------------------

def _v1_resolved_to_v2(flat: dict, recipe_json: dict = None) -> dict:
    """Convert a v1 flat resolved slot map into v2 structured format for the editor.

    V1 keys follow patterns: HEADER_NAME, HEADER_CONTACT_N, SUMMARY, HEADLINE,
    HIGHLIGHT_N, JOB_N_HEADER, JOB_N_BULLET_N, JOB_N_INTRO, SKILLS_N, KEYWORDS_N,
    EDUCATION_N, CERT_N, SLOT_N.

    If recipe_json is provided, career_history references are resolved to
    separate employer/title/dates fields instead of a single header blob.
    """
    v2 = {}

    # Header
    header = {}
    if "HEADER_NAME" in flat:
        header["name"] = flat["HEADER_NAME"]
    contacts = []
    for k in sorted(flat):
        if k.startswith("HEADER_CONTACT"):
            contacts.append(flat[k])
    if contacts:
        header["contact"] = " • ".join(contacts)
    if header:
        v2["header"] = header

    # Headline
    if "HEADLINE" in flat:
        v2["headline"] = flat["HEADLINE"]

    # Summary
    if "SUMMARY" in flat:
        v2["summary"] = flat["SUMMARY"]

    # Highlights
    highlights = []
    for k in sorted(flat, key=lambda x: (len(x), x)):
        if k.startswith("HIGHLIGHT_"):
            highlights.append(flat[k])
    if highlights:
        v2["highlights"] = highlights

    # Experience — group JOB_N_HEADER + JOB_N_BULLET_M + JOB_N_INTRO
    jobs = {}  # job_num -> {header, bullets[], intro}
    for k in flat:
        m = re.match(r"JOB_(\d+)_(HEADER|BULLET_\d+|INTRO)", k)
        if m:
            num = int(m.group(1))
            part = m.group(2)
            if num not in jobs:
                jobs[num] = {"bullets": []}
            if part == "HEADER":
                jobs[num]["header"] = flat[k]
            elif part == "INTRO":
                jobs[num]["synopsis"] = flat[k]
            elif part.startswith("BULLET_"):
                jobs[num]["bullets"].append((int(part.split("_")[1]), flat[k]))

    if jobs:

        def _fmt_date_v1(d):
            """Format date for v1 experience display."""
            if not d or not hasattr(d, "strftime"):
                return str(d) if d else ""
            return str(d.year) if d.month == 1 and d.day == 1 else d.strftime("%B %Y")

        experience = []
        for num in sorted(jobs):
            job = jobs[num]
            exp_item = {}
            header_text = job.get("header", "")

            # Try to get structured employer/title from the original recipe reference
            parsed = False
            if recipe_json:
                ref = recipe_json.get(f"JOB_{num}_HEADER")
                if ref and isinstance(ref, dict) and ref.get("table") == "career_history":
                    row_id = ref.get("id")
                    if row_id:
                        ch = db.query_one(
                            "SELECT employer, title, start_date, end_date, is_current, location "
                            "FROM career_history WHERE id = %s",
                            (row_id,),
                        )
                        if ch:
                            exp_item["employer"] = ch["employer"]
                            title_val = ch.get("title", "")
                            if title_val and title_val not in ("_COMPANY_OVERVIEW", "Unknown"):
                                exp_item["title"] = title_val
                            if ch.get("location"):
                                exp_item["location"] = ch["location"]
                            if ch.get("start_date"):
                                start = _fmt_date_v1(ch["start_date"])
                                if ch.get("is_current"):
                                    exp_item["dates"] = f"{start} \u2013 Present"
                                elif ch.get("end_date"):
                                    exp_item["dates"] = f"{start} \u2013 {_fmt_date_v1(ch['end_date'])}"
                                else:
                                    exp_item["dates"] = start
                            parsed = True

            if not parsed and header_text:
                # Skip pure page markers from PDF parsing
                stripped = header_text.strip()
                if re.match(r"^Page \d+ of \d+$", stripped):
                    continue

                # Try heuristic parsing for multiline LinkedIn PDF headers:
                # Employer\nTitle\nDates (duration)\nLocation
                lines = [ln.strip() for ln in header_text.split("\n")]
                lines = [ln for ln in lines if ln and not re.match(r"^Page \d+ of \d+", ln)]
                # Match dates at line start OR embedded after text (e.g., "Title January 2021 - January 2024")
                date_start_re = re.compile(r"^(?:\w+\s+)?\d{4}\s*-\s*(?:\w+\s+)?\d{4}")
                date_embed_re = re.compile(r"(\w+\s+\d{4}\s*-\s*\w+\s+\d{4}|\d{4}\s*-\s*\d{4})")
                date_idx = None
                for idx_l, ln in enumerate(lines):
                    if date_start_re.match(ln):
                        date_idx = idx_l
                        break
                # If no line starts with a date, look for embedded dates (skip first line = employer)
                if date_idx is None:
                    for idx_l, ln in enumerate(lines):
                        if idx_l == 0:
                            continue
                        m = date_embed_re.search(ln)
                        if m:
                            # Split the line: text before date = title, date part = dates
                            pre = ln[:m.start()].strip()
                            if pre:
                                lines[idx_l] = pre
                                lines.insert(idx_l + 1, m.group(0))
                                date_idx = idx_l + 1
                            else:
                                date_idx = idx_l
                            break
                if date_idx is not None and date_idx >= 1:
                    # Lines before date: employer (+ title if 2+)
                    before = lines[:date_idx]
                    if len(before) >= 2:
                        exp_item["employer"] = before[0]
                        exp_item["title"] = before[1]
                    else:
                        exp_item["employer"] = before[0]
                    # Date line: strip "(N years)" suffix
                    date_str = re.sub(r"\s*\(\d+\s*years?\).*", "", lines[date_idx]).strip()
                    if date_str:
                        exp_item["dates"] = date_str
                    # Lines after date: location
                    after = [ln for ln in lines[date_idx + 1:] if ln]
                    if after:
                        exp_item["location"] = after[0]
                elif lines:
                    exp_item["employer"] = header_text

            if job.get("synopsis"):
                exp_item["synopsis"] = job["synopsis"]
            # Sort bullets by index
            sorted_bullets = [b[1] for b in sorted(job["bullets"])]
            if sorted_bullets:
                exp_item["bullets"] = sorted_bullets
            experience.append(exp_item)
        v2["experience"] = experience

    # Skills + Keywords
    skills = []
    for k in sorted(flat, key=lambda x: (len(x), x)):
        if k.startswith("SKILLS_") or k.startswith("KEYWORDS_"):
            skills.append(flat[k])
    if skills:
        v2["skills"] = skills

    # Education
    education = []
    for k in sorted(flat, key=lambda x: (len(x), x)):
        if k.startswith("EDUCATION_"):
            education.append(flat[k])
    if education:
        v2["education"] = education

    # Certifications
    certs = []
    for k in sorted(flat, key=lambda x: (len(x), x)):
        if k.startswith("CERT_"):
            certs.append(flat[k])
    if certs:
        v2["certifications"] = certs

    # Additional experience (SLOT_N — often additional one-liner jobs)
    additional = []
    for k in sorted(flat, key=lambda x: (len(x), x)):
        if k.startswith("SLOT_"):
            additional.append(flat[k])
    if additional:
        v2["additional_experience"] = additional

    return v2


# ---------------------------------------------------------------------------
# Resume Recipes
# ---------------------------------------------------------------------------

@bp.route("/api/resume/recipes", methods=["GET"])
def list_recipes():
    """List resume recipes. Query params: template_id, is_active."""
    template_id = request.args.get("template_id", type=int)
    is_active = request.args.get("is_active", "true").lower() != "false"

    sql = "SELECT id, name, description, headline, template_id, application_id, is_active, created_at, updated_at FROM resume_recipes WHERE 1=1"
    params = []
    if template_id:
        sql += " AND template_id = %s"
        params.append(template_id)
    if is_active:
        sql += " AND is_active = TRUE"
    sql += " ORDER BY id"
    rows = db.query(sql, params)
    return jsonify({"recipes": rows, "count": len(rows)})


@bp.route("/api/resume/recipes/<int:recipe_id>", methods=["GET"])
def get_recipe(recipe_id):
    """Get a single recipe with full JSON. Query params: resolve=true for text preview."""
    row = db.query_one("SELECT * FROM resume_recipes WHERE id = %s", (recipe_id,))
    if not row:
        return jsonify({"error": f"Recipe id={recipe_id} not found"}), 404

    if request.args.get("resolve", "").lower() == "true":
        from mcp_tools_resume_gen import _resolve_recipe_db
        recipe_json = row["recipe"]
        if isinstance(recipe_json, str):
            recipe_json = json.loads(recipe_json)
        resolved = _resolve_recipe_db(recipe_json, recipe_version=row.get("recipe_version", 1))
        if row.get("headline"):
            resolved["HEADLINE"] = row["headline"]
        # Convert v1 flat slot map to v2 structure for the editor
        if row.get("recipe_version", 1) < 2:
            resolved = _v1_resolved_to_v2(resolved, recipe_json=recipe_json)
        row["resolved_preview"] = resolved

    # Include template_name alongside template_id
    template_name = ""
    if row.get("template_id"):
        tpl = db.query_one("SELECT name FROM resume_templates WHERE id = %s", (row["template_id"],))
        template_name = tpl["name"] if tpl else ""
    row["template_name"] = template_name

    return jsonify(row)


@bp.route("/api/resume/recipes/<int:recipe_id>/validate", methods=["GET"])
def validate_recipe(recipe_id):
    """Validate a recipe's references. Returns which IDs exist and which are missing."""
    row = db.query_one("SELECT * FROM resume_recipes WHERE id = %s", (recipe_id,))
    if not row:
        return jsonify({"error": f"Recipe id={recipe_id} not found"}), 404

    ALLOWED = {"bullets", "career_history", "skills", "summary_variants",
               "education", "certifications", "resume_header"}
    recipe_json = row["recipe"]
    if isinstance(recipe_json, str):
        recipe_json = json.loads(recipe_json)

    recipe_version = row.get("recipe_version", 1)
    if recipe_version >= 2:
        ALLOWED_V2 = {"bullets", "career_history", "skills", "summary_variants",
                       "education", "certifications", "resume_header"}
        errors = []
        # Check single-item refs
        for key in ("header", "headline", "summary"):
            if key in recipe_json and isinstance(recipe_json[key], dict):
                ref = recipe_json[key]
                if "ref" in ref:
                    table = ref["ref"]
                    if table not in ALLOWED_V2:
                        errors.append({"section": key, "error": f"Table '{table}' not allowed"})
                    else:
                        row_id = ref.get("id", 1)
                        found = db.query_one(f"SELECT id FROM {table} WHERE id = %s", (row_id,))
                        if not found:
                            errors.append({"section": key, "table": table, "id": row_id, "error": "Not found"})

        # Check experience array
        for job_idx, job in enumerate(recipe_json.get("experience", [])):
            if "ref" in job and job["ref"] == "career_history":
                found = db.query_one("SELECT id FROM career_history WHERE id = %s", (job.get("id"),))
                if not found:
                    errors.append({"section": f"experience[{job_idx}]", "table": "career_history", "id": job.get("id"), "error": "Not found"})
            for b_idx, bullet in enumerate(job.get("bullets", [])):
                if "ref" in bullet and bullet["ref"] in ALLOWED_V2:
                    found = db.query_one(f"SELECT id FROM {bullet['ref']} WHERE id = %s", (bullet.get("id"),))
                    if not found:
                        errors.append({"section": f"experience[{job_idx}].bullets[{b_idx}]", "table": bullet["ref"], "id": bullet.get("id"), "error": "Not found"})

        return jsonify({"valid": len(errors) == 0, "errors": errors, "recipe_version": 2})

    errors = []
    db_refs = 0
    literals = 0
    valid_refs = 0

    for slot_name, ref in recipe_json.items():
        if "literal" in ref:
            literals += 1
            continue
        if "ids" in ref:
            table = ref["table"]
            if table not in ALLOWED:
                errors.append({"slot": slot_name, "error": f"Table '{table}' not allowed"})
                continue
            ids = ref["ids"]
            db_refs += len(ids)
            found = db.query(f"SELECT id FROM {table} WHERE id = ANY(%s)", (ids,))
            found_ids = {r["id"] for r in found}
            missing = [i for i in ids if i not in found_ids]
            if missing:
                errors.append({"slot": slot_name, "table": table, "missing_ids": missing})
            else:
                valid_refs += len(ids)
        elif "table" in ref:
            table = ref["table"]
            if table not in ALLOWED:
                errors.append({"slot": slot_name, "error": f"Table '{table}' not allowed"})
                continue
            row_id = ref.get("id", 1)
            db_refs += 1
            found = db.query_one(f"SELECT id FROM {table} WHERE id = %s", (row_id,))
            if found:
                valid_refs += 1
            else:
                errors.append({"slot": slot_name, "table": table, "id": row_id, "error": "Not found"})

    return jsonify({
        "valid": len(errors) == 0,
        "errors": errors,
        "stats": {
            "total_slots": len(recipe_json),
            "db_refs": db_refs,
            "literals": literals,
            "valid_refs": valid_refs,
            "missing_refs": len(errors),
        },
    })


@bp.route("/api/resume/recipes", methods=["POST"])
def create_recipe():
    """Create a new recipe. Body: {name, headline, template_id, recipe, description?, application_id?}."""
    data = request.get_json(silent=True) or {}
    name = data.get("name")
    template_id = data.get("template_id")
    recipe = data.get("recipe")
    if not name or not template_id or not recipe:
        return jsonify({"error": "name, template_id, and recipe are required"}), 400

    row = db.execute_returning(
        """INSERT INTO resume_recipes (name, description, headline, template_id, recipe, application_id)
           VALUES (%s, %s, %s, %s, %s, %s) RETURNING *""",
        (name, data.get("description"), data.get("headline"),
         template_id, json.dumps(recipe), data.get("application_id")),
    )
    return jsonify(row), 201


@bp.route("/api/resume/recipes/<int:recipe_id>", methods=["PUT"])
def update_recipe(recipe_id):
    """Update a recipe. Body: partial fields to update."""
    existing = db.query_one("SELECT * FROM resume_recipes WHERE id = %s", (recipe_id,))
    if not existing:
        return jsonify({"error": f"Recipe id={recipe_id} not found"}), 404

    data = request.get_json(silent=True) or {}
    fields = []
    values = []
    for col in ["name", "description", "headline", "is_active"]:
        if col in data:
            fields.append(f"{col} = %s")
            values.append(data[col])
    if "recipe" in data:
        fields.append("recipe = %s")
        values.append(json.dumps(data["recipe"]))
    if "application_id" in data:
        fields.append("application_id = %s")
        values.append(data["application_id"])

    if not fields:
        return jsonify(existing)

    fields.append("updated_at = NOW()")
    values.append(recipe_id)
    row = db.execute_returning(
        f"UPDATE resume_recipes SET {', '.join(fields)} WHERE id = %s RETURNING *",
        values,
    )
    return jsonify(row)


@bp.route("/api/resume/recipes/<int:recipe_id>", methods=["DELETE"])
def delete_recipe(recipe_id):
    """Soft-delete a recipe (set is_active=false)."""
    existing = db.query_one("SELECT * FROM resume_recipes WHERE id = %s", (recipe_id,))
    if not existing:
        return jsonify({"error": f"Recipe id={recipe_id} not found"}), 404
    db.execute(
        "UPDATE resume_recipes SET is_active = FALSE, updated_at = NOW() WHERE id = %s",
        (recipe_id,),
    )
    return jsonify({"status": "deleted", "id": recipe_id})


@bp.route("/api/resume/recipes/<int:recipe_id>/clone", methods=["POST"])
def clone_recipe(recipe_id):
    """Clone a recipe as a new entry for tailoring."""
    existing = db.query_one("SELECT * FROM resume_recipes WHERE id = %s", (recipe_id,))
    if not existing:
        return jsonify({"error": f"Recipe id={recipe_id} not found"}), 404

    data = request.get_json(silent=True) or {}
    new_name = data.get("name", f"{existing['name']} (copy)")

    row = db.execute_returning(
        """INSERT INTO resume_recipes (name, description, headline, template_id, recipe, application_id)
           VALUES (%s, %s, %s, %s, %s, %s) RETURNING *""",
        (new_name, existing["description"], existing["headline"],
         existing["template_id"], json.dumps(existing["recipe"]),
         data.get("application_id")),
    )
    return jsonify(row), 201


@bp.route("/api/resume/recipes/<int:recipe_id>/clone-item", methods=["POST"])
def clone_recipe_item(recipe_id):
    """Clone a DB record. Frontend swaps the recipe ref and autosaves.
    Body: {"table": "bullets", "id": 26}
    Returns: {"id": 142, "table": "bullets", "text": "..."}
    """
    ALLOWED_CLONE = {"bullets", "summary_variants", "education", "certifications"}

    recipe = db.query_one("SELECT id FROM resume_recipes WHERE id = %s", (recipe_id,))
    if not recipe:
        return jsonify({"error": f"Recipe {recipe_id} not found"}), 404

    data = request.get_json(silent=True) or {}
    table = data.get("table")
    source_id = data.get("id")

    if not table or not source_id:
        return jsonify({"error": "table and id required"}), 400
    if table not in ALLOWED_CLONE:
        return jsonify({"error": f"Cannot clone from table '{table}'"}), 400

    original = db.query_one(f"SELECT * FROM {table} WHERE id = %s", (source_id,))
    if not original:
        return jsonify({"error": f"{table} id={source_id} not found"}), 404

    # Clone based on table type — append ' [copy]' to text fields with unique constraints
    if table == "bullets":
        clone = db.execute_returning(
            """INSERT INTO bullets (career_history_id, text, type, star_situation,
               star_task, star_action, star_result, tags, role_suitability,
               industry_suitability, display_order)
               SELECT career_history_id, text || ' [copy]', type, star_situation,
               star_task, star_action, star_result, tags, role_suitability,
               industry_suitability, display_order
               FROM bullets WHERE id = %s RETURNING *""",
            (source_id,),
        )
    elif table == "summary_variants":
        clone = db.execute_returning(
            """INSERT INTO summary_variants (role_type, text)
               SELECT role_type || ' (copy)', text
               FROM summary_variants WHERE id = %s RETURNING *""",
            (source_id,),
        )
    elif table == "education":
        clone = db.execute_returning(
            """INSERT INTO education (degree, field, institution, location, type, sort_order)
               SELECT degree, field, institution, location, type, sort_order
               FROM education WHERE id = %s RETURNING *""",
            (source_id,),
        )
    elif table == "certifications":
        clone = db.execute_returning(
            """INSERT INTO certifications (name, issuer, is_active, sort_order)
               SELECT name || ' (copy)', issuer, is_active, sort_order
               FROM certifications WHERE id = %s RETURNING *""",
            (source_id,),
        )

    if not clone:
        return jsonify({"error": "Clone failed"}), 500

    display_col = {"bullets": "text", "summary_variants": "text",
                   "education": "degree", "certifications": "name"}
    text = clone.get(display_col.get(table, "id"), "")

    return jsonify({"id": clone["id"], "table": table, "text": text}), 201


@bp.route("/api/resume/recipes/<int:recipe_id>/autosave", methods=["PUT"])
def autosave_recipe(recipe_id):
    """Debounced save of full recipe JSON + theme.
    Body: {"recipe": {...v2 JSON...}, "theme": {...theme JSON...}}
    """
    existing = db.query_one("SELECT id FROM resume_recipes WHERE id = %s", (recipe_id,))
    if not existing:
        return jsonify({"error": f"Recipe {recipe_id} not found"}), 404

    data = request.get_json(silent=True) or {}
    recipe_json = data.get("recipe")
    theme = data.get("theme")

    updates = ["updated_at = NOW()"]
    params = []

    if recipe_json is not None:
        updates.append("recipe = %s")
        params.append(json.dumps(recipe_json))
        # Only stamp as v2 if the recipe has v2 structure
        if isinstance(recipe_json.get("experience"), list) or "header" in recipe_json:
            updates.append("recipe_version = 2")

    if theme is not None:
        updates.append("theme = %s")
        params.append(json.dumps(theme))

    params.append(recipe_id)
    row = db.execute_returning(
        f"UPDATE resume_recipes SET {', '.join(updates)} WHERE id = %s RETURNING updated_at",
        tuple(params),
    )

    # Validation warnings
    warnings = []
    if recipe_json:
        REQUIRED = {"header", "headline", "summary", "experience"}
        for section in REQUIRED:
            if section not in recipe_json:
                warnings.append({"section": section, "message": f"{section.title()} section is missing"})
            elif isinstance(recipe_json[section], list) and len(recipe_json[section]) == 0:
                warnings.append({"section": section, "message": f"{section.title()} section is empty"})

    return jsonify({
        "saved": True,
        "warnings": warnings,
        "updated_at": row["updated_at"].isoformat() if row else None,
    })


@bp.route("/api/resume/recipes/<int:recipe_id>/generate", methods=["POST"])
def generate_from_recipe(recipe_id):
    """Generate a .docx resume from a recipe. Returns file download."""
    from mcp_tools_resume_gen import _resolve_recipe_db

    recipe_row = db.query_one("SELECT * FROM resume_recipes WHERE id = %s", (recipe_id,))
    if not recipe_row:
        return jsonify({"error": f"Recipe id={recipe_id} not found"}), 404

    tmpl = db.query_one(
        "SELECT template_blob, template_map FROM resume_templates WHERE id = %s",
        (recipe_row["template_id"],),
    )
    if not tmpl:
        return jsonify({"error": f"Template id={recipe_row['template_id']} not found"}), 404

    template_blob = bytes(tmpl["template_blob"])
    template_map = tmpl["template_map"] or {}
    recipe_json = recipe_row["recipe"]
    if isinstance(recipe_json, str):
        recipe_json = json.loads(recipe_json)

    resolved_content = _resolve_recipe_db(recipe_json, recipe_version=recipe_row.get("recipe_version", 1))
    if recipe_row.get("headline"):
        resolved_content["HEADLINE"] = recipe_row["headline"]

    # --- AI routing: enhance resolved recipe content ---
    ai_context = {
        "task_type": "generate_from_recipe",
        "recipe_id": recipe_id,
        "resolved_content": resolved_content,
    }

    def _python_recipe_content(ctx):
        return {"content": ctx["resolved_content"]}

    def _ai_recipe_content(ctx):
        # AI polishing disabled during generation to preserve content fidelity.
        # AI enhancement happens via the ai-review and ai-generate-slot endpoints instead.
        return {"content": ctx["resolved_content"]}

    gen_result = route_inference(
        task="generate_from_recipe",
        context=ai_context,
        python_fallback=_python_recipe_content,
        ai_handler=_ai_recipe_content,
    )
    content = gen_result["content"]

    # If v2 recipe, flatten resolved dict to placeholder map
    recipe_version = recipe_row.get("recipe_version", 1)
    if recipe_version >= 2 and isinstance(content, dict) and any(k in content for k in ("experience", "header", "highlights")):
        content = _flatten_v2_for_template(content)

    # Build slot info
    slot_info = {}
    for slot in template_map.get("slots", []):
        if slot.get("placeholder"):
            slot_info[slot["placeholder"]] = {
                "slot_type": slot.get("slot_type", ""),
                "formatting": slot.get("formatting", {}),
            }

    # Fill template
    PLACEHOLDER_RE = re.compile(r"\{\{([A-Z0-9_]+)\}\}")
    BOLD_SEPS = {
        "highlight": ": ", "job_bullet": ": ", "education": " | ",
        "certification": " | ", "additional_exp": " | ", "ref_link": " | ",
        "job_header": ", ",
    }

    doc = Document(io.BytesIO(template_blob))
    for para in doc.paragraphs:
        match = PLACEHOLDER_RE.search(para.text)
        if not match:
            continue
        placeholder = match.group(1)
        if placeholder not in content:
            _fill_simple(para, "")
            continue
        text = content[placeholder]
        info = slot_info.get(placeholder, {})
        slot_type = info.get("slot_type", "")
        formatting = info.get("formatting", {})
        if formatting.get("bold_label") and slot_type in BOLD_SEPS:
            _fill_bold_label(para, text, BOLD_SEPS[slot_type])
        else:
            _fill_simple(para, text)

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)

    # Build filename: {Name}_{TargetRole}_{Date}.docx
    import re as _re
    from datetime import date as _date
    _header_name = (resolved_content.get("HEADER_NAME") or "Resume").replace("\n", " ").replace("\r", "").strip()
    # Take only the first meaningful part (before bullet/pipe separators or commas after credentials)
    _name_part = _re.split(r'[•|]|\s{2,}', _header_name)[0].strip().replace(" ", "")
    _name_part = _re.sub(r'[^a-zA-Z0-9_,]', '', _name_part)[:40]
    _raw_role = recipe_row.get("headline") or recipe_row.get("name") or "General"
    _raw_role = _raw_role.replace("\n", " ").replace("\r", "")
    # Clean up onboard auto-names: strip "Onboard: " prefix, file extensions, hash chars
    _raw_role = _re.sub(r'^Onboard:\s*', '', _raw_role)
    _raw_role = _re.sub(r'\.docx$', '', _raw_role, flags=_re.IGNORECASE)
    _raw_role = _re.sub(r'[#]+', '', _raw_role)
    _role_part = _re.sub(r'[^a-zA-Z0-9_\-]', '_', _raw_role.strip())[:40].strip('_')
    _date_part = _date.today().strftime("%Y%m%d")
    filename = f"{_name_part}_{_role_part}_{_date_part}.docx"

    # If ?format=json, save to disk and return JSON metadata instead of file download
    if request.args.get("format") == "json":
        import os
        output_dir = os.path.join(os.getcwd(), "generated")
        os.makedirs(output_dir, exist_ok=True)
        output_path = os.path.join(output_dir, filename)
        with open(output_path, "wb") as f:
            f.write(output.read())
        return jsonify({"status": "ok", "output_path": output_path, "filename": filename})

    return send_file(
        output,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name=filename,
    )


# ---------------------------------------------------------------------------
# Resume Header
# ---------------------------------------------------------------------------

@bp.route("/api/resume/header", methods=["GET"])
def get_resume_header():
    """Get resume header info (name, credentials, contact details)."""
    row = db.query_one("SELECT * FROM resume_header LIMIT 1")
    return jsonify(row or {"error": "No header data found"})


@bp.route("/api/resume/header", methods=["PATCH"])
def update_resume_header():
    """Update resume header fields (upsert — creates if not exists)."""
    data = request.get_json(force=True)
    allowed = [
        "full_name", "credentials", "location", "location_note",
        "email", "phone", "linkedin_url", "website_url", "calendly_url",
    ]
    existing = db.query_one("SELECT id FROM resume_header LIMIT 1")
    if existing:
        sets, params = [], []
        for key in allowed:
            if key in data:
                sets.append(f"{key} = %s")
                params.append(data[key])
        if not sets:
            return jsonify({"error": "No valid fields to update"}), 400
        params.append(existing["id"])
        row = db.execute_returning(
            f"UPDATE resume_header SET {', '.join(sets)} WHERE id = %s RETURNING *",
            params,
        )
    else:
        row = db.execute_returning(
            """
            INSERT INTO resume_header (full_name, credentials, location, location_note,
                email, phone, linkedin_url, website_url, calendly_url)
            VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s)
            RETURNING *
            """,
            (
                data.get("full_name", ""), data.get("credentials"),
                data.get("location"), data.get("location_note"),
                data.get("email"), data.get("phone"),
                data.get("linkedin_url"), data.get("website_url"),
                data.get("calendly_url"),
            ),
        )
    return jsonify(row), 200


# ---------------------------------------------------------------------------
# Education
# ---------------------------------------------------------------------------

@bp.route("/api/education", methods=["GET"])
def get_education():
    """Get education entries."""
    rows = db.query("SELECT * FROM education ORDER BY sort_order")
    return jsonify({"education": rows, "count": len(rows)})


@bp.route("/api/education", methods=["POST"])
def create_education():
    """Add a new education entry."""
    data = request.get_json(force=True)
    if not data.get("degree") or not data.get("institution"):
        return jsonify({"error": "degree and institution are required"}), 400

    row = db.execute_returning(
        """
        INSERT INTO education (degree, field, institution, location, type, sort_order)
        VALUES (%s,%s,%s,%s,%s,%s)
        RETURNING *
        """,
        (
            data["degree"], data.get("field"), data["institution"],
            data.get("location"), data.get("type", "degree"),
            data.get("sort_order", 0),
        ),
    )
    return jsonify(row), 201


@bp.route("/api/education/<int:edu_id>", methods=["PATCH"])
def update_education(edu_id):
    """Update an education entry."""
    data = request.get_json(force=True)
    allowed = ["degree", "field", "institution", "location", "type", "sort_order"]
    sets, params = [], []
    for key in allowed:
        if key in data:
            sets.append(f"{key} = %s")
            params.append(data[key])
    if not sets:
        return jsonify({"error": "No valid fields to update"}), 400

    params.append(edu_id)
    row = db.execute_returning(
        f"UPDATE education SET {', '.join(sets)} WHERE id = %s RETURNING *",
        params,
    )
    if not row:
        return jsonify({"error": "Not found"}), 404
    return jsonify(row), 200


@bp.route("/api/education/<int:edu_id>", methods=["DELETE"])
def delete_education(edu_id):
    """Delete an education entry."""
    count = db.execute("DELETE FROM education WHERE id = %s", (edu_id,))
    if count == 0:
        return jsonify({"error": "Not found"}), 404
    return jsonify({"deleted": edu_id}), 200


# ---------------------------------------------------------------------------
# Certifications
# ---------------------------------------------------------------------------

@bp.route("/api/certifications", methods=["GET"])
def get_certifications():
    """Get certification entries."""
    rows = db.query("SELECT * FROM certifications WHERE is_active = TRUE ORDER BY sort_order")
    return jsonify({"certifications": rows, "count": len(rows)})


@bp.route("/api/certifications", methods=["POST"])
def create_certification():
    """Add a new certification."""
    data = request.get_json(force=True)
    if not data.get("name"):
        return jsonify({"error": "name is required"}), 400

    row = db.execute_returning(
        """
        INSERT INTO certifications (name, issuer, is_active, sort_order)
        VALUES (%s,%s,%s,%s)
        RETURNING *
        """,
        (data["name"], data.get("issuer"), data.get("is_active", True), data.get("sort_order", 0)),
    )
    return jsonify(row), 201


@bp.route("/api/certifications/<int:cert_id>", methods=["PATCH"])
def update_certification(cert_id):
    """Update a certification."""
    data = request.get_json(force=True)
    allowed = ["name", "issuer", "is_active", "sort_order"]
    sets, params = [], []
    for key in allowed:
        if key in data:
            sets.append(f"{key} = %s")
            params.append(data[key])
    if not sets:
        return jsonify({"error": "No valid fields to update"}), 400

    params.append(cert_id)
    row = db.execute_returning(
        f"UPDATE certifications SET {', '.join(sets)} WHERE id = %s RETURNING *",
        params,
    )
    if not row:
        return jsonify({"error": "Not found"}), 404
    return jsonify(row), 200


@bp.route("/api/certifications/<int:cert_id>", methods=["DELETE"])
def delete_certification(cert_id):
    """Delete a certification."""
    count = db.execute("DELETE FROM certifications WHERE id = %s", (cert_id,))
    if count == 0:
        return jsonify({"error": "Not found"}), 404
    return jsonify({"deleted": cert_id}), 200


# ---------------------------------------------------------------------------
# Languages
# ---------------------------------------------------------------------------

@bp.route("/api/languages", methods=["GET"])
def get_languages():
    """Get language entries."""
    rows = db.query("SELECT * FROM languages ORDER BY sort_order")
    return jsonify({"languages": rows, "count": len(rows)})


@bp.route("/api/languages", methods=["POST"])
def create_language():
    """Add a new language."""
    data = request.get_json(force=True)
    if not data.get("name"):
        return jsonify({"error": "name is required"}), 400

    valid_proficiencies = ["native", "fluent", "professional", "conversational", "basic"]
    proficiency = data.get("proficiency", "conversational")
    if proficiency not in valid_proficiencies:
        return jsonify({"error": f"proficiency must be one of: {', '.join(valid_proficiencies)}"}), 400

    row = db.execute_returning(
        """
        INSERT INTO languages (name, proficiency, sort_order)
        VALUES (%s,%s,%s)
        RETURNING *
        """,
        (data["name"], proficiency, data.get("sort_order", 0)),
    )
    return jsonify(row), 201


@bp.route("/api/languages/<int:lang_id>", methods=["PATCH"])
def update_language(lang_id):
    """Update a language."""
    data = request.get_json(force=True)
    allowed = ["name", "proficiency", "sort_order"]
    sets, params = [], []

    valid_proficiencies = ["native", "fluent", "professional", "conversational", "basic"]
    if "proficiency" in data and data["proficiency"] not in valid_proficiencies:
        return jsonify({"error": f"proficiency must be one of: {', '.join(valid_proficiencies)}"}), 400

    for key in allowed:
        if key in data:
            sets.append(f"{key} = %s")
            params.append(data[key])
    if not sets:
        return jsonify({"error": "No valid fields to update"}), 400

    params.append(lang_id)
    row = db.execute_returning(
        f"UPDATE languages SET {', '.join(sets)} WHERE id = %s RETURNING *",
        params,
    )
    if not row:
        return jsonify({"error": "Not found"}), 404
    return jsonify(row), 200


@bp.route("/api/languages/<int:lang_id>", methods=["DELETE"])
def delete_language(lang_id):
    """Delete a language."""
    count = db.execute("DELETE FROM languages WHERE id = %s", (lang_id,))
    if count == 0:
        return jsonify({"error": "Not found"}), 404
    return jsonify({"deleted": lang_id}), 200


# ---------------------------------------------------------------------------
# Summary Variants
# ---------------------------------------------------------------------------

@bp.route("/api/summary-variants", methods=["GET"])
def get_summary_variants():
    """Get summary variant entries."""
    rows = db.query("SELECT * FROM summary_variants ORDER BY sort_order")
    return jsonify({"summary_variants": rows, "count": len(rows)})


@bp.route("/api/summary-variants", methods=["POST"])
def create_summary_variant():
    """Add a new summary variant."""
    data = request.get_json(force=True)
    if not data.get("role_type") or not data.get("text"):
        return jsonify({"error": "role_type and text are required"}), 400

    row = db.execute_returning(
        """
        INSERT INTO summary_variants (role_type, text, sort_order)
        VALUES (%s,%s,%s)
        RETURNING *
        """,
        (data["role_type"], data["text"], data.get("sort_order", 0)),
    )
    return jsonify(row), 201


@bp.route("/api/summary-variants/<int:sv_id>", methods=["PATCH"])
def update_summary_variant(sv_id):
    """Update a summary variant."""
    data = request.get_json(force=True)
    allowed = ["role_type", "text", "sort_order"]
    sets, params = [], []
    for key in allowed:
        if key in data:
            sets.append(f"{key} = %s")
            params.append(data[key])
    if not sets:
        return jsonify({"error": "No valid fields to update"}), 400

    params.append(sv_id)
    row = db.execute_returning(
        f"UPDATE summary_variants SET {', '.join(sets)}, updated_at = NOW() WHERE id = %s RETURNING *",
        params,
    )
    if not row:
        return jsonify({"error": "Not found"}), 404
    return jsonify(row), 200


@bp.route("/api/summary-variants/<int:sv_id>", methods=["DELETE"])
def delete_summary_variant(sv_id):
    """Delete a summary variant."""
    count = db.execute("DELETE FROM summary_variants WHERE id = %s", (sv_id,))
    if count == 0:
        return jsonify({"error": "Not found"}), 404
    return jsonify({"deleted": sv_id}), 200


# ---------------------------------------------------------------------------
# Skills CRUD — lives in career.py (routes/career.py)
# ---------------------------------------------------------------------------

# ---------------------------------------------------------------------------
# References
# ---------------------------------------------------------------------------

@bp.route("/api/references", methods=["GET"])
def get_references():
    """Get professional references."""
    rows = db.query('SELECT * FROM "references" ORDER BY sort_order')
    return jsonify({"references": rows, "count": len(rows)})


@bp.route("/api/references", methods=["POST"])
def create_reference():
    """Add a new reference."""
    data = request.get_json(force=True)
    if not data.get("name"):
        return jsonify({"error": "name is required"}), 400

    row = db.execute_returning(
        """
        INSERT INTO "references" (name, title, company, relationship, email, phone,
                                  linkedin_url, notes, ok_to_contact, career_history_id, sort_order)
        VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)
        RETURNING *
        """,
        (
            data["name"], data.get("title"), data.get("company"),
            data.get("relationship"), data.get("email"), data.get("phone"),
            data.get("linkedin_url"), data.get("notes"),
            data.get("ok_to_contact", True),
            data.get("career_history_id"), data.get("sort_order", 0),
        ),
    )
    return jsonify(row), 201


@bp.route("/api/references/<int:ref_id>", methods=["PATCH"])
def update_reference(ref_id):
    """Update a reference."""
    data = request.get_json(force=True)
    allowed = ["name", "title", "company", "relationship", "email", "phone",
               "linkedin_url", "notes", "ok_to_contact", "career_history_id", "sort_order"]
    sets, params = [], []
    for key in allowed:
        if key in data:
            sets.append(f"{key} = %s")
            params.append(data[key])
    if not sets:
        return jsonify({"error": "No valid fields to update"}), 400

    params.append(ref_id)
    row = db.execute_returning(
        f'UPDATE "references" SET {", ".join(sets)} WHERE id = %s RETURNING *',
        params,
    )
    if not row:
        return jsonify({"error": "Not found"}), 404
    return jsonify(row), 200


@bp.route("/api/references/<int:ref_id>", methods=["DELETE"])
def delete_reference(ref_id):
    """Delete a reference."""
    count = db.execute('DELETE FROM "references" WHERE id = %s', (ref_id,))
    if count == 0:
        return jsonify({"error": "Not found"}), 404
    return jsonify({"deleted": ref_id}), 200


# ---------------------------------------------------------------------------
# Resume Templates
# ---------------------------------------------------------------------------

@bp.route("/api/resume/templates", methods=["GET"])
def list_templates():
    """List available resume templates (without blob data).

    Query params:
        dedup: if 'true', collapse exact duplicates (same content_hash) into one entry.
    """
    # Backfill content_hash for templates missing it
    missing_hash = db.query(
        "SELECT id, template_blob FROM resume_templates WHERE content_hash IS NULL AND template_blob IS NOT NULL"
    )
    for row in (missing_hash or []):
        h = compute_hash(bytes(row["template_blob"]))
        db.execute("UPDATE resume_templates SET content_hash = %s WHERE id = %s", (h, row["id"]))

    dedup = request.args.get("dedup", "true").lower() == "true"

    rows = db.query(
        """SELECT t.id, t.name, t.filename, t.description, t.is_active,
                  t.is_default, t.template_type, t.parser_version,
                  t.content_hash,
                  length(t.template_blob) as size_bytes,
                  t.preview_blob IS NOT NULL as has_thumbnail,
                  t.created_at,
                  COALESCE(rc.recipe_count, 0) as recipe_count
           FROM resume_templates t
           LEFT JOIN (
               SELECT template_id, COUNT(*) as recipe_count
               FROM resume_recipes
               GROUP BY template_id
           ) rc ON rc.template_id = t.id
           ORDER BY t.is_default DESC, t.name"""
    )

    if dedup and rows:
        seen_hashes = {}
        deduped = []
        duplicate_count = 0
        for r in rows:
            h = r.get("content_hash")
            if h and h in seen_hashes:
                duplicate_count += 1
                # Add to the kept entry's duplicate list
                seen_hashes[h]["duplicate_ids"] = seen_hashes[h].get("duplicate_ids", []) + [r["id"]]
                continue
            if h:
                seen_hashes[h] = r
            deduped.append(r)
        return jsonify({"templates": deduped, "count": len(deduped), "duplicates_hidden": duplicate_count})

    return jsonify({"templates": rows, "count": len(rows)})


@bp.route("/api/resume/templates/<int:template_id>", methods=["GET"])
def get_template_detail(template_id):
    """Get template detail including template_map slots."""
    row = db.query_one(
        """SELECT t.id, t.name, t.filename, t.description, t.is_active,
                  t.template_type, t.parser_version, t.template_map,
                  length(t.template_blob) as size_bytes,
                  t.preview_blob IS NOT NULL as has_thumbnail,
                  t.created_at,
                  COALESCE(rc.recipe_count, 0) as recipe_count
           FROM resume_templates t
           LEFT JOIN (
               SELECT template_id, COUNT(*) as recipe_count
               FROM resume_recipes
               GROUP BY template_id
           ) rc ON rc.template_id = t.id
           WHERE t.id = %s""",
        (template_id,),
    )
    if not row:
        return jsonify({"error": "Template not found"}), 404
    # Also fetch recipes referencing this template
    recipes = db.query(
        "SELECT id, name, description, is_active, created_at FROM resume_recipes WHERE template_id = %s ORDER BY name",
        (template_id,),
    )
    row["recipes"] = recipes
    return jsonify(row)


@bp.route("/api/resume/templates/<int:template_id>", methods=["DELETE"])
def delete_template(template_id):
    """Delete a template with optional recipe handling strategy.

    Body (optional JSON):
        reassign_to: dict mapping recipe_id (str) -> new_template_id (int)
        delete_recipes: bool — if true, delete all linked recipes
    If recipes exist and no strategy provided, returns 409 with affected list.
    """
    row = db.query_one("SELECT id, name, is_default FROM resume_templates WHERE id = %s", (template_id,))
    if not row:
        return jsonify({"error": "Template not found"}), 404

    if row.get("is_default"):
        return jsonify({"error": "Cannot delete the default template"}), 403

    # Check for linked recipes
    linked = db.query(
        "SELECT id, name FROM resume_recipes WHERE template_id = %s",
        (template_id,),
    ) or []

    if linked:
        data = request.get_json(silent=True) or {}
        reassign_to = data.get("reassign_to")
        delete_recipes = data.get("delete_recipes", False)

        if delete_recipes:
            db.execute("DELETE FROM resume_recipes WHERE template_id = %s", (template_id,))
        elif reassign_to:
            for rid_str, new_tid in reassign_to.items():
                rid = int(rid_str)
                target = db.query_one("SELECT id FROM resume_templates WHERE id = %s", (new_tid,))
                if not target:
                    return jsonify({"error": f"Target template {new_tid} not found for recipe {rid}"}), 400
                db.execute(
                    "UPDATE resume_recipes SET template_id = %s WHERE id = %s",
                    (new_tid, rid),
                )
        else:
            return jsonify({
                "error": f"Cannot delete: {len(linked)} recipe(s) reference this template. Provide reassign_to or delete_recipes.",
                "affected_recipes": [{"id": r["id"], "name": r["name"]} for r in linked],
            }), 409

    db.execute("DELETE FROM resume_templates WHERE id = %s", (template_id,))
    return jsonify({"deleted": template_id, "name": row["name"]})


@bp.route("/api/resume/templates/upload", methods=["POST"])
def upload_template():
    """Upload a .docx file as a new template. Detects duplicates (exact + near).

    Query params:
        force=true — skip duplicate check and insert anyway
        check_only=true — only check for duplicates, don't insert
    """
    if "file" not in request.files:
        return jsonify({"error": "No file uploaded. Use multipart form with 'file' field."}), 400
    f = request.files["file"]
    if not f.filename or not f.filename.lower().endswith(".docx"):
        return jsonify({"error": "Only .docx files are supported"}), 400
    blob = f.read()
    name = request.form.get("name", f.filename.rsplit(".", 1)[0])
    force = request.args.get("force", "").lower() == "true"
    check_only = request.args.get("check_only", "").lower() == "true"

    # Duplicate detection
    if not force:
        dedup = check_duplicates(blob, "full")
        if dedup["exact_match"]:
            return jsonify({
                "duplicate": True,
                "exact_match": dedup["exact_match"],
                "message": "This template already exists. Use ?force=true to upload anyway.",
                "analysis_mode": dedup["analysis_mode"],
            }), 409
        if dedup["near_matches"] and check_only:
            return jsonify({
                "duplicate": False,
                "near_matches": dedup["near_matches"],
                "content_hash": dedup["content_hash"],
                "message": "Similar templates found. Review matches or use ?force=true to upload.",
                "analysis_mode": dedup["analysis_mode"],
            }), 200
        if check_only:
            return jsonify({
                "duplicate": False,
                "near_matches": [],
                "content_hash": dedup["content_hash"],
                "message": "No duplicates found.",
                "analysis_mode": dedup["analysis_mode"],
            }), 200

    content_hash = compute_hash(blob)
    # Force upload: set hash to NULL to bypass unique constraint (intentional duplicate)
    store_hash = None if force else content_hash
    row = db.query_one(
        """INSERT INTO resume_templates (name, filename, template_blob, template_type, is_active, content_hash)
           VALUES (%s, %s, %s, 'full', true, %s)
           RETURNING id, name, filename, created_at""",
        (name, f.filename, blob, store_hash),
    )
    resp = dict(row)
    if force:
        resp["forced"] = True
    return jsonify(resp), 201


@bp.route("/api/resume/templates/<int:template_id>/download", methods=["GET"])
def download_template(template_id):
    """Download a resume template .docx file."""
    row = db.query_one(
        "SELECT filename, template_blob FROM resume_templates WHERE id = %s", (template_id,)
    )
    if not row:
        return jsonify({"error": "Template not found"}), 404
    return send_file(
        io.BytesIO(row["template_blob"]),
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name=row["filename"],
    )


# ---------------------------------------------------------------------------
# Resume Versions & Specs
# ---------------------------------------------------------------------------

@bp.route("/api/resume/versions", methods=["GET"])
def list_resume_versions():
    """List resume versions with their specs."""
    rows = db.query(
        """SELECT id, version, variant, is_current, spec IS NOT NULL as has_spec,
                  docx_path, pdf_path, summary, target_role_type, created_at
           FROM resume_versions ORDER BY created_at DESC"""
    )
    return jsonify({"versions": rows, "count": len(rows)})


@bp.route("/api/resume/versions/<int:version_id>/spec", methods=["GET"])
def get_resume_spec(version_id):
    """Get the full spec for a resume version."""
    row = db.query_one("SELECT * FROM resume_versions WHERE id = %s", (version_id,))
    if not row:
        return jsonify({"error": "Version not found"}), 404
    return jsonify(row)


# ---------------------------------------------------------------------------
# Resume Data (full reconstruction data for a version)
# ---------------------------------------------------------------------------

@bp.route("/api/resume/data", methods=["GET"])
def get_resume_data():
    """Get all data needed to reconstruct a resume.

    Query params:
        version: resume version (default: v32)
        variant: resume variant (default: base)
        format: 'full' returns everything, 'spec_only' returns just the spec

    Returns header, education, certs, career history with bullets, spec, etc.
    """
    version = request.args.get("version", "v32")
    variant = request.args.get("variant", "base")
    fmt = request.args.get("format", "full")

    # Get the resume spec
    rv = db.query_one(
        "SELECT * FROM resume_versions WHERE version = %s AND variant = %s AND spec IS NOT NULL",
        (version, variant),
    )
    if not rv:
        return jsonify({"error": f"No spec found for {version}/{variant}"}), 404

    if fmt == "spec_only":
        return jsonify(rv)

    spec = rv["spec"] if isinstance(rv["spec"], dict) else json.loads(rv["spec"])

    # Get header
    header = db.query_one("SELECT * FROM resume_header LIMIT 1")

    # Get education
    education = db.query("SELECT * FROM education ORDER BY sort_order")

    # Get certifications
    certifications = db.query("SELECT * FROM certifications WHERE is_active = TRUE ORDER BY sort_order")

    # Get experience sections with bullets
    experience = []
    for employer_name in spec.get("experience_employers", []):
        ch = db.query_one(
            "SELECT * FROM career_history WHERE employer ILIKE %s",
            (f"%{employer_name}%",),
        )
        if ch:
            bullets = db.query(
                """SELECT id, text, type, tags, role_suitability, industry_suitability, metrics_json
                   FROM bullets WHERE career_history_id = %s ORDER BY id""",
                (ch["id"],),
            )
            ch["bullets"] = bullets
            experience.append(ch)

    return jsonify({
        "version": version,
        "variant": variant,
        "spec": spec,
        "header": header,
        "education": education,
        "certifications": certifications,
        "experience": experience,
        "template_available": db.query_one(
            "SELECT id, name FROM resume_templates LIMIT 1"
        ),
    })


# ---------------------------------------------------------------------------
# Resume Generation
# ---------------------------------------------------------------------------

PLACEHOLDER_RE = re.compile(r"\{\{([A-Z0-9_]+)\}\}")

BOLD_SEPARATORS = {
    "highlight": ": ",
    "job_bullet": ": ",
    "education": " | ",
    "certification": " | ",
    "additional_exp": " | ",
    "ref_link": " | ",
    "job_header": ", ",
}


def _flatten_v2_for_template(resolved: dict) -> dict:
    """Convert v2 resolved content to flat placeholder->text map for .docx filling."""
    flat = {}

    # Header
    if "header" in resolved and isinstance(resolved["header"], dict):
        h = resolved["header"]
        flat["HEADER_NAME"] = f"{h.get('full_name', '')}, {h.get('credentials', '')}"
        parts = [h.get("location", "")]
        if h.get("location_note"):
            parts[0] += f" ({h['location_note']})"
        parts.extend([h.get("email", ""), h.get("phone", "")])
        if h.get("linkedin_url"):
            parts.append(h["linkedin_url"])
        flat["HEADER_CONTACT"] = " \u2022 ".join(p for p in parts if p)

    # Single text slots
    for key in ("headline", "summary"):
        if key in resolved:
            flat[key.upper()] = resolved[key] if isinstance(resolved[key], str) else str(resolved[key])

    # Highlights
    if "highlights" in resolved:
        for i, h in enumerate(resolved["highlights"]):
            flat[f"HIGHLIGHT_{i+1}"] = h if isinstance(h, str) else str(h)

    # Experience
    if "experience" in resolved:
        for i, job in enumerate(resolved["experience"]):
            n = i + 1
            if isinstance(job, dict):
                flat[f"JOB_{n}_HEADER"] = job.get("employer", "")
                flat[f"JOB_{n}_TITLE"] = job.get("title", "")
                dates = f"{job.get('start_date', '')} - {job.get('end_date', 'Present')}"
                flat[f"JOB_{n}_DATES"] = dates
                flat[f"JOB_{n}_INTRO"] = job.get("synopsis", "")
                for j, bullet in enumerate(job.get("bullets", [])):
                    flat[f"JOB_{n}_BULLET_{j+1}"] = bullet if isinstance(bullet, str) else str(bullet)

    # Skills
    if "skills" in resolved:
        skills_list = resolved["skills"]
        if isinstance(skills_list, list):
            # Try to fill TECH_SKILLS and OTHER_SKILLS_N and KEYWORDS_N
            flat["TECH_SKILLS"] = ", ".join(str(s) for s in skills_list if s)
            for i, s in enumerate(skills_list):
                flat[f"OTHER_SKILLS_{i+1}"] = str(s) if s else ""
                flat[f"KEYWORDS_{i+1}"] = str(s) if s else ""

    # Education
    if "education" in resolved:
        for i, edu in enumerate(resolved["education"]):
            flat[f"EDUCATION_{i+1}"] = edu if isinstance(edu, str) else str(edu)

    # Certifications
    if "certifications" in resolved:
        for i, cert in enumerate(resolved["certifications"]):
            flat[f"CERT_{i+1}"] = cert if isinstance(cert, str) else str(cert)

    # Additional experience
    if "additional_experience" in resolved:
        for i, ae in enumerate(resolved["additional_experience"]):
            flat[f"ADDL_EXP_{i+1}"] = ae if isinstance(ae, str) else str(ae)

    # Custom catch-all
    if "custom" in resolved:
        for k, v in resolved["custom"].items():
            flat[k] = v if isinstance(v, str) else str(v)

    return flat


def _fill_simple(paragraph, text):
    """Replace placeholder text, preserving first run formatting."""
    from docx.oxml.ns import qn
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.text = text
    # Remove hyperlink elements that carry leftover text from the original
    for hl in paragraph._element.findall(qn('w:hyperlink')):
        paragraph._element.remove(hl)


def _fill_bold_label(paragraph, text, separator=": "):
    """Replace with bold label + non-bold body, split at separator."""
    idx = text.find(separator)
    if idx < 0 or not paragraph.runs:
        _fill_simple(paragraph, text)
        return
    paragraph.runs[0].text = text[:idx]
    paragraph.runs[0].bold = True
    if len(paragraph.runs) > 1:
        paragraph.runs[1].text = text[idx:]
        paragraph.runs[1].bold = None  # inherit = not bold
        for run in paragraph.runs[2:]:
            run.text = ""
            run.bold = None
    else:
        paragraph.runs[0].text = text


def _build_content_map(spec, header, education, certifications, career, template_map):
    """Build placeholder_name -> content_text mapping."""
    content = {}

    # Seed from template_map original_text
    for slot in template_map.get("slots", []):
        placeholder = slot.get("placeholder")
        original = slot.get("original_text")
        if placeholder and original:
            content[placeholder] = original

    # Overlay with spec-derived content
    if header:
        content["HEADER_NAME"] = f"{header['full_name']}, {header['credentials']}"
        parts = [header["location"]]
        if header.get("location_note"):
            parts[0] += f" ({header['location_note']})"
        parts.append(header["email"])
        parts.append(header["phone"])
        if header.get("linkedin_url"):
            parts.append(header["linkedin_url"])
        content["HEADER_CONTACT"] = " \u2022 ".join(parts)

    if "headline" in spec:
        content["HEADLINE"] = spec["headline"]
    if "summary_text" in spec:
        content["SUMMARY"] = spec["summary_text"]

    for i, bullet in enumerate(spec.get("highlight_bullets", []), 1):
        content[f"HIGHLIGHT_{i}"] = bullet

    if "keywords" in spec:
        content["KEYWORDS"] = " | ".join(spec["keywords"])

    # Experience blocks
    employers = spec.get("experience_employers", [])
    exp_bullets = spec.get("experience_bullets", {})

    for job_n, emp_name in enumerate(employers, 1):
        emp_data = career.get(emp_name, {})
        bullets_raw = exp_bullets.get(emp_name, [])

        if emp_data.get("intro_text"):
            content[f"JOB_{job_n}_INTRO"] = emp_data["intro_text"]
        elif bullets_raw and len(bullets_raw[0]) > 200 and ": " not in bullets_raw[0][:80]:
            content[f"JOB_{job_n}_INTRO"] = bullets_raw[0]

        bullet_texts = []
        subtitle_texts = {content.get(f"JOB_{job_n}_SUBTITLE_1", ""),
                          content.get(f"JOB_{job_n}_SUBTITLE_2", "")}

        for b in bullets_raw:
            if b == content.get(f"JOB_{job_n}_INTRO"):
                continue
            if b in subtitle_texts:
                continue
            bullet_texts.append(b)

        for i, bt in enumerate(bullet_texts, 1):
            content[f"JOB_{job_n}_BULLET_{i}"] = bt

    if "executive_keywords" in spec:
        content["EXEC_KEYWORDS"] = " | ".join(spec["executive_keywords"])
    if "technical_keywords" in spec:
        content["TECH_KEYWORDS"] = " | ".join(spec["technical_keywords"])

    for i, ref_section in enumerate(spec.get("references", []), 1):
        for j, link in enumerate(ref_section.get("links", []), 1):
            content[f"REF_{i}_LINK_{j}"] = f"{link['text']} | {link['desc']}"

    return content


@bp.route("/api/resume/generate", methods=["POST"])
def generate_resume():
    """Generate a .docx resume from placeholder template + spec.

    POST body (JSON):
        version: resume version (default: v32)
        variant: resume variant (default: base)
        template_name: template to use (default: V32 Placeholder)
        overrides: optional dict of placeholder_name -> text overrides

    Returns: .docx file download
    """
    data = request.get_json(silent=True) or {}
    version = data.get("version", "v32")
    variant = data.get("variant", "base")
    template_name = data.get("template_name", "V32 Placeholder")
    overrides = data.get("overrides", {})

    # Load template
    tmpl = db.query_one(
        "SELECT template_blob, template_map FROM resume_templates WHERE name = %s",
        (template_name,),
    )
    if not tmpl:
        return jsonify({"error": f"Template '{template_name}' not found"}), 404

    template_blob = bytes(tmpl["template_blob"])
    template_map = tmpl["template_map"] or {}

    # Load spec
    rv = db.query_one(
        "SELECT spec FROM resume_versions WHERE version = %s AND variant = %s AND spec IS NOT NULL",
        (version, variant),
    )
    if not rv:
        return jsonify({"error": f"No spec found for {version}/{variant}"}), 404
    spec = rv["spec"] if isinstance(rv["spec"], dict) else json.loads(rv["spec"])

    # Load supporting data
    header = db.query_one("SELECT * FROM resume_header LIMIT 1")
    education = db.query("SELECT * FROM education ORDER BY sort_order")
    certifications = db.query(
        "SELECT * FROM certifications WHERE is_active = TRUE ORDER BY sort_order"
    )

    employers = spec.get("experience_employers", [])
    career = {}
    for emp in employers:
        ch = db.query_one(
            "SELECT * FROM career_history WHERE employer ILIKE %s ORDER BY start_date DESC LIMIT 1",
            (f"%{emp}%",),
        )
        if ch:
            career[emp] = ch

    # Build content map
    base_content_map = _build_content_map(spec, header, education, certifications, career, template_map)
    base_content_map.update(overrides)

    # --- AI routing: enhance resume content ---
    ai_context = {
        "task_type": "generate_resume",
        "content_map": base_content_map,
        "version": version,
        "variant": variant,
    }

    def _python_resume_content(ctx):
        return {"content_map": ctx["content_map"]}

    def _ai_resume_content(ctx):
        from ai_providers import get_provider
        provider = get_provider()
        enhanced = dict(ctx["content_map"])
        # AI-polish summary slots
        for key, val in enhanced.items():
            if "SUMMARY" in key and val:
                prompt = (
                    f"Polish this resume summary for impact and clarity. "
                    f"Keep the same facts and metrics. No buzzwords. No em dashes. "
                    f"Under the same word count:\n\n{val}"
                )
                try:
                    enhanced[key] = provider.generate(prompt)
                except Exception:
                    pass
        return {"content_map": enhanced}

    gen_result = route_inference(
        task="generate_resume",
        context=ai_context,
        python_fallback=_python_resume_content,
        ai_handler=_ai_resume_content,
    )
    content_map = gen_result["content_map"]

    # Build slot info lookup
    slot_info = {}
    for slot in template_map.get("slots", []):
        if slot.get("placeholder"):
            slot_info[slot["placeholder"]] = {
                "slot_type": slot.get("slot_type", ""),
                "formatting": slot.get("formatting", {}),
            }

    # Fill template
    doc = Document(io.BytesIO(template_blob))
    filled = 0
    for para in doc.paragraphs:
        match = PLACEHOLDER_RE.search(para.text)
        if not match:
            continue
        placeholder = match.group(1)
        if placeholder not in content_map:
            _fill_simple(para, "")
            continue

        text = content_map[placeholder]
        info = slot_info.get(placeholder, {})
        slot_type = info.get("slot_type", "")
        formatting = info.get("formatting", {})

        if formatting.get("bold_label") and slot_type in BOLD_SEPARATORS:
            _fill_bold_label(para, text, BOLD_SEPARATORS[slot_type])
        else:
            _fill_simple(para, text)
        filled += 1

    # Return as downloadable .docx
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)

    # Build filename: {Name}_{TargetRole}_{Date}.docx
    from datetime import date as _date
    _name_part = (content_map.get("HEADER_NAME") or "Resume").strip().replace(" ", "")
    _role_part = (variant or version or "General").strip().replace(" ", "_")[:40]
    _date_part = _date.today().strftime("%Y%m%d")
    filename = f"{_name_part}_{_role_part}_{_date_part}.docx"
    return send_file(
        output,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name=filename,
    )


# ---------------------------------------------------------------------------
# S10: Resume content validation — checks bullet quality, placeholders, voice
# ---------------------------------------------------------------------------

# Patterns that flag a bullet as lacking a metric
_METRIC_PATTERNS = re.compile(
    r"\d+%|\$[\d,]+|[\d,]+\s*(users|customers|engineers|employees|teams?|"
    r"services?|systems?|projects?|products?|repos?|pipelines?)|"
    r"\d+x|reduced|increased|improved|saved|grew|cut|boosted|delivered|"
    r"launched|scaled|migrated|built|led\s+\d",
    re.IGNORECASE,
)
_PLACEHOLDER_PATTERNS = re.compile(
    r"\[.*?\]|\bTBD\b|\bXXX\b|\bPLACEHOLDER\b|\bINSERT\b|\bTODO\b",
    re.IGNORECASE,
)
_REQUIRED_SECTIONS = {
    "summary", "experience", "skills", "education",
}
_BANNED_PHRASES = [
    "proven track record", "results-driven", "detail-oriented",
    "team player", "go-to", "utilize", "leverage", "synergy",
    "innovative", "dynamic", "passionate", "guru", "ninja", "rockstar",
]


@bp.route("/api/resume/recipes/<int:recipe_id>/validate-content", methods=["GET"])
def validate_recipe_content(recipe_id):
    """Deep content validation of a recipe: bullet quality, placeholders, voice.

    Checks:
    - All required sections present (summary, experience, skills, education)
    - No placeholder text ([...], TBD, XXX) in resolved bullet text
    - Bullets have measurable metrics (numbers, %, $, outcome verbs)
    - No banned voice patterns in bullet text

    Returns:
        {"valid": bool, "score": 0-100, "issues": [...], "stats": {...}}
    """
    row = db.query_one("SELECT * FROM resume_recipes WHERE id = %s", (recipe_id,))
    if not row:
        return jsonify({"error": f"Recipe id={recipe_id} not found"}), 404

    recipe_json = row.get("recipe") or {}
    if isinstance(recipe_json, str):
        recipe_json = json.loads(recipe_json)

    issues = []
    stats = {
        "bullets_checked": 0,
        "bullets_with_metrics": 0,
        "bullets_missing_metrics": 0,
        "placeholders_found": 0,
        "voice_violations": 0,
        "sections_present": [],
        "sections_missing": [],
    }

    # 1. Check required sections
    slot_keys_lower = {k.lower() for k in recipe_json.keys()}
    for section in _REQUIRED_SECTIONS:
        found = any(section in k for k in slot_keys_lower)
        if found:
            stats["sections_present"].append(section)
        else:
            stats["sections_missing"].append(section)
            issues.append({
                "type": "missing_section",
                "severity": "error",
                "message": f"Required section '{section}' not found in recipe slots",
            })

    # 2. Resolve bullet IDs and check content
    for slot_name, ref in recipe_json.items():
        if not isinstance(ref, dict):
            continue
        table = ref.get("table", "")
        ids = ref.get("ids", [])
        if table == "bullets" and ids:
            bullets = db.query(
                "SELECT id, text, metrics_json FROM bullets WHERE id = ANY(%s)",
                (ids,),
            )
            for b in bullets:
                text = b.get("text") or ""
                stats["bullets_checked"] += 1

                # Placeholder check
                if _PLACEHOLDER_PATTERNS.search(text):
                    stats["placeholders_found"] += 1
                    issues.append({
                        "type": "placeholder_text",
                        "severity": "error",
                        "bullet_id": b["id"],
                        "slot": slot_name,
                        "message": f"Bullet {b['id']} contains placeholder text",
                        "text_preview": text[:80],
                    })

                # Metric check
                has_metric = bool(_METRIC_PATTERNS.search(text)) or bool(b.get("metrics_json"))
                if has_metric:
                    stats["bullets_with_metrics"] += 1
                else:
                    stats["bullets_missing_metrics"] += 1
                    issues.append({
                        "type": "missing_metric",
                        "severity": "warning",
                        "bullet_id": b["id"],
                        "slot": slot_name,
                        "message": f"Bullet {b['id']} has no measurable metric or outcome",
                        "text_preview": text[:80],
                    })

                # Voice check
                text_lower = text.lower()
                for phrase in _BANNED_PHRASES:
                    if phrase in text_lower:
                        stats["voice_violations"] += 1
                        issues.append({
                            "type": "voice_violation",
                            "severity": "warning",
                            "bullet_id": b["id"],
                            "slot": slot_name,
                            "message": f"Banned phrase '{phrase}' in bullet {b['id']}",
                            "text_preview": text[:80],
                        })
                        break

    # 3. Score: start at 100, deduct per issue type
    errors = [i for i in issues if i["severity"] == "error"]
    warnings = [i for i in issues if i["severity"] == "warning"]
    score = max(0, 100 - (len(errors) * 15) - (len(warnings) * 5))

    return jsonify({
        "valid": len(errors) == 0,
        "score": score,
        "issue_count": {"errors": len(errors), "warnings": len(warnings)},
        "issues": issues,
        "stats": stats,
    }), 200


@bp.route("/api/resume/recipes/compare", methods=["POST"])
def compare_recipes():
    """Compare two recipes side-by-side: bullet diffs, skill changes, slot differences.

    Body (JSON):
        recipe_a_id (int): first recipe ID
        recipe_b_id (int): second recipe ID

    Returns:
        {"added": [...], "removed": [...], "shared": [...], "skill_diff": {...}}
    """
    data = request.get_json(force=True)
    id_a = data.get("recipe_a_id")
    id_b = data.get("recipe_b_id")
    if not id_a or not id_b:
        return jsonify({"error": "recipe_a_id and recipe_b_id are required"}), 400

    def _load_recipe(rid):
        r = db.query_one("SELECT id, name, recipe FROM resume_recipes WHERE id = %s", (rid,))
        if not r:
            return None, None
        rj = r["recipe"]
        if isinstance(rj, str):
            rj = json.loads(rj)
        return r, rj

    rec_a, json_a = _load_recipe(id_a)
    rec_b, json_b = _load_recipe(id_b)
    if not rec_a:
        return jsonify({"error": f"Recipe {id_a} not found"}), 404
    if not rec_b:
        return jsonify({"error": f"Recipe {id_b} not found"}), 404

    def _extract_ids(recipe_json, table_name):
        ids = set()
        for ref in recipe_json.values():
            if isinstance(ref, dict) and ref.get("table") == table_name:
                ids.update(ref.get("ids", []))
        return ids

    # Bullet diff
    bullets_a = _extract_ids(json_a, "bullets")
    bullets_b = _extract_ids(json_b, "bullets")
    added_ids = bullets_b - bullets_a
    removed_ids = bullets_a - bullets_b
    shared_ids = bullets_a & bullets_b

    def _fetch_bullets(ids):
        if not ids:
            return []
        return db.query(
            "SELECT b.id, b.text, b.tags, ch.employer FROM bullets b LEFT JOIN career_history ch ON ch.id = b.career_history_id WHERE b.id = ANY(%s)",
            (list(ids),),
        )

    # Skill diff
    skills_a = _extract_ids(json_a, "skills")
    skills_b = _extract_ids(json_b, "skills")

    def _fetch_skills(ids):
        if not ids:
            return []
        return db.query("SELECT id, name, category FROM skills WHERE id = ANY(%s)", (list(ids),))

    # Slot key diff
    slots_a = set(json_a.keys())
    slots_b = set(json_b.keys())

    return jsonify({
        "recipe_a": {"id": id_a, "name": rec_a["name"]},
        "recipe_b": {"id": id_b, "name": rec_b["name"]},
        "bullet_diff": {
            "added_in_b": _fetch_bullets(added_ids),
            "removed_in_b": _fetch_bullets(removed_ids),
            "shared_count": len(shared_ids),
        },
        "skill_diff": {
            "added_in_b": _fetch_skills(skills_b - skills_a),
            "removed_in_b": _fetch_skills(skills_a - skills_b),
            "shared_count": len(skills_a & skills_b),
        },
        "slot_diff": {
            "only_in_a": sorted(slots_a - slots_b),
            "only_in_b": sorted(slots_b - slots_a),
            "shared": sorted(slots_a & slots_b),
        },
    }), 200


# ---------------------------------------------------------------------------
# Template Thumbnail
# ---------------------------------------------------------------------------

def _generate_template_thumbnail(template_blob, template_name):
    """Generate a 200x260 PNG thumbnail from a .docx template blob.

    Uses LibreOffice headless to render the actual document as a real preview.
    Falls back to a simple placeholder image if LibreOffice is unavailable.
    Returns PNG bytes.
    """
    from utils.thumbnail import generate_thumbnail

    # Try real rendering first
    result = generate_thumbnail(template_blob)
    if result:
        return result

    # Fallback: simple placeholder with template name
    WIDTH, HEIGHT = 200, 260
    img = Image.new("RGB", (WIDTH, HEIGHT), (240, 240, 240))
    draw = ImageDraw.Draw(img)
    try:
        font = ImageFont.truetype("arial.ttf", 12)
    except (OSError, IOError):
        font = ImageFont.load_default()
    draw.text((10, 120), template_name[:30], fill="black", font=font)
    draw.text((10, 140), "(preview unavailable)", fill="gray", font=font)
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()


@bp.route("/api/resume/templates/<int:template_id>/thumbnail", methods=["GET"])
def get_template_thumbnail(template_id):
    """Return a PNG thumbnail preview of a resume template.

    If a cached preview_blob exists, returns it directly.
    Otherwise generates a layout diagram from the template .docx,
    caches it in preview_blob, and returns the PNG.
    """
    row = db.query_one(
        "SELECT id, name, preview_blob, template_blob FROM resume_templates WHERE id = %s",
        (template_id,),
    )
    if not row:
        return jsonify({"error": "Template not found"}), 404

    # Return cached thumbnail if available
    if row.get("preview_blob"):
        preview_bytes = bytes(row["preview_blob"])
        return Response(preview_bytes, mimetype="image/png")

    # Generate thumbnail from template_blob
    if not row.get("template_blob"):
        return jsonify({"error": "Template has no .docx blob to generate thumbnail from"}), 404

    template_blob = bytes(row["template_blob"])
    png_bytes = _generate_template_thumbnail(template_blob, row.get("name", "Template"))

    # Cache in DB
    try:
        db.execute(
            "UPDATE resume_templates SET preview_blob = %s WHERE id = %s",
            (png_bytes, template_id),
        )
    except Exception as e:
        logger.warning("Failed to cache template thumbnail: %s", e)

    return Response(png_bytes, mimetype="image/png")


# ---------------------------------------------------------------------------
# ATS Score for Recipes (Phase 4)
# ---------------------------------------------------------------------------

# Stopwords for keyword extraction (shared with resume_tailoring)
_ATS_STOPWORDS = frozenset(
    "a an the and or but in on at to for of is it by with as from that this "
    "be are was were been have has had do does did will would shall should may "
    "might can could not no nor so if then than too very also about above after "
    "again all am any because before between both but each few further get got "
    "here how into just more most no only other our out over own same she some "
    "such their them there these they through under until up us we what when "
    "where which while who whom why you your able must need per via etc".split()
)


def _ats_extract_keywords(text: str) -> list[dict]:
    """Extract keywords from text, returning [{keyword, count}] sorted by count desc."""
    words = re.findall(r"[a-zA-Z][a-zA-Z+#.\-]{1,}", text.lower())
    freq: dict[str, int] = {}
    for w in words:
        w_clean = w.strip(".-")
        if w_clean and w_clean not in _ATS_STOPWORDS and len(w_clean) > 2:
            freq[w_clean] = freq.get(w_clean, 0) + 1
    return sorted(
        [{"keyword": k, "count": v} for k, v in freq.items()],
        key=lambda x: x["count"],
        reverse=True,
    )


def _resolved_to_text(resolved: dict) -> str:
    """Convert a resolved recipe dict to plain text by joining all string values and list items."""
    parts: list[str] = []

    def _collect(obj):
        if isinstance(obj, str):
            parts.append(obj)
        elif isinstance(obj, list):
            for item in obj:
                _collect(item)
        elif isinstance(obj, dict):
            for v in obj.values():
                _collect(v)

    _collect(resolved)
    return " ".join(parts)


@bp.route("/api/resume/recipes/<int:recipe_id>/ats-score", methods=["POST"])
def recipe_ats_score(recipe_id):
    """Score a recipe-based resume against a JD for ATS compatibility.

    Body (JSON):
        jd_text: optional job description text
        application_id: optional — pulls JD from applications table
    """
    row = db.query_one("SELECT * FROM resume_recipes WHERE id = %s", (recipe_id,))
    if not row:
        return jsonify({"error": f"Recipe id={recipe_id} not found"}), 404

    data = request.get_json(force=True) if request.is_json else {}
    jd_text = data.get("jd_text")
    application_id = data.get("application_id")

    # If application_id provided and no jd_text, fetch JD from applications
    if application_id and not jd_text:
        app_row = db.query_one(
            "SELECT jd_text FROM applications WHERE id = %s", (application_id,)
        )
        if app_row and app_row.get("jd_text"):
            jd_text = app_row["jd_text"]

    # If still no JD, fall back to target_roles from settings preferences
    if not jd_text:
        settings_row = db.query_one(
            "SELECT preferences FROM settings WHERE id = 1"
        )
        if settings_row and settings_row.get("preferences"):
            prefs = settings_row["preferences"]
            if isinstance(prefs, str):
                prefs = json.loads(prefs)
            target_roles = prefs.get("target_roles", [])
            if target_roles:
                jd_text = " ".join(target_roles)

    if not jd_text:
        return jsonify({"error": "No JD text available. Provide jd_text, application_id, or set target_roles in settings."}), 400

    # Resolve recipe to full text
    from mcp_tools_resume_gen import _resolve_recipe_db

    recipe_json = row["recipe"]
    if isinstance(recipe_json, str):
        recipe_json = json.loads(recipe_json)
    recipe_version = row.get("recipe_version", 1)
    resolved = _resolve_recipe_db(recipe_json, recipe_version=recipe_version)
    resume_text = _resolved_to_text(resolved)

    if not resume_text.strip():
        return jsonify({"error": "Recipe resolved to empty text"}), 400

    # Extract keywords from JD
    jd_keywords = _ats_extract_keywords(jd_text)
    resume_lower = resume_text.lower()

    # Check each JD keyword against resume
    keyword_matches: dict[str, bool] = {}
    found_count = 0
    missing_keywords: list[str] = []
    for kw in jd_keywords[:50]:
        word = kw["keyword"]
        matched = bool(re.search(r"\b" + re.escape(word) + r"\b", resume_lower))
        keyword_matches[word] = matched
        if matched:
            found_count += 1
        else:
            missing_keywords.append(word)

    total_checked = len(keyword_matches)
    match_percentage = round((found_count / total_checked * 100), 1) if total_checked else 0

    # Format score: recipe text is plain text (no HTML), so 100
    format_score = 100
    formatting_flags: list[str] = []
    if "<table" in resume_lower:
        format_score = 60
        formatting_flags.append("Contains HTML tables")
    if "<img" in resume_lower:
        format_score = 60
        formatting_flags.append("Contains images")

    ats_score_val = round(match_percentage * 0.8 + format_score * 0.2)
    ats_score_val = min(ats_score_val, 100)

    python_result = {
        "ats_score": ats_score_val,
        "keyword_matches": keyword_matches,
        "match_percentage": match_percentage,
        "keywords_found": found_count,
        "keywords_checked": total_checked,
        "formatting_flags": formatting_flags,
        "missing_keywords": missing_keywords,
        "analysis_mode": "rule-based",
    }

    # AI enhancement via route_inference
    ai_context = {
        "resume_text": resume_text[:3000],
        "jd_text": jd_text[:3000],
        "python_result": python_result,
    }

    def _python_fallback(ctx):
        return ctx["python_result"]

    def _ai_handler(ctx):
        from ai_providers import get_provider
        provider = get_provider()
        prompt = f"""Analyze this resume against the job description for ATS compatibility.

Return a JSON object with these keys:
- suggestions: array of 3-5 actionable improvements to boost the ATS score

RESUME (first 3000 chars):
{ctx["resume_text"]}

JOB DESCRIPTION (first 3000 chars):
{ctx["jd_text"]}

Current keyword match: {ctx["python_result"]["match_percentage"]}%
Missing keywords: {", ".join(ctx["python_result"]["missing_keywords"][:15])}"""
        ai_result = provider.generate(prompt, response_format="json")
        merged = {**ctx["python_result"], "analysis_mode": "ai-enhanced"}
        if isinstance(ai_result, dict) and ai_result.get("suggestions"):
            merged["suggestions"] = ai_result["suggestions"]
        return merged

    result = route_inference("recipe_ats_score", ai_context, _python_fallback, _ai_handler)
    return jsonify(result), 200


# ---------------------------------------------------------------------------
# AI Review (generic quality + target-role fit)
# ---------------------------------------------------------------------------

_METRICS_RE = re.compile(
    r"\$[\d,.]+[MmKkBb]?"
    r"|[\d,.]+%"
    r"|[\d,.]+[Xx]"
    r"|\d+\+?\s*(?:users|employees|team|members|clients|customers|reports|engineers|people)",
    re.IGNORECASE,
)


@bp.route("/api/resume/recipes/<int:recipe_id>/ai-review", methods=["POST"])
def recipe_ai_review(recipe_id):
    """Two-layer AI review: generic quality + target-role fit scoring."""
    from mcp_tools_resume_gen import _resolve_recipe_db

    recipe_row = db.query_one("SELECT * FROM resume_recipes WHERE id = %s", (recipe_id,))
    if not recipe_row:
        return jsonify({"error": f"Recipe id={recipe_id} not found"}), 404

    recipe_json = recipe_row["recipe"]
    if isinstance(recipe_json, str):
        recipe_json = json.loads(recipe_json)

    resolved = _resolve_recipe_db(recipe_json, recipe_version=recipe_row.get("recipe_version", 1))
    resume_text = _resolved_to_text(resolved)

    # --- Layer 1: Generic quality review (rule-based) ---
    feedback = []
    strengths = []
    score = 70

    # Bullet count per job
    experience = resolved.get("experience", [])
    if isinstance(experience, list):
        for i, job in enumerate(experience):
            job_label = f"Job {i + 1}"
            if isinstance(job, dict):
                employer = job.get("employer", job_label)
                job_label = employer if employer else job_label
                bullets = job.get("bullets", [])
            elif isinstance(job, list):
                bullets = job
            else:
                bullets = []
            n = len(bullets)
            if n < 3:
                feedback.append({"section": job_label, "issue": f"Only {n} bullets — aim for 3-6", "severity": "medium"})
                score -= 3
            elif n > 8:
                feedback.append({"section": job_label, "issue": f"{n} bullets may be too many — trim to 5-6", "severity": "low"})
                score -= 1
            else:
                strengths.append(f"Good bullet count ({n}) for {job_label}")
                score += 2

            # Metrics scan per job
            if bullets:
                bullet_texts = [b if isinstance(b, str) else b.get("text", str(b)) for b in bullets]
                metrics_count = sum(1 for b in bullet_texts if _METRICS_RE.search(b))
                pct = metrics_count / len(bullet_texts) if bullet_texts else 0
                if pct < 0.5:
                    feedback.append({
                        "section": job_label,
                        "issue": f"Only {metrics_count}/{len(bullet_texts)} bullets have metrics — aim for 50%+",
                        "severity": "high",
                    })
                    score -= 5
                else:
                    strengths.append(f"Strong metrics in {job_label} ({metrics_count}/{len(bullet_texts)})")
                    score += 2

    # Summary length
    summary = resolved.get("summary", "")
    if isinstance(summary, str):
        slen = len(summary)
        if slen < 50:
            feedback.append({"section": "summary", "issue": "Summary is too short — flesh it out", "severity": "medium"})
            score -= 3
        elif slen > 500:
            feedback.append({"section": "summary", "issue": "Summary is too long — trim to 2-3 sentences", "severity": "medium"})
            score -= 3
        elif slen > 0:
            strengths.append("Summary length is appropriate")
            score += 2

    # Total word count
    word_count = len(resume_text.split())
    if word_count > 800:
        feedback.append({"section": "overall", "issue": f"Resume is {word_count} words — consider trimming", "severity": "low"})
        score -= 1

    score = max(0, min(100, score))

    generic_result = {
        "score": score,
        "feedback": feedback,
        "strengths": strengths,
    }

    # --- Layer 2: Target role fit ---
    target_roles = []
    try:
        settings_row = db.query_one("SELECT preferences FROM settings WHERE id = 1")
        if settings_row and settings_row.get("preferences"):
            prefs = settings_row["preferences"]
            if isinstance(prefs, str):
                prefs = json.loads(prefs)
            target_roles = prefs.get("target_roles", [])
    except Exception:
        pass

    resume_words = set(resume_text.lower().split())
    target_role_results = []
    for role in target_roles:
        role_name = role if isinstance(role, str) else role.get("title", str(role))
        role_words = set(role_name.lower().split())
        overlap = role_words & resume_words
        role_score = min(100, 50 + len(overlap) * 15)
        target_role_results.append({
            "role": role_name,
            "score": role_score,
            "gaps": [],
            "suggestions": [],
        })

    python_result = {
        "generic": generic_result,
        "target_roles": target_role_results,
        "analysis_mode": "rule_based",
    }

    # AI enhancement
    ai_context = {
        "resume_text": resume_text[:4000],
        "target_roles": target_roles,
        "python_result": python_result,
    }

    def _python_fallback(ctx):
        return ctx["python_result"]

    def _ai_handler(ctx):
        from ai_providers import get_provider
        provider = get_provider()
        roles_str = ", ".join(
            r if isinstance(r, str) else r.get("title", str(r))
            for r in ctx["target_roles"]
        ) if ctx["target_roles"] else "general tech leadership"

        prompt = f"""Review this resume and provide structured feedback.

Return a JSON object with exactly these keys:
- "generic": object with "score" (0-100), "feedback" (array of objects with "section", "issue", "severity"), "strengths" (array of strings)
- "target_roles": array of objects with "role", "score" (0-100), "gaps" (array of strings), "suggestions" (array of strings)

Target roles to evaluate against: {roles_str}

RESUME TEXT (first 4000 chars):
{ctx["resume_text"]}

Be specific and actionable. Severity must be "high", "medium", or "low"."""

        try:
            ai_result = provider.generate(prompt, response_format="json")
            if isinstance(ai_result, dict):
                result = {"analysis_mode": "ai"}
                if "generic" in ai_result:
                    result["generic"] = ai_result["generic"]
                else:
                    result["generic"] = ctx["python_result"]["generic"]
                if "target_roles" in ai_result:
                    result["target_roles"] = ai_result["target_roles"]
                else:
                    result["target_roles"] = ctx["python_result"]["target_roles"]
                return result
        except Exception as exc:
            logger.warning("AI review failed, using rule-based: %s", exc)
        return ctx["python_result"]

    result = route_inference("recipe_ai_review", ai_context, _python_fallback, _ai_handler)
    return jsonify(result), 200


# ---------------------------------------------------------------------------
# AI Generate Slot
# ---------------------------------------------------------------------------

@bp.route("/api/resume/recipes/<int:recipe_id>/ai-generate-slot", methods=["POST"])
def ai_generate_slot(recipe_id):
    """Generate content suggestions for a recipe slot (bullet, summary, highlight, job_intro).

    Uses route_inference: AI when available, Python fallback otherwise.
    """
    row = db.query_one("SELECT * FROM resume_recipes WHERE id = %s", (recipe_id,))
    if not row:
        return jsonify({"error": f"Recipe id={recipe_id} not found"}), 404

    data = request.get_json(force=True) or {}
    slot_type = data.get("slot_type", "bullet")
    ctx = data.get("context", {})
    job_id = ctx.get("job_id")
    existing_bullets = ctx.get("existing_bullets", [])
    target_role = ctx.get("target_role", "")
    instructions = ctx.get("instructions", "")

    # Build context for route_inference
    ai_context = {
        "slot_type": slot_type,
        "job_id": job_id,
        "existing_bullets": existing_bullets,
        "target_role": target_role,
        "instructions": instructions,
        "recipe_id": recipe_id,
    }

    def _python_fallback(ctx):
        suggestions = []
        st = ctx["slot_type"]
        jid = ctx["job_id"]
        existing = set(ctx["existing_bullets"])

        if st == "bullet" and jid:
            rows = db.query(
                "SELECT id, text FROM bullets WHERE career_history_id = %s ORDER BY display_order",
                (jid,),
            )
            for r in rows:
                if r["text"] not in existing:
                    suggestions.append({
                        "text": r["text"],
                        "confidence": 0.7,
                        "source": "existing_bullet",
                        "bullet_id": r["id"],
                    })

        elif st == "summary":
            rows = db.query(
                "SELECT id, text, role_type FROM summary_variants ORDER BY id"
            )
            for r in rows:
                suggestions.append({
                    "text": r["text"],
                    "confidence": 0.7,
                    "source": f"summary_{r.get('role_type', 'base')}",
                })

        elif st == "highlight":
            rows = db.query(
                """SELECT b.id, b.text, ch.employer
                   FROM bullets b
                   JOIN career_history ch ON ch.id = b.career_history_id
                   WHERE b.text ~ %s
                   ORDER BY b.display_order""",
                (r"\$[\d,.]+|\d+%|\d+x",),
            )
            for r in rows:
                suggestions.append({
                    "text": r["text"],
                    "confidence": 0.7,
                    "source": f"metric_bullet_{r.get('employer', '')}",
                    "bullet_id": r["id"],
                })

        elif st == "job_intro" and jid:
            ch = db.query_one(
                "SELECT intro_text FROM career_history WHERE id = %s",
                (jid,),
            )
            if ch and ch.get("intro_text"):
                intro = ch["intro_text"]
                if isinstance(intro, dict):
                    intro = intro.get("text", str(intro))
                suggestions.append({
                    "text": str(intro),
                    "confidence": 0.8,
                    "source": "career_history_intro",
                })

        return {"suggestions": suggestions, "analysis_mode": "rule_based"}

    def _ai_handler(ctx):
        from ai_providers import get_provider
        provider = get_provider()

        st = ctx["slot_type"]
        jid = ctx["job_id"]

        # Build job context
        job_context = ""
        if jid:
            ch = db.query_one(
                "SELECT employer, title FROM career_history WHERE id = %s",
                (jid,),
            )
            if ch:
                job_context = f"Company: {ch['employer']}, Title: {ch['title']}"

        # Load voice rules
        voice_rows = db.query(
            "SELECT rule_text FROM voice_rules WHERE category = 'resume_rule' ORDER BY sort_order LIMIT 10"
        )
        voice_rules = "\n".join(r["rule_text"] for r in voice_rows)

        # Load banned words
        banned_rows = db.query(
            "SELECT rule_text FROM voice_rules WHERE category = 'banned_word' ORDER BY sort_order"
        )
        banned_words = [r["rule_text"].lower() for r in banned_rows]

        existing_str = "\n".join(f"- {b}" for b in ctx["existing_bullets"]) if ctx["existing_bullets"] else "None"
        target = ctx["target_role"] or "general tech leadership"
        instr = ctx["instructions"] or "No special instructions"

        prompt = f"""Generate 3-5 resume {st} suggestions for a {target} role.

{f"Job context: {job_context}" if job_context else ""}
Target role: {target}
Special instructions: {instr}

Existing content (do NOT duplicate):
{existing_str}

Voice rules to follow:
{voice_rules}

Each suggestion MUST include a concrete metric or measurable outcome.
Return a JSON object with key "suggestions" containing an array of objects with "text" and "confidence" (0.0-1.0).
"""
        try:
            ai_result = provider.generate(prompt, response_format="json")
            if isinstance(ai_result, dict) and "suggestions" in ai_result:
                filtered = []
                for s in ai_result["suggestions"]:
                    text = s.get("text", "")
                    text_lower = text.lower()
                    if not any(bw in text_lower for bw in banned_words):
                        filtered.append({
                            "text": text,
                            "confidence": s.get("confidence", 0.8),
                            "source": "ai_generated",
                        })
                return {"suggestions": filtered, "analysis_mode": "ai"}
        except Exception as exc:
            logger.warning("AI generate-slot failed, using fallback: %s", exc)

        return _python_fallback(ctx)

    result = route_inference("recipe_ai_generate_slot", ai_context, _python_fallback, _ai_handler)
    return jsonify(result), 200


# ---------------------------------------------------------------------------
# Best Picks — rank bullets & jobs by JD relevance
# ---------------------------------------------------------------------------

@bp.route("/api/resume/recipes/<int:recipe_id>/best-picks", methods=["POST"])
def recipe_best_picks(recipe_id):
    """Rank bullets and jobs by relevance to a JD. Input: {jd_text?, application_id?, limit: 10}."""
    row = db.query_one("SELECT id FROM resume_recipes WHERE id = %s", (recipe_id,))
    if not row:
        return jsonify({"error": f"Recipe id={recipe_id} not found"}), 404

    body = request.get_json(force=True) or {}
    jd_text = body.get("jd_text")
    application_id = body.get("application_id")
    limit = min(body.get("limit", 10), 50)

    if not jd_text and not application_id:
        return jsonify({"error": "One of jd_text or application_id is required"}), 400

    # Resolve JD text from application if needed
    if not jd_text and application_id:
        app_row = db.query_one("SELECT jd_text FROM applications WHERE id = %s", (application_id,))
        if not app_row or not app_row.get("jd_text"):
            return jsonify({"error": f"Application id={application_id} has no jd_text"}), 400
        jd_text = app_row["jd_text"]

    def _python_fallback(_ctx=None):
        # 1. Load all bullets with job info
        bullets_rows = db.query(
            "SELECT b.id, b.text, b.career_history_id, ch.employer, ch.title "
            "FROM bullets b JOIN career_history ch ON b.career_history_id = ch.id "
            "ORDER BY ch.start_date DESC NULLS LAST, b.display_order"
        )

        # 2. Load all jobs
        jobs_rows = db.query(
            "SELECT id, employer, title, start_date, end_date "
            "FROM career_history ORDER BY start_date DESC NULLS LAST"
        )

        # 3. Extract keywords from JD
        keywords = _ats_extract_keywords(jd_text)
        kw_set = {kw["keyword"].lower() for kw in keywords}
        if not kw_set:
            return {
                "ranked_bullets": [],
                "ranked_jobs": [],
                "suggested_skills": [],
                "analysis_mode": "rule_based",
            }

        # 4. Score each bullet by keyword overlap
        scored_bullets = []
        for b in bullets_rows:
            text_lower = b["text"].lower()
            matched = [kw for kw in kw_set if kw in text_lower]
            relevance = len(matched) / len(kw_set) if kw_set else 0
            scored_bullets.append({
                "bullet_id": b["id"],
                "text": b["text"],
                "relevance": round(relevance, 4),
                "job": f"{b['employer']} — {b['title']}",
                "career_history_id": b["career_history_id"],
                "matched_keywords": sorted(matched),
            })

        # 5. Rank bullets by relevance
        scored_bullets.sort(key=lambda x: x["relevance"], reverse=True)
        ranked_bullets = scored_bullets[:limit]

        # 6. Rank jobs by average bullet relevance
        from collections import defaultdict
        job_scores = defaultdict(list)
        for sb in scored_bullets:
            job_scores[sb["career_history_id"]].append(sb["relevance"])

        ranked_jobs = []
        for j in jobs_rows:
            scores = job_scores.get(j["id"], [])
            avg = sum(scores) / len(scores) if scores else 0
            ranked_jobs.append({
                "career_history_id": j["id"],
                "company": j["employer"],
                "title": j["title"],
                "relevance": round(avg, 4),
                "reason": f"{len(scores)} bullets, avg relevance {round(avg * 100, 1)}%",
            })
        ranked_jobs.sort(key=lambda x: x["relevance"], reverse=True)

        # 7. Suggest missing skills: JD keywords not in skills table
        skill_rows = db.query("SELECT LOWER(name) AS name FROM skills")
        skill_names = {r["name"] for r in skill_rows}
        # Also check what's in the bullet text
        all_bullet_text = " ".join(b["text"].lower() for b in bullets_rows)
        suggested_skills = [
            kw for kw in kw_set
            if kw not in skill_names and kw not in all_bullet_text
        ]
        suggested_skills.sort()

        return {
            "ranked_bullets": ranked_bullets,
            "ranked_jobs": ranked_jobs[:10],
            "suggested_skills": suggested_skills[:20],
            "analysis_mode": "rule_based",
        }

    def _ai_handler(ctx):
        from ai_providers import get_provider
        provider = get_provider()
        # Get top 20 bullets from fallback for AI to re-rank
        fallback = _python_fallback()
        if provider is None:
            return fallback
        top_bullets = fallback["ranked_bullets"][:20]

        if not top_bullets:
            return fallback

        bullets_text = "\n".join(
            f"[ID:{b['bullet_id']}] ({b['job']}) {b['text']}"
            for b in top_bullets
        )

        prompt = f"""You are a resume optimization expert. Given a job description and a list of resume bullets, rank them by semantic relevance to the JD.

JOB DESCRIPTION:
{jd_text[:3000]}

CANDIDATE BULLETS (top 20 by keyword match):
{bullets_text}

Return a JSON object with:
- "ranked_bullet_ids": array of bullet IDs in order of relevance (most relevant first)
- "job_reasons": object mapping career_history_id to a 1-sentence reason why that job is relevant
- "suggested_skills": array of skills mentioned in the JD that the candidate should add

Be precise and specific in your reasoning."""

        try:
            ai_result = provider.generate(prompt, response_format="json")
            if isinstance(ai_result, dict):
                # Re-order bullets by AI ranking
                id_order = ai_result.get("ranked_bullet_ids", [])
                bullet_map = {b["bullet_id"]: b for b in top_bullets}
                reranked = []
                for bid in id_order:
                    if bid in bullet_map:
                        reranked.append(bullet_map.pop(bid))
                # Append any not mentioned by AI
                reranked.extend(bullet_map.values())
                # Re-score by position
                for i, b in enumerate(reranked):
                    b["relevance"] = round(1.0 - (i / max(len(reranked), 1)), 4)

                # Enhance job reasons from AI
                job_reasons = ai_result.get("job_reasons", {})
                for j in fallback["ranked_jobs"]:
                    chid = str(j["career_history_id"])
                    if chid in job_reasons:
                        j["reason"] = job_reasons[chid]

                return {
                    "ranked_bullets": reranked[:limit],
                    "ranked_jobs": fallback["ranked_jobs"],
                    "suggested_skills": ai_result.get("suggested_skills", fallback["suggested_skills"]),
                    "analysis_mode": "ai",
                }
        except Exception as exc:
            logger.warning("AI best-picks failed, using fallback: %s", exc)

        return fallback

    result = route_inference("recipe_best_picks", {"jd_text": jd_text}, _python_fallback, _ai_handler)
    return jsonify(result), 200
