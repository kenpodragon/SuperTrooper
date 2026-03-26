"""MCP tool functions for resume generation: recipes, generate_resume, document tools, onboarding.

Orchestrator note: call register_resume_gen_tools(mcp) in mcp_server.py.
"""

from __future__ import annotations

import sys
import os
from pathlib import Path

import db


def register_resume_gen_tools(mcp):
    """Register all resume generation MCP tools with the given MCP server instance."""

    @mcp.tool()
    def list_recipes(template_id: int = 0, is_active: bool = True) -> dict:
        """List available resume recipes.

        Args:
            template_id: Filter by template ID (0 = all templates).
            is_active: Filter by active status (default True).
        """
        sql = "SELECT id, name, description, headline, template_id, application_id, is_active, created_at FROM resume_recipes WHERE 1=1"
        params = []
        if template_id > 0:
            sql += " AND template_id = %s"
            params.append(template_id)
        if is_active:
            sql += " AND is_active = TRUE"
        sql += " ORDER BY id"
        rows = db.query(sql, params)
        return {"recipes": rows, "count": len(rows)}

    # ------------------------------------------------------------------
    # Template Management
    # ------------------------------------------------------------------

    @mcp.tool()
    def list_templates(active_only: bool = True) -> dict:
        """List available resume templates.

        Args:
            active_only: If True, only return active templates (default True).
        """
        if active_only:
            rows = db.query(
                "SELECT id, name, filename, description, is_active, template_type, created_at "
                "FROM resume_templates WHERE is_active = TRUE ORDER BY id"
            )
        else:
            rows = db.query(
                "SELECT id, name, filename, description, is_active, template_type, created_at "
                "FROM resume_templates ORDER BY id"
            )
        return {"templates": rows, "count": len(rows)}

    @mcp.tool()
    def upload_template(name: str = "", file_path: str = "", description: str = "",
                        template_type: str = "full") -> dict:
        """Register a new resume template from a .docx file path.

        Args:
            name: Template name (required).
            file_path: Path to the .docx template file (required).
            description: Optional description.
            template_type: Template type (full, section, header). Default 'full'.
        """
        if not name or not file_path:
            return {"error": "name and file_path are required"}

        from pathlib import Path as _Path
        p = _Path(file_path)
        if not p.exists():
            return {"error": f"File not found: {file_path}"}

        with open(file_path, "rb") as f:
            blob = f.read()

        row = db.execute_returning(
            """INSERT INTO resume_templates (name, filename, template_blob, description, template_type)
               VALUES (%s, %s, %s, %s, %s) RETURNING id, name, filename, description, is_active, template_type, created_at""",
            (name, p.name, blob, description or None, template_type),
        )
        return row or {"error": "Failed to upload template"}

    @mcp.tool()
    def activate_template(template_id: int = 0) -> dict:
        """Activate a resume template, making it available for use.

        Args:
            template_id: Template ID to activate.
        """
        if template_id <= 0:
            return {"error": "template_id is required"}
        row = db.execute_returning(
            "UPDATE resume_templates SET is_active = TRUE, updated_at = NOW() WHERE id = %s "
            "RETURNING id, name, is_active",
            (template_id,),
        )
        if not row:
            return {"error": f"Template id={template_id} not found"}
        return row

    @mcp.tool()
    def deactivate_template(template_id: int = 0) -> dict:
        """Deactivate a resume template, hiding it from active use.

        Args:
            template_id: Template ID to deactivate.
        """
        if template_id <= 0:
            return {"error": "template_id is required"}
        row = db.execute_returning(
            "UPDATE resume_templates SET is_active = FALSE, updated_at = NOW() WHERE id = %s "
            "RETURNING id, name, is_active",
            (template_id,),
        )
        if not row:
            return {"error": f"Template id={template_id} not found"}
        return row

    @mcp.tool()
    def get_recipe(recipe_id: int = 0) -> dict:
        """Get a single resume recipe with full JSON.

        Args:
            recipe_id: Recipe ID to fetch.
        """
        if recipe_id <= 0:
            return {"error": "recipe_id is required"}
        row = db.query_one("SELECT * FROM resume_recipes WHERE id = %s", (recipe_id,))
        if not row:
            return {"error": f"Recipe id={recipe_id} not found"}
        return row

    @mcp.tool()
    def create_recipe(name: str = "", headline: str = "", template_id: int = 0,
                      recipe_json: str = "", description: str = "", application_id: int = 0) -> dict:
        """Create a new resume recipe.

        Args:
            name: Recipe name (e.g. "V32 AI Architect - Optum").
            headline: Resume headline text.
            template_id: ID of the template to use.
            recipe_json: JSON string of slot-to-source mappings.
            description: Optional description.
            application_id: Optional linked application ID (0 = none).
        """
        if not name or not template_id or not recipe_json:
            return {"error": "name, template_id, and recipe_json are required"}
        import json as _json
        try:
            recipe = _json.loads(recipe_json) if isinstance(recipe_json, str) else recipe_json
        except _json.JSONDecodeError as e:
            return {"error": f"Invalid recipe JSON: {e}"}

        app_id = application_id if application_id > 0 else None
        row = db.execute_returning(
            """INSERT INTO resume_recipes (name, description, headline, template_id, recipe, application_id)
               VALUES (%s, %s, %s, %s, %s, %s) RETURNING *""",
            (name, description or None, headline or None, template_id, _json.dumps(recipe), app_id),
        )
        return row or {"error": "Failed to create recipe"}

    @mcp.tool()
    def update_recipe(recipe_id: int = 0, name: str = "", headline: str = "",
                      recipe_json: str = "", description: str = "", is_active: bool = True) -> dict:
        """Update an existing resume recipe.

        Args:
            recipe_id: Recipe ID to update.
            name: New name (empty = keep current).
            headline: New headline (empty = keep current).
            recipe_json: New recipe JSON (empty = keep current).
            description: New description (empty = keep current).
            is_active: Active status.
        """
        if recipe_id <= 0:
            return {"error": "recipe_id is required"}
        existing = db.query_one("SELECT * FROM resume_recipes WHERE id = %s", (recipe_id,))
        if not existing:
            return {"error": f"Recipe id={recipe_id} not found"}

        import json as _json
        updates = {}
        if name:
            updates["name"] = name
        if headline:
            updates["headline"] = headline
        if description:
            updates["description"] = description
        if recipe_json:
            try:
                updates["recipe"] = _json.dumps(
                    _json.loads(recipe_json) if isinstance(recipe_json, str) else recipe_json
                )
            except _json.JSONDecodeError as e:
                return {"error": f"Invalid recipe JSON: {e}"}
        updates["is_active"] = is_active

        if not updates:
            return existing

        set_clauses = ", ".join(f"{k} = %s" for k in updates)
        set_clauses += ", updated_at = NOW()"
        values = list(updates.values()) + [recipe_id]
        row = db.execute_returning(
            f"UPDATE resume_recipes SET {set_clauses} WHERE id = %s RETURNING *",
            values,
        )
        return row or {"error": "Failed to update recipe"}

    @mcp.tool()
    def generate_resume(version: str = "v32", variant: str = "base", output_path: str = "",
                        recipe_id: int = 0) -> dict:
        """Generate a .docx resume from a recipe or legacy spec.

        When recipe_id is provided, uses recipe-based generation (pointer references).
        Otherwise falls back to legacy spec-based generation.

        Args:
            version: Resume version for legacy path (default v32).
            variant: Resume variant for legacy path (default base).
            output_path: Where to save the .docx. Defaults to Output/resume_{version}_{variant}.docx.
            recipe_id: Recipe ID from resume_recipes (0 = use legacy spec path).
        """
        import io as _io
        import json as _json
        import re as _re
        from pathlib import Path
        from docx import Document

        PLACEHOLDER_RE = _re.compile(r"\{\{([A-Z0-9_]+)\}\}")
        BOLD_SEPS = {
            "highlight": ": ", "job_bullet": ": ", "education": " | ",
            "certification": " | ", "additional_exp": " | ", "ref_link": " | ",
            "job_header": ", ",
        }

        # === RECIPE PATH ===
        if recipe_id > 0:
            recipe_row = db.query_one("SELECT * FROM resume_recipes WHERE id = %s", (recipe_id,))
            if not recipe_row:
                return {"error": f"Recipe id={recipe_id} not found"}

            tmpl = db.query_one(
                "SELECT name, template_blob, template_map FROM resume_templates WHERE id = %s AND is_active = TRUE",
                (recipe_row["template_id"],),
            )
            if not tmpl:
                return {"error": f"Template id={recipe_row['template_id']} not found"}

            template_blob = bytes(tmpl["template_blob"])
            template_map = tmpl["template_map"] or {}
            recipe_json = recipe_row["recipe"]
            if isinstance(recipe_json, str):
                recipe_json = _json.loads(recipe_json)

            content = _resolve_recipe_db(recipe_json)
            if recipe_row.get("headline"):
                content["HEADLINE"] = recipe_row["headline"]

            slot_info = {}
            for slot in template_map.get("slots", []):
                if slot.get("placeholder"):
                    slot_info[slot["placeholder"]] = {
                        "slot_type": slot.get("slot_type", ""),
                        "formatting": slot.get("formatting", {}),
                    }

            doc = Document(_io.BytesIO(template_blob))
            filled = 0
            for para in doc.paragraphs:
                match = PLACEHOLDER_RE.search(para.text)
                if not match:
                    continue
                placeholder = match.group(1)
                if placeholder not in content:
                    if para.runs:
                        para.runs[0].text = ""
                        for run in para.runs[1:]:
                            run.text = ""
                    continue
                text = content[placeholder]
                info = slot_info.get(placeholder, {})
                slot_type = info.get("slot_type", "")
                formatting = info.get("formatting", {})
                if formatting.get("bold_label") and slot_type in BOLD_SEPS:
                    sep = BOLD_SEPS[slot_type]
                    idx = text.find(sep)
                    if idx >= 0 and para.runs:
                        para.runs[0].text = text[:idx]
                        para.runs[0].bold = True
                        if len(para.runs) > 1:
                            para.runs[1].text = text[idx:]
                            para.runs[1].bold = None
                            for run in para.runs[2:]:
                                run.text = ""
                                run.bold = None
                        else:
                            para.runs[0].text = text
                    else:
                        if para.runs:
                            para.runs[0].text = text
                            for run in para.runs[1:]:
                                run.text = ""
                else:
                    if para.runs:
                        para.runs[0].text = text
                        for run in para.runs[1:]:
                            run.text = ""
                    else:
                        para.text = text
                filled += 1

            if not output_path:
                import datetime
                today = datetime.date.today().isoformat()
                # Get candidate name from resume_header
                header = db.query_one("SELECT full_name FROM resume_header LIMIT 1")
                cand_name = (header.get("full_name") or "Resume").replace(" ", "_").replace("/", "_") if header else "Resume"

                if recipe_row.get("application_id"):
                    app = db.query_one("SELECT company_name, role, date_applied FROM applications WHERE id = %s",
                                       (recipe_row["application_id"],))
                    if app:
                        company = (app.get("company_name") or "Unknown").replace(" ", "_").replace("/", "_")
                        role = (app.get("role") or "Role").replace(" ", "_").replace("/", "_")
                        date = (app.get("date_applied") or today)
                        if hasattr(date, "isoformat"):
                            date = date.isoformat()
                        output_path = f"Output/{company}_{role}_{date}/{cand_name}_{role}_{date}.docx"
                    else:
                        target = (recipe_row.get("target_role") or "General").replace(" ", "_").replace("/", "_")
                        output_path = f"Output/{cand_name}_{target}_{today}.docx"
                else:
                    target = (recipe_row.get("target_role") or "General").replace(" ", "_").replace("/", "_")
                    output_path = f"Output/{cand_name}_{target}_{today}.docx"
            out = Path(output_path)
            out.parent.mkdir(parents=True, exist_ok=True)
            doc.save(str(out))

            return {
                "status": "generated",
                "output_path": str(out),
                "recipe_id": recipe_id,
                "recipe_name": recipe_row["name"],
                "slots_filled": filled,
                "total_content": len(content),
            }

        # === LEGACY SPEC PATH ===
        tmpl = db.query_one(
            "SELECT template_blob, template_map FROM resume_templates "
            "WHERE name = 'V32 Placeholder' AND is_active = TRUE"
        )
        if not tmpl:
            return {"error": "Placeholder template not found in DB"}

        template_blob = bytes(tmpl["template_blob"])
        template_map = tmpl["template_map"] or {}

        rv = db.query_one(
            "SELECT spec FROM resume_versions WHERE version = %s AND variant = %s AND spec IS NOT NULL",
            (version, variant),
        )
        if not rv:
            return {"error": f"No spec found for {version}/{variant}"}
        spec = rv["spec"] if isinstance(rv["spec"], dict) else _json.loads(rv["spec"])

        header = db.query_one("SELECT * FROM resume_header LIMIT 1")
        education = db.query("SELECT * FROM education ORDER BY sort_order")
        certifications = db.query("SELECT * FROM certifications WHERE is_active = TRUE ORDER BY sort_order")

        employers = spec.get("experience_employers", [])
        career = {}
        for emp in employers:
            ch = db.query_one(
                "SELECT * FROM career_history WHERE employer ILIKE %s ORDER BY start_date DESC LIMIT 1",
                (f"%{emp}%",),
            )
            if ch:
                career[emp] = ch

        content = {}
        for slot in template_map.get("slots", []):
            ph = slot.get("placeholder")
            orig = slot.get("original_text")
            if ph and orig:
                content[ph] = orig

        if header:
            content["HEADER_NAME"] = f"{header['full_name']}, {header['credentials']}"
            parts = [header["location"]]
            if header.get("location_note"):
                parts[0] += f" ({header['location_note']})"
            parts.append(header["email"])
            parts.append(header["phone"])
            if header.get("linkedin_url"):
                parts.append(header["linkedin_url"])
            content["HEADER_CONTACT"] = " \u2022 ".join(parts)

        for key in ["headline", "summary_text"]:
            if key in spec:
                target = "HEADLINE" if key == "headline" else "SUMMARY"
                content[target] = spec[key]

        for i, b in enumerate(spec.get("highlight_bullets", []), 1):
            content[f"HIGHLIGHT_{i}"] = b

        if "keywords" in spec:
            content["KEYWORDS"] = " | ".join(spec["keywords"])
        if "executive_keywords" in spec:
            content["EXEC_KEYWORDS"] = " | ".join(spec["executive_keywords"])
        if "technical_keywords" in spec:
            content["TECH_KEYWORDS"] = " | ".join(spec["technical_keywords"])

        exp_bullets = spec.get("experience_bullets", {})
        for job_n, emp_name in enumerate(employers, 1):
            emp_data = career.get(emp_name, {})
            bullets_raw = exp_bullets.get(emp_name, [])

            if emp_data.get("intro_text"):
                content[f"JOB_{job_n}_INTRO"] = emp_data["intro_text"]

            subtitle_texts = {content.get(f"JOB_{job_n}_SUBTITLE_1", ""),
                              content.get(f"JOB_{job_n}_SUBTITLE_2", "")}
            bullet_texts = [b for b in bullets_raw
                            if b != content.get(f"JOB_{job_n}_INTRO") and b not in subtitle_texts]
            for i, bt in enumerate(bullet_texts, 1):
                content[f"JOB_{job_n}_BULLET_{i}"] = bt

        for i, ref in enumerate(spec.get("references", []), 1):
            for j, link in enumerate(ref.get("links", []), 1):
                content[f"REF_{i}_LINK_{j}"] = f"{link['text']} | {link['desc']}"

        slot_info = {}
        for slot in template_map.get("slots", []):
            if slot.get("placeholder"):
                slot_info[slot["placeholder"]] = {
                    "slot_type": slot.get("slot_type", ""),
                    "formatting": slot.get("formatting", {}),
                }

        doc = Document(_io.BytesIO(template_blob))
        filled = 0
        for para in doc.paragraphs:
            match = PLACEHOLDER_RE.search(para.text)
            if not match:
                continue
            placeholder = match.group(1)
            if placeholder not in content:
                if para.runs:
                    para.runs[0].text = ""
                    for run in para.runs[1:]:
                        run.text = ""
                continue

            text = content[placeholder]
            info = slot_info.get(placeholder, {})
            slot_type = info.get("slot_type", "")
            formatting = info.get("formatting", {})

            if formatting.get("bold_label") and slot_type in BOLD_SEPS:
                sep = BOLD_SEPS[slot_type]
                idx = text.find(sep)
                if idx >= 0 and para.runs:
                    para.runs[0].text = text[:idx]
                    para.runs[0].bold = True
                    if len(para.runs) > 1:
                        para.runs[1].text = text[idx:]
                        para.runs[1].bold = None
                        for run in para.runs[2:]:
                            run.text = ""
                            run.bold = None
                    else:
                        para.runs[0].text = text
                else:
                    if para.runs:
                        para.runs[0].text = text
                        for run in para.runs[1:]:
                            run.text = ""
            else:
                if para.runs:
                    para.runs[0].text = text
                    for run in para.runs[1:]:
                        run.text = ""
                else:
                    para.text = text
            filled += 1

        if not output_path:
            import datetime
            header = db.query_one("SELECT full_name FROM resume_header LIMIT 1")
            cand_name = (header.get("full_name") or "Resume").replace(" ", "_").replace("/", "_") if header else "Resume"
            today = datetime.date.today().isoformat()
            output_path = f"Output/{cand_name}_{version}_{variant}_{today}.docx"
        out = Path(output_path)
        out.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(out))

        return {
            "status": "generated",
            "output_path": str(out),
            "version": version,
            "variant": variant,
            "slots_filled": filled,
            "total_content": len(content),
        }

    @mcp.tool()
    def mcp_read_docx(file_path: str) -> dict:
        """Extract text from a .docx file.

        Args:
            file_path: Path to the .docx file.

        Returns:
            {"text": str, "paragraphs": int}
        """
        sys.path.insert(0, str(Path(__file__).parent.parent / "utils"))
        from read_docx import read_full_text
        text = read_full_text(file_path)
        return {"text": text, "paragraphs": len([p for p in text.split("\n") if p.strip()])}

    @mcp.tool()
    def mcp_read_pdf(file_path: str, pages: str | None = None) -> dict:
        """Extract text from a .pdf file.

        Args:
            file_path: Path to the .pdf file.
            pages: Optional page range (e.g., "1-5"). Default reads all.

        Returns:
            {"text": str}
        """
        sys.path.insert(0, str(Path(__file__).parent.parent / "utils"))
        from read_pdf import read_pdf_text
        text = read_pdf_text(file_path, pages=pages)
        return {"text": text}

    @mcp.tool()
    def mcp_templatize_resume(file_path: str, output_dir: str = "/tmp", layout: str = "auto") -> dict:
        """Convert a .docx resume into a placeholder template.

        Args:
            file_path: Path to the .docx resume.
            output_dir: Directory for output files. Defaults to /tmp.
            layout: Template layout name. Default 'auto'.

        Returns:
            {"template_path": str, "map_path": str, "slots": int}
        """
        import json
        sys.path.insert(0, str(Path(__file__).parent.parent / "utils"))
        from templatize_resume import templatize

        stem = Path(file_path).stem
        out_docx = os.path.join(output_dir, f"{stem}_placeholder.docx")
        out_map = os.path.join(output_dir, f"{stem}_map.json")
        templatize(file_path, out_docx, out_map, layout_name=layout)

        with open(out_map) as f:
            tmap = json.load(f)
        return {"template_path": out_docx, "map_path": out_map, "slots": len(tmap)}

    @mcp.tool()
    def mcp_compare_docs(file_a: str, file_b: str) -> dict:
        """Compare two .docx documents and return a match score + diff.

        Args:
            file_a: Path to first .docx document.
            file_b: Path to second .docx document.

        Returns:
            {"match_percentage": float, "diff_count": int, "diff_text": str}
        """
        sys.path.insert(0, str(Path(__file__).parent.parent / "utils"))
        from compare_docs import extract_paragraphs, compare_text

        paras_a = extract_paragraphs(file_a)
        paras_b = extract_paragraphs(file_b)
        diff = compare_text(paras_a, paras_b)
        total = max(len(paras_a), len(paras_b), 1)
        matching = sum(1 for a, b in zip(paras_a, paras_b) if a.strip() == b.strip())
        return {
            "match_percentage": round((matching / total) * 100, 1),
            "diff_count": len([l for l in diff.split("\n") if l.startswith("+") or l.startswith("-")]),
            "diff_text": diff,
        }

    @mcp.tool()
    def mcp_docx_to_pdf(file_path: str, output_path: str | None = None) -> dict:
        """Convert a .docx file to .pdf.

        Args:
            file_path: Path to the .docx file.
            output_path: Optional output path. Defaults to same name with .pdf extension.

        Returns:
            {"pdf_path": str}
        """
        sys.path.insert(0, str(Path(__file__).parent.parent / "utils"))
        from docx_to_pdf import docx_to_pdf as _docx_to_pdf
        pdf_path = _docx_to_pdf(file_path, output_path=output_path)
        return {"pdf_path": pdf_path}

    @mcp.tool()
    def mcp_edit_docx(file_path: str, find_text: str, replace_text: str, output_path: str | None = None, replace_all: bool = False) -> dict:
        """Find and replace text in a .docx file.

        Args:
            file_path: Path to the .docx file.
            find_text: Text to find.
            replace_text: Replacement text.
            output_path: Optional output path. Defaults to overwriting original.
            replace_all: Replace all occurrences. Default False.

        Returns:
            {"replacements": int}
        """
        sys.path.insert(0, str(Path(__file__).parent.parent / "utils"))
        from edit_docx import find_replace
        count = find_replace(file_path, find_text, replace_text, output_path=output_path, replace_all=replace_all)
        return {"replacements": count}

    @mcp.tool()
    def onboard_resume(file_path: str) -> dict:
        """Run the full onboarding pipeline on a resume file.

        Parses resume into career data, creates template + recipe, verifies reconstruction.

        Args:
            file_path: Path to .docx or .pdf resume file.

        Returns:
            Full pipeline report with inserted row counts, template/recipe IDs, match score.
        """
        from pathlib import PurePath

        with open(file_path, "rb") as f:
            file_bytes = f.read()

        filename = PurePath(file_path).name
        file_ext = PurePath(file_path).suffix.lower()

        from routes.onboard import _process_file
        return _process_file(filename, file_bytes, file_ext)


_V2_ALLOWED_TABLES = {
    "bullets", "career_history", "skills", "summary_variants",
    "education", "certifications", "resume_header",
}

_V2_DEFAULT_COLUMNS = {
    "bullets": "text",
    "skills": "name",
    "summary_variants": "text",
    "education": "description",
    "certifications": "name",
}


def _resolve_single_ref(ref: dict):
    """Resolve one v2 ref dict to its display value.

    Handles: literal, resume_header (full row), career_history (full row),
    ids (plural), and standard table+id lookups.
    """
    # Literal text — return directly
    if "literal" in ref:
        return ref["literal"]

    table = ref.get("ref") or ref.get("table")
    if not table:
        return None

    if table not in _V2_ALLOWED_TABLES:
        return f"[blocked table: {table}]"

    # Plural ids — resolve each individually
    if "ids" in ref:
        column = _V2_DEFAULT_COLUMNS.get(table, "name")
        ids = ref["ids"]
        rows = db.query(
            f"SELECT id, {column} FROM {table} WHERE id = ANY(%s)",
            (ids,),
        )
        by_id = {r["id"]: r[column] for r in rows}
        return [by_id.get(i, "") for i in ids]

    row_id = ref.get("id", 1)

    # Special case: resume_header → full dict
    if table == "resume_header":
        h = db.query_one("SELECT * FROM resume_header WHERE id = %s", (row_id,))
        if not h:
            return {}
        return {
            "full_name": h.get("full_name", ""),
            "credentials": h.get("credentials", ""),
            "location": h.get("location", ""),
            "location_note": h.get("location_note", ""),
            "email": h.get("email", ""),
            "phone": h.get("phone", ""),
            "linkedin_url": h.get("linkedin_url", ""),
        }

    # Special case: career_history → full dict
    if table == "career_history":
        r = db.query_one("SELECT * FROM career_history WHERE id = %s", (row_id,))
        if not r:
            return {}
        return {
            "id": r["id"],
            "employer": r.get("employer", ""),
            "title": r.get("title", ""),
            "start_date": r.get("start_date"),
            "end_date": r.get("end_date"),
            "location": r.get("location", ""),
            "industry": r.get("industry", ""),
            "synopsis": r.get("intro_text", ""),
        }

    # Standard table lookup with default column
    column = _V2_DEFAULT_COLUMNS.get(table, "name")
    row = db.query_one(f"SELECT {column} FROM {table} WHERE id = %s", (row_id,))
    return row[column] if row and row.get(column) else ""


def _resolve_recipe_v2(recipe_json: dict) -> dict:
    """Walk a v2 recipe structure and resolve all refs to display values."""
    resolved = {}

    # Single-item slots
    for key in ("header", "headline", "summary"):
        if key in recipe_json:
            ref = recipe_json[key]
            if isinstance(ref, dict):
                resolved[key] = _resolve_single_ref(ref)
            else:
                resolved[key] = ref

    # Experience array — each entry may have:
    #   - top-level ref/table (career_history ref) OR sub-fields (header, title, synopsis)
    #   - bullets array
    if "experience" in recipe_json:
        exp_list = []
        for entry in recipe_json["experience"]:
            exp_item = {}

            # Case 1: top-level career_history ref (spec format)
            if "ref" in entry and entry["ref"] == "career_history":
                job_data = _resolve_single_ref(entry)
                if isinstance(job_data, dict):
                    exp_item.update(job_data)
            # Case 2: sub-fields from migration (header, title as separate literals/refs)
            elif "header" in entry or "title" in entry:
                if "header" in entry:
                    h = _resolve_single_ref(entry["header"]) if isinstance(entry["header"], dict) else entry["header"]
                    exp_item["employer"] = h if isinstance(h, str) else (h.get("employer", "") if isinstance(h, dict) else str(h))
                if "title" in entry:
                    t = _resolve_single_ref(entry["title"]) if isinstance(entry["title"], dict) else entry["title"]
                    exp_item["title"] = t if isinstance(t, str) else str(t)
                if "titles" in entry:
                    titles = [_resolve_single_ref(t) if isinstance(t, dict) else t for t in entry["titles"]]
                    exp_item["title"] = " / ".join(str(t) for t in titles if t)
                if "subtitle" in entry:
                    s = _resolve_single_ref(entry["subtitle"]) if isinstance(entry["subtitle"], dict) else entry["subtitle"]
                    exp_item["subtitle"] = s if isinstance(s, str) else str(s)
                if "subtitles" in entry:
                    subs = [_resolve_single_ref(s) if isinstance(s, dict) else s for s in entry["subtitles"]]
                    exp_item["subtitle"] = " / ".join(str(s) for s in subs if s)

            # Synopsis — may resolve to a career_history dict; extract just the text
            if "synopsis" in entry:
                syn = entry["synopsis"]
                resolved_syn = _resolve_single_ref(syn) if isinstance(syn, dict) else syn
                if isinstance(resolved_syn, dict):
                    exp_item["synopsis"] = resolved_syn.get("synopsis", "") or resolved_syn.get("intro_text", "")
                else:
                    exp_item["synopsis"] = resolved_syn

            # Bullets
            if "bullets" in entry:
                bullet_ref = entry["bullets"]
                if isinstance(bullet_ref, dict):
                    resolved_bullets = _resolve_single_ref(bullet_ref)
                    exp_item["bullets"] = resolved_bullets if isinstance(resolved_bullets, list) else [resolved_bullets]
                elif isinstance(bullet_ref, list):
                    exp_item["bullets"] = [
                        _resolve_single_ref(b) if isinstance(b, dict) else b
                        for b in bullet_ref
                    ]

            exp_list.append(exp_item)
        resolved["experience"] = exp_list

    # Array slots
    for key in ("skills", "education", "certifications", "highlights", "additional_experience"):
        if key in recipe_json:
            items = recipe_json[key]
            if isinstance(items, list):
                resolved[key] = [
                    _resolve_single_ref(item) if isinstance(item, dict) else item
                    for item in items
                ]
            elif isinstance(items, dict):
                resolved[key] = _resolve_single_ref(items)

    # Custom slot — dict of key→ref
    if "custom" in recipe_json:
        custom = recipe_json["custom"]
        if isinstance(custom, dict):
            resolved["custom"] = {
                k: _resolve_single_ref(v) if isinstance(v, dict) else v
                for k, v in custom.items()
            }

    return resolved


def _resolve_recipe_db(recipe_json: dict, recipe_version: int = 1) -> dict:
    """Resolve a recipe JSON into a content_map using db module.

    Shared helper for MCP tool and Flask route.
    Dispatches to v2 resolver when recipe_version >= 2.
    """
    if recipe_version >= 2:
        return _resolve_recipe_v2(recipe_json)

    ALLOWED = {"bullets", "career_history", "skills", "summary_variants",
               "education", "certifications", "resume_header"}
    content = {}
    for slot_name, ref in recipe_json.items():
        if "literal" in ref:
            content[slot_name] = ref["literal"]
        elif "ids" in ref:
            table = ref["table"]
            if table not in ALLOWED:
                continue
            ids = ref["ids"]
            column = ref.get("column", "name")
            rows = db.query(
                f"SELECT id, {column} FROM {table} WHERE id = ANY(%s)",
                (ids,),
            )
            by_id = {r["id"]: r[column] for r in rows}
            values = [by_id.get(i, "") for i in ids]
            content[slot_name] = " | ".join(v for v in values if v)
        elif "table" in ref:
            table = ref["table"]
            if table not in ALLOWED:
                continue
            row_id = ref.get("id", 1)
            column = ref.get("column") or ref.get("slot")

            if table == "resume_header":
                h = db.query_one("SELECT * FROM resume_header WHERE id = %s", (row_id,))
                if not h:
                    content[slot_name] = ""
                    continue
                if column in ("name",) or slot_name == "HEADER_NAME":
                    content[slot_name] = f"{h['full_name']}, {h['credentials']}"
                elif column in ("contact",) or slot_name == "HEADER_CONTACT":
                    parts = [h["location"]]
                    if h.get("location_note"):
                        parts[0] += f" ({h['location_note']})"
                    parts.append(h["email"])
                    parts.append(h["phone"])
                    if h.get("linkedin_url"):
                        parts.append(h["linkedin_url"])
                    content[slot_name] = " \u2022 ".join(parts)
                else:
                    row = db.query_one(f"SELECT {column} FROM {table} WHERE id = %s", (row_id,))
                    content[slot_name] = row[column] if row else ""
            elif column is None or column == "":
                if table == "career_history":
                    r = db.query_one(
                        "SELECT employer, title, location, industry, notes, "
                        "start_date, end_date, is_current FROM career_history WHERE id = %s",
                        (row_id,),
                    )
                    def _fmt_date(d):
                        """Format date: bare year if Jan 1, otherwise Month Year."""
                        if not d or not hasattr(d, "strftime"):
                            return str(d) if d else ""
                        if d.month == 1 and d.day == 1:
                            return str(d.year)  # Bare year (e.g., 2012)
                        return d.strftime("%B %Y")  # Full month (e.g., August 2021)

                    if r:
                        # Build date suffix
                        date_suffix = ""
                        if ref.get("include_dates") and r.get("start_date"):
                            start = _fmt_date(r["start_date"])
                            is_bare_year = (hasattr(r["start_date"], "month") and r["start_date"].month == 1 and r["start_date"].day == 1)
                            sep = "-" if is_bare_year else " - "  # No spaces for year-only dates
                            if r.get("is_current"):
                                date_suffix = f"\t{start}{sep}Present"
                            elif r.get("end_date"):
                                end = _fmt_date(r["end_date"])
                                date_suffix = f"\t{start}{sep}{end}"

                        if ref.get("format") == "oneliner":
                            # One-liner: Title | Company (Industry)\tDates
                            title = r.get("title", "")
                            employer = r["employer"]
                            industry = r.get("industry", "")
                            if industry:
                                content[slot_name] = f"{title} | {employer} ({industry}){date_suffix}"
                            else:
                                content[slot_name] = f"{title} | {employer}{date_suffix}"
                        else:
                            # Standard: Company, Location (Onsite) {Industry}\tDates
                            parts = [r["employer"]]
                            if r.get("location") and r["location"] not in r["employer"]:
                                parts.append(f", {r['location']}")
                            if r.get("notes") and any(w in (r["notes"] or "").lower() for w in ("onsite", "remote", "hybrid")):
                                parts.append(f" ({r['notes']})")
                            if r.get("industry"):
                                parts.append(f" {{{r['industry']}}}")
                            content[slot_name] = "".join(parts) + date_suffix
                elif table == "education":
                    r = db.query_one("SELECT degree, field, institution, location FROM education WHERE id = %s", (row_id,))
                    if r:
                        parts = [p for p in [r.get("degree"), r.get("field")] if p]
                        result = ", ".join(parts)
                        if r.get("institution"):
                            result += f" | {r['institution']}"
                        if r.get("location"):
                            result += f" \u2014 {r['location']}"
                        content[slot_name] = result
                elif table == "certifications":
                    r = db.query_one("SELECT name, issuer FROM certifications WHERE id = %s", (row_id,))
                    if r:
                        content[slot_name] = f"{r['name']} | {r['issuer']}" if r.get("issuer") else r["name"]
                else:
                    content[slot_name] = ""
            else:
                if table == "career_history" and column == "title":
                    # Title with optional date range (multi-role companies)
                    def _fmt_dt(d):
                        if not d or not hasattr(d, "strftime"):
                            return str(d) if d else ""
                        return str(d.year) if d.month == 1 and d.day == 1 else d.strftime("%B %Y")

                    row = db.query_one(
                        "SELECT title, start_date, end_date, is_current FROM career_history WHERE id = %s",
                        (row_id,),
                    )
                    if row and row.get("title"):
                        title_text = row["title"]
                        if ref.get("include_dates") and row.get("start_date"):
                            start = _fmt_dt(row["start_date"])
                            if row.get("is_current"):
                                title_text += f", {start} \u2013 Present"
                            elif row.get("end_date"):
                                end = _fmt_dt(row["end_date"])
                                title_text += f", {start} \u2013 {end}"
                        content[slot_name] = title_text
                    else:
                        content[slot_name] = ""
                else:
                    row = db.query_one(f"SELECT {column} FROM {table} WHERE id = %s", (row_id,))
                    if row and row.get(column):
                        value = row[column]
                        # Handle dict/JSON values (e.g. intro_text stored as {"text": "..."} from v1 import)
                        if isinstance(value, dict):
                            content[slot_name] = value.get("text", str(value))
                        else:
                            content[slot_name] = value
                    else:
                        content[slot_name] = ""
    return content
