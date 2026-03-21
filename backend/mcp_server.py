"""SuperTroopers MCP Server — exposes DB tools for Claude Code.

Run:  python mcp_server.py   # starts MCP server on stdio
"""

import sys
import os
import re
from collections import Counter
from datetime import date, timedelta
from pathlib import Path

# Ensure the backend directory is on sys.path
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from mcp.server.fastmcp import FastMCP
import db

mcp = FastMCP(
    "supertroopers",
    instructions="SuperTroopers hiring platform database tools.",
    host="0.0.0.0",
    port=int(os.environ.get("MCP_PORT", 8056)),
)


# ---------------------------------------------------------------------------
# Career / Bullets / Skills
# ---------------------------------------------------------------------------

@mcp.tool()
def search_bullets(
    query: str = "",
    tags: list[str] | None = None,
    role_type: str = "",
    industry: str = "",
    limit: int = 20,
) -> dict:
    """Search resume bullets by text, tags, role_type, or industry.

    Args:
        query: Text to search for in bullet text (ILIKE).
        tags: List of tags to filter by (array overlap).
        role_type: Filter by role suitability (e.g. CTO, VP Eng, Director).
        industry: Filter by industry suitability (e.g. defense, manufacturing).
        limit: Max results to return (default 20).
    """
    clauses, params = [], []
    if query:
        clauses.append("b.text ILIKE %s")
        params.append(f"%{query}%")
    if tags:
        clauses.append("b.tags && %s")
        params.append(tags)
    if role_type:
        clauses.append("%s = ANY(b.role_suitability)")
        params.append(role_type)
    if industry:
        clauses.append("%s = ANY(b.industry_suitability)")
        params.append(industry)

    where = f"WHERE {' AND '.join(clauses)}" if clauses else ""
    rows = db.query(
        f"""
        SELECT b.id, b.text, b.type, b.tags, b.role_suitability,
               b.industry_suitability, b.metrics_json, b.detail_recall,
               ch.employer, ch.title
        FROM bullets b
        LEFT JOIN career_history ch ON ch.id = b.career_history_id
        {where}
        ORDER BY b.id
        LIMIT %s
        """,
        params + [limit],
    )
    return {"count": len(rows), "bullets": rows}


@mcp.tool()
def get_career_history(employer: str = "") -> dict:
    """Get career history for an employer (or all if blank).

    Args:
        employer: Employer name to search (ILIKE match). Leave blank for all.
    """
    if employer:
        rows = db.query(
            """
            SELECT ch.*, array_agg(json_build_object(
                'id', b.id, 'text', b.text, 'type', b.type, 'tags', b.tags
            )) FILTER (WHERE b.id IS NOT NULL) AS bullets
            FROM career_history ch
            LEFT JOIN bullets b ON b.career_history_id = ch.id
            WHERE ch.employer ILIKE %s
            GROUP BY ch.id
            ORDER BY ch.start_date DESC NULLS LAST
            """,
            (f"%{employer}%",),
        )
    else:
        rows = db.query(
            """
            SELECT id, employer, title, start_date, end_date, location,
                   industry, team_size, budget_usd, revenue_impact, is_current
            FROM career_history
            ORDER BY start_date DESC NULLS LAST
            """
        )
    return {"count": len(rows), "career_history": rows}


@mcp.tool()
def get_summary_variant(role_type: str) -> dict:
    """Get a professional summary variant for a target role type.

    Args:
        role_type: Target role (e.g. CTO, VP Eng, Director, AI Architect, SW Architect, PM, Sr SWE).
    """
    row = db.query_one(
        "SELECT id, role_type, text, updated_at FROM summary_variants WHERE role_type = %s",
        (role_type,),
    )
    if not row:
        # Try ILIKE fallback
        row = db.query_one(
            "SELECT id, role_type, text, updated_at FROM summary_variants WHERE role_type ILIKE %s",
            (f"%{role_type}%",),
        )
    return row or {"error": f"No summary variant found for role_type '{role_type}'"}


@mcp.tool()
def get_skills(category: str = "") -> dict:
    """List skills, optionally filtered by category.

    Args:
        category: Skill category filter (language, framework, platform, methodology, tool).
    """
    if category:
        rows = db.query(
            "SELECT id, name, category, proficiency, last_used_year FROM skills WHERE category = %s ORDER BY name",
            (category,),
        )
    else:
        rows = db.query(
            "SELECT id, name, category, proficiency, last_used_year FROM skills ORDER BY category, name"
        )
    return {"count": len(rows), "skills": rows}


# ---------------------------------------------------------------------------
# JD Matching / Gap Analysis
# ---------------------------------------------------------------------------

@mcp.tool()
def match_jd(jd_text: str) -> dict:
    """Match a job description against resume bullets. Returns best matches and gaps.

    Args:
        jd_text: Full job description text to analyze.
    """
    stop_words = {
        "the", "a", "an", "and", "or", "but", "in", "on", "at", "to", "for",
        "of", "with", "by", "from", "is", "are", "was", "were", "be", "been",
        "being", "have", "has", "had", "do", "does", "did", "will", "would",
        "could", "should", "may", "might", "shall", "can", "need", "must",
        "that", "this", "these", "those", "it", "its", "we", "our", "you",
        "your", "they", "their", "them", "he", "she", "his", "her", "as",
        "if", "not", "no", "so", "up", "out", "about", "into", "over",
        "after", "before", "between", "through", "during", "such", "each",
        "which", "who", "whom", "what", "where", "when", "how", "all", "any",
        "both", "few", "more", "most", "other", "some", "than", "too", "very",
        "just", "also", "well", "able", "etc", "including", "experience",
        "work", "working", "role", "position", "team", "company", "years",
        "strong", "required", "preferred", "requirements", "qualifications",
        "responsibilities", "job", "description", "candidate", "ideal",
    }

    words = re.findall(r"[a-z][a-z+#/.]+", jd_text.lower())
    keywords = [w for w in words if w not in stop_words and len(w) > 2]
    top_keywords = [k for k, _ in Counter(keywords).most_common(30)]

    matched_bullets = []
    seen_ids = set()
    for kw in top_keywords[:15]:
        rows = db.query(
            """
            SELECT b.id, b.text, b.type, b.tags, b.role_suitability,
                   ch.employer, ch.title
            FROM bullets b
            LEFT JOIN career_history ch ON ch.id = b.career_history_id
            WHERE b.text ILIKE %s
            LIMIT 5
            """,
            (f"%{kw}%",),
        )
        for row in rows:
            if row["id"] not in seen_ids:
                seen_ids.add(row["id"])
                row["matched_keyword"] = kw
                matched_bullets.append(row)

    matched_skills = []
    for kw in top_keywords:
        skills = db.query(
            "SELECT id, name, category, proficiency FROM skills WHERE name ILIKE %s",
            (f"%{kw}%",),
        )
        for s in skills:
            if s["id"] not in {ms["id"] for ms in matched_skills}:
                s["matched_keyword"] = kw
                matched_skills.append(s)

    covered = {b["matched_keyword"] for b in matched_bullets} | {s["matched_keyword"] for s in matched_skills}
    gaps = [kw for kw in top_keywords if kw not in covered]

    return {
        "jd_keywords": top_keywords,
        "matched_bullets": matched_bullets[:30],
        "matched_skills": matched_skills,
        "gaps": gaps,
        "coverage_pct": round(len(covered) / max(len(top_keywords), 1) * 100, 1),
    }


# ---------------------------------------------------------------------------
# Applications
# ---------------------------------------------------------------------------

@mcp.tool()
def search_applications(
    status: str = "",
    company: str = "",
    source: str = "",
    limit: int = 50,
) -> dict:
    """Search job applications by status, company, or source.

    Args:
        status: Filter by status (Applied, Interview, Rejected, etc.).
        company: Filter by company name (ILIKE match).
        source: Filter by source (Indeed, LinkedIn, etc.).
        limit: Max results (default 50).
    """
    clauses, params = [], []
    if status:
        clauses.append("status = %s")
        params.append(status)
    if company:
        clauses.append("company_name ILIKE %s")
        params.append(f"%{company}%")
    if source:
        clauses.append("source = %s")
        params.append(source)

    where = f"WHERE {' AND '.join(clauses)}" if clauses else ""
    rows = db.query(
        f"""
        SELECT id, company_id, company_name, role, date_applied, source, status,
               resume_version, jd_url, contact_name, notes,
               last_status_change, created_at
        FROM applications
        {where}
        ORDER BY date_applied DESC NULLS LAST
        LIMIT %s
        """,
        params + [limit],
    )
    return {"count": len(rows), "applications": rows}


@mcp.tool()
def add_application(
    company_name: str,
    role: str,
    source: str = "Direct",
    status: str = "Applied",
    notes: str = "",
    company_id: int | None = None,
    date_applied: str | None = None,
    jd_url: str = "",
    jd_text: str = "",
) -> dict:
    """Add a new job application to the tracker.

    Args:
        company_name: Company name.
        role: Job title / role applied for.
        source: Application source (Indeed, LinkedIn, Dice, ZipRecruiter, Direct, Recruiter, Referral).
        status: Initial status (default Applied).
        notes: Any notes.
        company_id: Optional company ID if known.
        date_applied: Date applied (YYYY-MM-DD). Defaults to today.
        jd_url: Job description URL.
        jd_text: Job description text.
    """
    row = db.execute_returning(
        """
        INSERT INTO applications (company_id, company_name, role, date_applied,
            source, status, jd_url, jd_text, notes, last_status_change)
        VALUES (%s, %s, %s, COALESCE(%s::date, CURRENT_DATE), %s, %s, %s, %s, %s, NOW())
        RETURNING *
        """,
        (company_id, company_name, role, date_applied, source, status, jd_url, jd_text, notes),
    )
    return row


@mcp.tool()
def update_application(
    id: int,
    status: str = "",
    notes: str = "",
) -> dict:
    """Update an application's status and/or notes.

    Args:
        id: Application ID.
        status: New status value.
        notes: Updated notes (replaces existing).
    """
    sets, params = [], []
    if status:
        sets.append("status = %s")
        params.append(status)
        sets.append("last_status_change = NOW()")
    if notes:
        sets.append("notes = %s")
        params.append(notes)
    if not sets:
        return {"error": "Provide status or notes to update"}

    params.append(id)
    row = db.execute_returning(
        f"UPDATE applications SET {', '.join(sets)} WHERE id = %s RETURNING *",
        params,
    )
    return row or {"error": f"Application {id} not found"}


# ---------------------------------------------------------------------------
# Companies
# ---------------------------------------------------------------------------

@mcp.tool()
def search_companies(
    query: str = "",
    priority: str = "",
    sector: str = "",
    limit: int = 50,
) -> dict:
    """Search target companies by name, priority, or sector.

    Args:
        query: Company name search (ILIKE).
        priority: Priority tier (A, B, C).
        sector: Sector filter (ILIKE).
        limit: Max results (default 50).
    """
    clauses, params = [], []
    if query:
        clauses.append("name ILIKE %s")
        params.append(f"%{query}%")
    if priority:
        clauses.append("priority = %s")
        params.append(priority)
    if sector:
        clauses.append("sector ILIKE %s")
        params.append(f"%{sector}%")

    where = f"WHERE {' AND '.join(clauses)}" if clauses else ""
    rows = db.query(
        f"""
        SELECT id, name, sector, hq_location, size, stage, fit_score, priority,
               target_role, resume_variant, melbourne_relevant, comp_range, notes
        FROM companies
        {where}
        ORDER BY fit_score DESC NULLS LAST, name
        LIMIT %s
        """,
        params + [limit],
    )
    return {"count": len(rows), "companies": rows}


@mcp.tool()
def get_company_dossier(name: str) -> dict:
    """Get full company info including applications, contacts, and related emails.

    Args:
        name: Company name (ILIKE match).
    """
    company = db.query_one(
        "SELECT * FROM companies WHERE name ILIKE %s",
        (f"%{name}%",),
    )
    if not company:
        return {"error": f"Company '{name}' not found"}

    applications = db.query(
        """
        SELECT id, role, date_applied, source, status, resume_version, notes,
               last_status_change
        FROM applications
        WHERE company_name ILIKE %s
        ORDER BY date_applied DESC NULLS LAST
        """,
        (f"%{name}%",),
    )

    contacts = db.query(
        """
        SELECT id, name, title, relationship, email, linkedin_url,
               relationship_strength, last_contact
        FROM contacts
        WHERE company ILIKE %s
        ORDER BY relationship_strength, name
        """,
        (f"%{name}%",),
    )

    emails = db.query(
        """
        SELECT id, date, from_name, subject, snippet, category
        FROM emails
        WHERE subject ILIKE %s OR from_address ILIKE %s OR from_name ILIKE %s
        ORDER BY date DESC NULLS LAST
        LIMIT 20
        """,
        (f"%{name}%", f"%{name}%", f"%{name}%"),
    )

    company["applications"] = applications
    company["contacts"] = contacts
    company["emails"] = emails
    return company


# ---------------------------------------------------------------------------
# Contacts
# ---------------------------------------------------------------------------

@mcp.tool()
def search_contacts(company: str = "", name: str = "") -> dict:
    """Search contacts by company or name.

    Args:
        company: Company name filter (ILIKE).
        name: Contact name filter (ILIKE).
    """
    clauses, params = [], []
    if company:
        clauses.append("company ILIKE %s")
        params.append(f"%{company}%")
    if name:
        clauses.append("name ILIKE %s")
        params.append(f"%{name}%")

    where = f"WHERE {' AND '.join(clauses)}" if clauses else ""
    rows = db.query(
        f"""
        SELECT id, name, company, title, relationship, email, phone,
               linkedin_url, relationship_strength, last_contact, notes
        FROM contacts
        {where}
        ORDER BY relationship_strength, name
        """,
        params,
    )
    return {"count": len(rows), "contacts": rows}


@mcp.tool()
def network_check(company: str) -> dict:
    """Find contacts and related emails for a company. Useful for warm intro research.

    Args:
        company: Company name to check network for.
    """
    contacts = db.query(
        """
        SELECT id, name, title, relationship, email, linkedin_url,
               relationship_strength, last_contact, notes
        FROM contacts
        WHERE company ILIKE %s
        ORDER BY relationship_strength, name
        """,
        (f"%{company}%",),
    )

    emails = db.query(
        """
        SELECT id, date, from_name, from_address, subject, snippet, category
        FROM emails
        WHERE from_name ILIKE %s OR from_address ILIKE %s OR subject ILIKE %s
        ORDER BY date DESC NULLS LAST
        LIMIT 20
        """,
        (f"%{company}%", f"%{company}%", f"%{company}%"),
    )

    applications = db.query(
        """
        SELECT id, role, status, date_applied, contact_name, contact_email
        FROM applications
        WHERE company_name ILIKE %s
        ORDER BY date_applied DESC NULLS LAST
        """,
        (f"%{company}%",),
    )

    return {
        "company": company,
        "contacts": contacts,
        "related_emails": emails,
        "applications": applications,
        "has_warm_intro": any(
            c.get("relationship_strength") in ("strong", "warm") for c in contacts
        ),
    }


# ---------------------------------------------------------------------------
# Emails
# ---------------------------------------------------------------------------

@mcp.tool()
def search_emails(
    query: str = "",
    category: str = "",
    after: str = "",
    before: str = "",
    limit: int = 20,
) -> dict:
    """Search emails by text, category, and date range.

    Args:
        query: Text search in subject, snippet, body (ILIKE).
        category: Email category (application, rejection, interview, recruiter, reference, other).
        after: Date filter - emails after this date (YYYY-MM-DD).
        before: Date filter - emails before this date (YYYY-MM-DD).
        limit: Max results (default 20).
    """
    clauses, params = [], []
    if query:
        clauses.append("(subject ILIKE %s OR snippet ILIKE %s OR body ILIKE %s)")
        params.extend([f"%{query}%", f"%{query}%", f"%{query}%"])
    if category:
        clauses.append("category = %s")
        params.append(category)
    if after:
        clauses.append("date >= %s")
        params.append(after)
    if before:
        clauses.append("date <= %s")
        params.append(before)

    where = f"WHERE {' AND '.join(clauses)}" if clauses else ""
    rows = db.query(
        f"""
        SELECT id, date, from_name, from_address, subject, snippet, category,
               application_id
        FROM emails
        {where}
        ORDER BY date DESC NULLS LAST
        LIMIT %s
        """,
        params + [limit],
    )
    return {"count": len(rows), "emails": rows}


# ---------------------------------------------------------------------------
# Analytics
# ---------------------------------------------------------------------------

@mcp.tool()
def get_analytics() -> dict:
    """Get pipeline statistics: funnel, source effectiveness, monthly activity, and summary counts."""
    funnel = db.query("SELECT status, count, pct FROM application_funnel")
    sources = db.query(
        "SELECT source, total_apps, got_response, response_rate_pct, got_interview, interview_rate_pct FROM source_effectiveness"
    )
    monthly = db.query(
        "SELECT month, applications, interviews, rejections, ghosted, offers FROM monthly_activity"
    )
    summary = db.query_one(
        """
        SELECT
            (SELECT COUNT(*) FROM applications) AS total_applications,
            (SELECT COUNT(*) FROM applications WHERE status IN ('Phone Screen','Interview','Technical','Final')) AS in_progress,
            (SELECT COUNT(*) FROM applications WHERE status = 'Offer') AS offers,
            (SELECT COUNT(*) FROM applications WHERE status = 'Rejected') AS rejected,
            (SELECT COUNT(*) FROM applications WHERE status = 'Ghosted') AS ghosted,
            (SELECT COUNT(*) FROM interviews) AS total_interviews,
            (SELECT COUNT(*) FROM companies) AS total_companies,
            (SELECT COUNT(*) FROM contacts) AS total_contacts
        """
    )
    return {
        "summary": summary,
        "funnel": funnel,
        "sources": sources,
        "monthly": monthly,
    }


# ---------------------------------------------------------------------------
# Content Sections (Candidate Profile, Rejection Analysis, etc.)
# ---------------------------------------------------------------------------

@mcp.tool()
def get_candidate_profile(section: str = "", format: str = "sections") -> dict:
    """Get candidate profile data. Returns full profile or specific sections.

    Args:
        section: Filter by section name (e.g. "Identity", "Career Narrative", "Target Roles", "Compensation", "References"). Leave blank for all.
        format: "sections" for structured data, "text" for reconstructed markdown.
    """
    return _get_content_document("candidate_profile", section, format)


@mcp.tool()
def get_rejection_analysis(section: str = "", format: str = "sections") -> dict:
    """Get rejection and ghosting analysis data.

    Args:
        section: Filter by section (e.g. "Companies with Confirmed Interview Activity", "Pattern Analysis"). Leave blank for all.
        format: "sections" for structured data, "text" for reconstructed markdown.
    """
    return _get_content_document("rejection_analysis", section, format)


@mcp.tool()
def get_voice_rules(category: str = "", part: int = 0, format: str = "rules") -> dict:
    """Get voice guide rules for content generation. Use to check writing quality.

    Args:
        category: Filter by rule type: banned_word, banned_construction, caution_word, structural_tell, resume_rule, cover_letter_rule, final_check, linkedin_pattern, stephen_ism, context_pattern, quick_reference. Leave blank for all.
        part: Filter by Voice Guide part number (1-8). 0 = all parts.
        format: "rules" for structured data, "text" for reconstructed guide.
    """
    clauses, params = [], []
    if category:
        clauses.append("category = %s")
        params.append(category)
    if part:
        clauses.append("part = %s")
        params.append(part)

    where = f"WHERE {' AND '.join(clauses)}" if clauses else ""
    rows = db.query(
        f"""
        SELECT id, part, part_title, category, subcategory, rule_text,
               explanation, examples_bad, examples_good
        FROM voice_rules
        {where}
        ORDER BY sort_order
        """,
        params,
    )

    if format == "text":
        parts = {}
        for row in rows:
            key = f"Part {row['part']}: {row['part_title']}"
            if key not in parts:
                parts[key] = []
            parts[key].append(row)

        text_parts = []
        for key, rules_list in parts.items():
            text_parts.append(f"## {key}\n")
            for r in rules_list:
                if r["category"] == "banned_word":
                    text_parts.append(f"- {r['rule_text']}")
                else:
                    text_parts.append(f"### {r['rule_text']}")
                    if r["explanation"]:
                        text_parts.append(r["explanation"])
                text_parts.append("")
        return {"text": "\n".join(text_parts), "rule_count": len(rows)}

    return {"count": len(rows), "rules": rows}


@mcp.tool()
def check_voice(text: str) -> dict:
    """Check text against voice guide banned words and constructions. Returns violations.

    Args:
        text: The text to check against voice rules.
    """
    text_lower = text.lower()
    violations = []

    banned_words = db.query(
        "SELECT rule_text, subcategory FROM voice_rules WHERE category = 'banned_word'"
    )
    for bw in banned_words:
        word = bw["rule_text"].lower()
        if word in text_lower:
            violations.append({
                "type": "banned_word",
                "match": bw["rule_text"],
                "subcategory": bw["subcategory"],
            })

    banned_constructions = db.query(
        "SELECT rule_text, subcategory, explanation FROM voice_rules WHERE category = 'banned_construction'"
    )
    for bc in banned_constructions:
        pattern = bc["rule_text"].lower()
        if pattern in text_lower:
            violations.append({
                "type": "banned_construction",
                "match": bc["rule_text"],
                "subcategory": bc["subcategory"],
                "explanation": bc["explanation"],
            })

    return {
        "text_length": len(text),
        "violations": violations,
        "violation_count": len(violations),
        "clean": len(violations) == 0,
    }


@mcp.tool()
def get_salary_data(role: str = "", tier: int = 0) -> dict:
    """Get salary benchmarks and COLA market data for target roles.

    Args:
        role: Search by role title (e.g. "CTO", "VP", "Director"). Leave blank for all.
        tier: Filter by tier (1=Executive, 2=Director, 3=Senior IC, 4=PM, 5=Academia). 0 = all.
    """
    clauses, params = [], []
    if role:
        clauses.append("role_title ILIKE %s")
        params.append(f"%{role}%")
    if tier:
        clauses.append("tier = %s")
        params.append(tier)

    where = f"WHERE {' AND '.join(clauses)}" if clauses else ""
    benchmarks = db.query(
        f"""
        SELECT role_title, tier, tier_name, national_median_range,
               melbourne_range, remote_range, hcol_range, target_realistic
        FROM salary_benchmarks
        {where}
        ORDER BY sort_order
        """,
        params,
    )

    markets = db.query(
        "SELECT market_name, col_index_approx, cola_factor, melbourne_200k_equiv, melbourne_250k_equiv FROM cola_markets ORDER BY cola_factor"
    )

    return {
        "benchmarks": benchmarks,
        "benchmark_count": len(benchmarks),
        "cola_markets": markets,
        "cola_formula": "Melbourne-equivalent = Posted Salary x (97 / Market COL Index)",
    }


def _get_content_document(doc_name: str, section: str, format: str) -> dict:
    """Shared helper for content section retrieval."""
    clauses = ["source_document = %s"]
    params = [doc_name]

    if section:
        clauses.append("(section ILIKE %s OR subsection ILIKE %s)")
        params.extend([f"%{section}%", f"%{section}%"])

    where = f"WHERE {' AND '.join(clauses)}"
    rows = db.query(
        f"""
        SELECT id, section, subsection, sort_order, content, content_format
        FROM content_sections
        {where}
        ORDER BY sort_order
        """,
        params,
    )

    if not rows:
        return {"error": f"No content found for '{doc_name}'" + (f" section '{section}'" if section else "")}

    if format == "text":
        parts = []
        for row in rows:
            if row["subsection"]:
                parts.append(f"### {row['subsection']}\n\n{row['content']}")
            else:
                parts.append(f"## {row['section']}\n\n{row['content']}")
        return {
            "document": doc_name,
            "text": "\n\n---\n\n".join(parts),
            "section_count": len(rows),
        }

    return {
        "document": doc_name,
        "sections": rows,
        "count": len(rows),
    }


# ---------------------------------------------------------------------------
# Resume Data
# ---------------------------------------------------------------------------

@mcp.tool()
def get_resume_data(version: str = "v32", variant: str = "base", section: str = "") -> dict:
    """Get full resume data for reconstruction or querying. Returns header, spec, experience, education, certs.

    Args:
        version: Resume version (default v32).
        variant: Resume variant (default base).
        section: Optional filter: "header", "education", "certifications", "experience", "spec", "keywords". Leave blank for all.
    """
    import json as _json

    if section == "header":
        return db.query_one("SELECT * FROM resume_header LIMIT 1") or {"error": "No header"}

    if section == "education":
        rows = db.query("SELECT * FROM education ORDER BY sort_order")
        return {"education": rows, "count": len(rows)}

    if section == "certifications":
        rows = db.query("SELECT * FROM certifications WHERE is_active = TRUE ORDER BY sort_order")
        return {"certifications": rows, "count": len(rows)}

    # Get the spec
    rv = db.query_one(
        "SELECT * FROM resume_versions WHERE version = %s AND variant = %s AND spec IS NOT NULL",
        (version, variant),
    )
    if not rv:
        return {"error": f"No spec found for {version}/{variant}"}

    spec = rv["spec"] if isinstance(rv["spec"], dict) else _json.loads(rv["spec"])

    if section == "spec":
        return {"version": version, "variant": variant, "spec": spec}

    if section == "keywords":
        return {
            "keywords": spec.get("keywords", []),
            "executive_keywords": spec.get("executive_keywords", []),
            "technical_keywords": spec.get("technical_keywords", []),
        }

    if section == "experience":
        experience = []
        for employer_name in spec.get("experience_employers", []):
            ch = db.query_one(
                "SELECT * FROM career_history WHERE employer ILIKE %s",
                (f"%{employer_name}%",),
            )
            if ch:
                bullets = db.query(
                    "SELECT id, text, type, tags FROM bullets WHERE career_history_id = %s ORDER BY id",
                    (ch["id"],),
                )
                ch["bullets"] = bullets
                experience.append(ch)
        return {"experience": experience, "count": len(experience)}

    # Full data
    header = db.query_one("SELECT * FROM resume_header LIMIT 1")
    education = db.query("SELECT * FROM education ORDER BY sort_order")
    certifications = db.query("SELECT * FROM certifications WHERE is_active = TRUE ORDER BY sort_order")

    experience = []
    for employer_name in spec.get("experience_employers", []):
        ch = db.query_one(
            "SELECT * FROM career_history WHERE employer ILIKE %s",
            (f"%{employer_name}%",),
        )
        if ch:
            bullets = db.query(
                "SELECT id, text, type, tags FROM bullets WHERE career_history_id = %s ORDER BY id",
                (ch["id"],),
            )
            ch["bullets"] = bullets
            experience.append(ch)

    return {
        "version": version,
        "variant": variant,
        "header": header,
        "spec": spec,
        "experience": experience,
        "education": education,
        "certifications": certifications,
    }


# ---------------------------------------------------------------------------
# Resume Generation
# ---------------------------------------------------------------------------

@mcp.tool()
def list_recipes(template_id: int = 0, is_active: bool = True) -> dict:
    """List available resume recipes.

    Args:
        template_id: Filter by template ID (0 = all templates).
        is_active: Filter by active status (default True).
    """
    sql = "SELECT id, name, description, headline, template_id, application_id, is_active, created_at FROM resume_recipes WHERE 1=1"
    params = []
    if template_id > 0:
        sql += " AND template_id = %s"
        params.append(template_id)
    if is_active:
        sql += " AND is_active = TRUE"
    sql += " ORDER BY id"
    rows = db.query(sql, params)
    return {"recipes": rows, "count": len(rows)}


@mcp.tool()
def get_recipe(recipe_id: int = 0) -> dict:
    """Get a single resume recipe with full JSON.

    Args:
        recipe_id: Recipe ID to fetch.
    """
    if recipe_id <= 0:
        return {"error": "recipe_id is required"}
    row = db.query_one("SELECT * FROM resume_recipes WHERE id = %s", (recipe_id,))
    if not row:
        return {"error": f"Recipe id={recipe_id} not found"}
    return row


@mcp.tool()
def create_recipe(name: str = "", headline: str = "", template_id: int = 0,
                  recipe_json: str = "", description: str = "", application_id: int = 0) -> dict:
    """Create a new resume recipe.

    Args:
        name: Recipe name (e.g. "V32 AI Architect - Optum").
        headline: Resume headline text.
        template_id: ID of the template to use.
        recipe_json: JSON string of slot-to-source mappings.
        description: Optional description.
        application_id: Optional linked application ID (0 = none).
    """
    if not name or not template_id or not recipe_json:
        return {"error": "name, template_id, and recipe_json are required"}
    import json as _json
    try:
        recipe = _json.loads(recipe_json) if isinstance(recipe_json, str) else recipe_json
    except _json.JSONDecodeError as e:
        return {"error": f"Invalid recipe JSON: {e}"}

    app_id = application_id if application_id > 0 else None
    row = db.execute_returning(
        """INSERT INTO resume_recipes (name, description, headline, template_id, recipe, application_id)
           VALUES (%s, %s, %s, %s, %s, %s) RETURNING *""",
        (name, description or None, headline or None, template_id, _json.dumps(recipe), app_id),
    )
    return row or {"error": "Failed to create recipe"}


@mcp.tool()
def update_recipe(recipe_id: int = 0, name: str = "", headline: str = "",
                  recipe_json: str = "", description: str = "", is_active: bool = True) -> dict:
    """Update an existing resume recipe.

    Args:
        recipe_id: Recipe ID to update.
        name: New name (empty = keep current).
        headline: New headline (empty = keep current).
        recipe_json: New recipe JSON (empty = keep current).
        description: New description (empty = keep current).
        is_active: Active status.
    """
    if recipe_id <= 0:
        return {"error": "recipe_id is required"}
    existing = db.query_one("SELECT * FROM resume_recipes WHERE id = %s", (recipe_id,))
    if not existing:
        return {"error": f"Recipe id={recipe_id} not found"}

    import json as _json
    updates = {}
    if name:
        updates["name"] = name
    if headline:
        updates["headline"] = headline
    if description:
        updates["description"] = description
    if recipe_json:
        try:
            updates["recipe"] = _json.dumps(
                _json.loads(recipe_json) if isinstance(recipe_json, str) else recipe_json
            )
        except _json.JSONDecodeError as e:
            return {"error": f"Invalid recipe JSON: {e}"}
    updates["is_active"] = is_active

    if not updates:
        return existing

    set_clauses = ", ".join(f"{k} = %s" for k in updates)
    set_clauses += ", updated_at = NOW()"
    values = list(updates.values()) + [recipe_id]
    row = db.execute_returning(
        f"UPDATE resume_recipes SET {set_clauses} WHERE id = %s RETURNING *",
        values,
    )
    return row or {"error": "Failed to update recipe"}


def _resolve_recipe_db(recipe_json: dict) -> dict:
    """Resolve a recipe JSON into a content_map using db module.

    Shared helper for MCP tool and Flask route.
    """
    ALLOWED = {"bullets", "career_history", "skills", "summary_variants",
               "education", "certifications", "resume_header"}
    content = {}
    for slot_name, ref in recipe_json.items():
        if "literal" in ref:
            content[slot_name] = ref["literal"]
        elif "ids" in ref:
            table = ref["table"]
            if table not in ALLOWED:
                continue
            ids = ref["ids"]
            column = ref.get("column", "name")
            rows = db.query(
                f"SELECT id, {column} FROM {table} WHERE id = ANY(%s)",
                (ids,),
            )
            by_id = {r["id"]: r[column] for r in rows}
            values = [by_id.get(i, "") for i in ids]
            content[slot_name] = " | ".join(v for v in values if v)
        elif "table" in ref:
            table = ref["table"]
            if table not in ALLOWED:
                continue
            row_id = ref.get("id", 1)
            column = ref.get("column") or ref.get("slot")

            if table == "resume_header":
                h = db.query_one("SELECT * FROM resume_header WHERE id = %s", (row_id,))
                if not h:
                    content[slot_name] = ""
                    continue
                if column in ("name",) or slot_name == "HEADER_NAME":
                    content[slot_name] = f"{h['full_name']}, {h['credentials']}"
                elif column in ("contact",) or slot_name == "HEADER_CONTACT":
                    parts = [h["location"]]
                    if h.get("location_note"):
                        parts[0] += f" ({h['location_note']})"
                    parts.append(h["email"])
                    parts.append(h["phone"])
                    if h.get("linkedin_url"):
                        parts.append(h["linkedin_url"])
                    content[slot_name] = " \u2022 ".join(parts)
                else:
                    row = db.query_one(f"SELECT {column} FROM {table} WHERE id = %s", (row_id,))
                    content[slot_name] = row[column] if row else ""
            elif column is None or column == "":
                if table == "career_history":
                    r = db.query_one("SELECT employer, location, industry FROM career_history WHERE id = %s", (row_id,))
                    if r:
                        parts = [r["employer"]]
                        if r.get("location"):
                            parts.append(f", {r['location']}")
                        if r.get("industry"):
                            parts.append(f" {{{r['industry']}}}")
                        content[slot_name] = "".join(parts)
                elif table == "education":
                    r = db.query_one("SELECT degree, field, institution, location FROM education WHERE id = %s", (row_id,))
                    if r:
                        parts = [p for p in [r.get("degree"), r.get("field")] if p]
                        result = ", ".join(parts)
                        if r.get("institution"):
                            result += f" | {r['institution']}"
                        if r.get("location"):
                            result += f" \u2014 {r['location']}"
                        content[slot_name] = result
                elif table == "certifications":
                    r = db.query_one("SELECT name, issuer FROM certifications WHERE id = %s", (row_id,))
                    if r:
                        content[slot_name] = f"{r['name']} | {r['issuer']}" if r.get("issuer") else r["name"]
                else:
                    content[slot_name] = ""
            else:
                row = db.query_one(f"SELECT {column} FROM {table} WHERE id = %s", (row_id,))
                content[slot_name] = row[column] if row and row.get(column) else ""
    return content


@mcp.tool()
def generate_resume(version: str = "v32", variant: str = "base", output_path: str = "",
                    recipe_id: int = 0) -> dict:
    """Generate a .docx resume from a recipe or legacy spec.

    When recipe_id is provided, uses recipe-based generation (pointer references).
    Otherwise falls back to legacy spec-based generation.

    Args:
        version: Resume version for legacy path (default v32).
        variant: Resume variant for legacy path (default base).
        output_path: Where to save the .docx. Defaults to Output/resume_{version}_{variant}.docx.
        recipe_id: Recipe ID from resume_recipes (0 = use legacy spec path).
    """
    import io as _io
    import json as _json
    import re as _re
    from pathlib import Path
    from docx import Document

    PLACEHOLDER_RE = _re.compile(r"\{\{([A-Z0-9_]+)\}\}")
    BOLD_SEPS = {
        "highlight": ": ", "job_bullet": ": ", "education": " | ",
        "certification": " | ", "additional_exp": " | ", "ref_link": " | ",
        "job_header": ", ",
    }

    # === RECIPE PATH ===
    if recipe_id > 0:
        recipe_row = db.query_one("SELECT * FROM resume_recipes WHERE id = %s", (recipe_id,))
        if not recipe_row:
            return {"error": f"Recipe id={recipe_id} not found"}

        # Load template from recipe's template_id
        tmpl = db.query_one(
            "SELECT name, template_blob, template_map FROM resume_templates WHERE id = %s AND is_active = TRUE",
            (recipe_row["template_id"],),
        )
        if not tmpl:
            return {"error": f"Template id={recipe_row['template_id']} not found"}

        template_blob = bytes(tmpl["template_blob"])
        template_map = tmpl["template_map"] or {}
        recipe_json = recipe_row["recipe"]
        if isinstance(recipe_json, str):
            recipe_json = _json.loads(recipe_json)

        content = _resolve_recipe_db(recipe_json)
        if recipe_row.get("headline"):
            content["HEADLINE"] = recipe_row["headline"]

        # Build slot info and fill template (same as legacy path below)
        slot_info = {}
        for slot in template_map.get("slots", []):
            if slot.get("placeholder"):
                slot_info[slot["placeholder"]] = {
                    "slot_type": slot.get("slot_type", ""),
                    "formatting": slot.get("formatting", {}),
                }

        doc = Document(_io.BytesIO(template_blob))
        filled = 0
        for para in doc.paragraphs:
            match = PLACEHOLDER_RE.search(para.text)
            if not match:
                continue
            placeholder = match.group(1)
            if placeholder not in content:
                if para.runs:
                    para.runs[0].text = ""
                    for run in para.runs[1:]:
                        run.text = ""
                continue
            text = content[placeholder]
            info = slot_info.get(placeholder, {})
            slot_type = info.get("slot_type", "")
            formatting = info.get("formatting", {})
            if formatting.get("bold_label") and slot_type in BOLD_SEPS:
                sep = BOLD_SEPS[slot_type]
                idx = text.find(sep)
                if idx >= 0 and para.runs:
                    para.runs[0].text = text[:idx]
                    para.runs[0].bold = True
                    if len(para.runs) > 1:
                        para.runs[1].text = text[idx:]
                        para.runs[1].bold = None
                        for run in para.runs[2:]:
                            run.text = ""
                            run.bold = None
                    else:
                        para.runs[0].text = text
                else:
                    if para.runs:
                        para.runs[0].text = text
                        for run in para.runs[1:]:
                            run.text = ""
            else:
                if para.runs:
                    para.runs[0].text = text
                    for run in para.runs[1:]:
                        run.text = ""
                else:
                    para.text = text
            filled += 1

        if not output_path:
            # Organize output by company/role/date when application_id linked
            if recipe_row.get("application_id"):
                app = db.query_one("SELECT company_name, role, date_applied FROM applications WHERE id = %s",
                                   (recipe_row["application_id"],))
                if app:
                    import datetime
                    company = (app.get("company_name") or "Unknown").replace(" ", "_").replace("/", "_")
                    role = (app.get("role") or "Role").replace(" ", "_").replace("/", "_")
                    date = (app.get("date_applied") or datetime.date.today()).isoformat()
                    output_path = f"Output/{company}_{role}_{date}/resume.docx"
                else:
                    output_path = f"Output/resume_recipe_{recipe_id}.docx"
            else:
                output_path = f"Output/resume_recipe_{recipe_id}.docx"
        out = Path(output_path)
        out.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(out))

        return {
            "status": "generated",
            "output_path": str(out),
            "recipe_id": recipe_id,
            "recipe_name": recipe_row["name"],
            "slots_filled": filled,
            "total_content": len(content),
        }

    # === LEGACY SPEC PATH ===
    # Load template
    tmpl = db.query_one(
        "SELECT template_blob, template_map FROM resume_templates "
        "WHERE name = 'V32 Placeholder' AND is_active = TRUE"
    )
    if not tmpl:
        return {"error": "Placeholder template not found in DB"}

    template_blob = bytes(tmpl["template_blob"])
    template_map = tmpl["template_map"] or {}

    # Load spec
    rv = db.query_one(
        "SELECT spec FROM resume_versions WHERE version = %s AND variant = %s AND spec IS NOT NULL",
        (version, variant),
    )
    if not rv:
        return {"error": f"No spec found for {version}/{variant}"}
    spec = rv["spec"] if isinstance(rv["spec"], dict) else _json.loads(rv["spec"])

    # Load supporting data
    header = db.query_one("SELECT * FROM resume_header LIMIT 1")
    education = db.query("SELECT * FROM education ORDER BY sort_order")
    certifications = db.query("SELECT * FROM certifications WHERE is_active = TRUE ORDER BY sort_order")

    employers = spec.get("experience_employers", [])
    career = {}
    for emp in employers:
        ch = db.query_one(
            "SELECT * FROM career_history WHERE employer ILIKE %s ORDER BY start_date DESC LIMIT 1",
            (f"%{emp}%",),
        )
        if ch:
            career[emp] = ch

    # Build content map (seed from template_map, overlay with spec)
    content = {}
    for slot in template_map.get("slots", []):
        ph = slot.get("placeholder")
        orig = slot.get("original_text")
        if ph and orig:
            content[ph] = orig

    # Overlay dynamic content from spec
    if header:
        content["HEADER_NAME"] = f"{header['full_name']}, {header['credentials']}"
        parts = [header["location"]]
        if header.get("location_note"):
            parts[0] += f" ({header['location_note']})"
        parts.append(header["email"])
        parts.append(header["phone"])
        if header.get("linkedin_url"):
            parts.append(header["linkedin_url"])
        content["HEADER_CONTACT"] = " \u2022 ".join(parts)

    for key in ["headline", "summary_text"]:
        if key in spec:
            target = "HEADLINE" if key == "headline" else "SUMMARY"
            content[target] = spec[key]

    for i, b in enumerate(spec.get("highlight_bullets", []), 1):
        content[f"HIGHLIGHT_{i}"] = b

    if "keywords" in spec:
        content["KEYWORDS"] = " | ".join(spec["keywords"])
    if "executive_keywords" in spec:
        content["EXEC_KEYWORDS"] = " | ".join(spec["executive_keywords"])
    if "technical_keywords" in spec:
        content["TECH_KEYWORDS"] = " | ".join(spec["technical_keywords"])

    exp_bullets = spec.get("experience_bullets", {})
    for job_n, emp_name in enumerate(employers, 1):
        emp_data = career.get(emp_name, {})
        bullets_raw = exp_bullets.get(emp_name, [])

        if emp_data.get("intro_text"):
            content[f"JOB_{job_n}_INTRO"] = emp_data["intro_text"]

        subtitle_texts = {content.get(f"JOB_{job_n}_SUBTITLE_1", ""),
                          content.get(f"JOB_{job_n}_SUBTITLE_2", "")}
        bullet_texts = [b for b in bullets_raw
                        if b != content.get(f"JOB_{job_n}_INTRO") and b not in subtitle_texts]
        for i, bt in enumerate(bullet_texts, 1):
            content[f"JOB_{job_n}_BULLET_{i}"] = bt

    for i, ref in enumerate(spec.get("references", []), 1):
        for j, link in enumerate(ref.get("links", []), 1):
            content[f"REF_{i}_LINK_{j}"] = f"{link['text']} | {link['desc']}"

    # Build slot info
    slot_info = {}
    for slot in template_map.get("slots", []):
        if slot.get("placeholder"):
            slot_info[slot["placeholder"]] = {
                "slot_type": slot.get("slot_type", ""),
                "formatting": slot.get("formatting", {}),
            }

    # Fill template
    doc = Document(_io.BytesIO(template_blob))
    filled = 0
    for para in doc.paragraphs:
        match = PLACEHOLDER_RE.search(para.text)
        if not match:
            continue
        placeholder = match.group(1)
        if placeholder not in content:
            if para.runs:
                para.runs[0].text = ""
                for run in para.runs[1:]:
                    run.text = ""
            continue

        text = content[placeholder]
        info = slot_info.get(placeholder, {})
        slot_type = info.get("slot_type", "")
        formatting = info.get("formatting", {})

        if formatting.get("bold_label") and slot_type in BOLD_SEPS:
            sep = BOLD_SEPS[slot_type]
            idx = text.find(sep)
            if idx >= 0 and para.runs:
                para.runs[0].text = text[:idx]
                para.runs[0].bold = True
                if len(para.runs) > 1:
                    para.runs[1].text = text[idx:]
                    para.runs[1].bold = None
                    for run in para.runs[2:]:
                        run.text = ""
                        run.bold = None
                else:
                    para.runs[0].text = text
            else:
                if para.runs:
                    para.runs[0].text = text
                    for run in para.runs[1:]:
                        run.text = ""
        else:
            if para.runs:
                para.runs[0].text = text
                for run in para.runs[1:]:
                    run.text = ""
            else:
                para.text = text
        filled += 1

    # Save
    if not output_path:
        output_path = f"Output/resume_{version}_{variant}.docx"
    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))

    return {
        "status": "generated",
        "output_path": str(out),
        "version": version,
        "variant": variant,
        "slots_filled": filled,
        "total_content": len(content),
    }


# ---------------------------------------------------------------------------
# Saved Jobs
# ---------------------------------------------------------------------------

@mcp.tool()
def save_job(title: str, company: str = "", url: str = "", jd_text: str = "",
             source: str = "manual", fit_score: float = 0, notes: str = "") -> dict:
    """Save a job to the evaluation queue.

    Args:
        title: Job title (required).
        company: Company name.
        url: Job posting URL.
        jd_text: Full job description text.
        source: Where the job was found (indeed, linkedin, manual, etc.).
        fit_score: Initial fit score (0-10).
        notes: Any notes about the job.
    """
    # Try to find matching company
    company_id = None
    if company:
        co = db.query_one("SELECT id FROM companies WHERE name ILIKE %s", (f"%{company}%",))
        if co:
            company_id = co["id"]

    row = db.execute_returning(
        """
        INSERT INTO saved_jobs (title, company, company_id, url, jd_text, source, fit_score, status, notes)
        VALUES (%s,%s,%s,%s,%s,%s,%s,'saved',%s)
        RETURNING id, title, company, status, fit_score, created_at
        """,
        (title, company, company_id, url, jd_text, source, fit_score, notes),
    )
    return row


@mcp.tool()
def list_saved_jobs(status: str = "", limit: int = 20) -> list:
    """List saved jobs in the evaluation queue.

    Args:
        status: Filter by status (saved, evaluating, applying, applied, passed). Empty = all.
        limit: Max results (default 20).
    """
    if status:
        rows = db.query(
            "SELECT id, title, company, source, fit_score, status, created_at FROM saved_jobs WHERE status = %s ORDER BY fit_score DESC NULLS LAST LIMIT %s",
            (status, limit),
        )
    else:
        rows = db.query(
            "SELECT id, title, company, source, fit_score, status, created_at FROM saved_jobs ORDER BY created_at DESC LIMIT %s",
            (limit,),
        )
    return rows


@mcp.tool()
def update_saved_job(job_id: int, status: str = "", fit_score: float = 0, notes: str = "") -> dict:
    """Update a saved job's status, score, or notes.

    Args:
        job_id: Saved job ID (required).
        status: New status (saved, evaluating, applying, applied, passed).
        fit_score: Updated fit score (0-10).
        notes: Updated notes.
    """
    sets, params = [], []
    if status:
        sets.append("status = %s")
        params.append(status)
    if fit_score > 0:
        sets.append("fit_score = %s")
        params.append(fit_score)
    if notes:
        sets.append("notes = %s")
        params.append(notes)
    if not sets:
        return {"error": "No fields to update"}
    params.append(job_id)
    row = db.execute_returning(
        f"UPDATE saved_jobs SET {', '.join(sets)} WHERE id = %s RETURNING *",
        params,
    )
    return row or {"error": f"Saved job id={job_id} not found"}


# ---------------------------------------------------------------------------
# Gap Analysis Persistence
# ---------------------------------------------------------------------------

@mcp.tool()
def save_gap_analysis(jd_text: str = "", application_id: int = 0, saved_job_id: int = 0,
                      strong_matches: str = "", partial_matches: str = "", gaps: str = "",
                      bonus_value: str = "", fit_scores: str = "",
                      overall_score: float = 0, recommendation: str = "",
                      notes: str = "") -> dict:
    """Save a gap analysis result to the database.

    All JSON fields (strong_matches, partial_matches, gaps, bonus_value, fit_scores)
    should be passed as JSON strings.

    Args:
        jd_text: The job description text analyzed.
        application_id: Link to an application (0 = none).
        saved_job_id: Link to a saved job (0 = none).
        strong_matches: JSON string of strong match items.
        partial_matches: JSON string of partial match items.
        gaps: JSON string of gap items.
        bonus_value: JSON string of bonus value items.
        fit_scores: JSON string of fit score breakdown.
        overall_score: Overall fit score (0-10).
        recommendation: strong_apply, apply_with_tailoring, stretch, or pass.
        notes: Additional notes.
    """
    import json as _json

    row = db.execute_returning(
        """
        INSERT INTO gap_analyses (application_id, saved_job_id, jd_text,
            strong_matches, partial_matches, gaps, bonus_value,
            fit_scores, overall_score, recommendation, notes)
        VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)
        RETURNING id, overall_score, recommendation, created_at
        """,
        (
            application_id or None, saved_job_id or None, jd_text,
            strong_matches or None, partial_matches or None,
            gaps or None, bonus_value or None, fit_scores or None,
            overall_score or None, recommendation or None, notes or None,
        ),
    )

    # Link to application
    if application_id and row:
        db.execute(
            "UPDATE applications SET gap_analysis_id = %s WHERE id = %s",
            (row["id"], application_id),
        )

    return row


@mcp.tool()
def get_gap_analysis(gap_id: int = 0, application_id: int = 0, saved_job_id: int = 0) -> dict:
    """Retrieve a gap analysis by ID or by linked application/saved job.

    Args:
        gap_id: Gap analysis ID (direct lookup).
        application_id: Find gap analysis linked to this application.
        saved_job_id: Find gap analysis linked to this saved job.
    """
    if gap_id:
        row = db.query_one("SELECT * FROM gap_analyses WHERE id = %s", (gap_id,))
    elif application_id:
        row = db.query_one(
            "SELECT * FROM gap_analyses WHERE application_id = %s ORDER BY created_at DESC LIMIT 1",
            (application_id,),
        )
    elif saved_job_id:
        row = db.query_one(
            "SELECT * FROM gap_analyses WHERE saved_job_id = %s ORDER BY created_at DESC LIMIT 1",
            (saved_job_id,),
        )
    else:
        return {"error": "Provide gap_id, application_id, or saved_job_id"}
    return row or {"error": "Gap analysis not found"}


# ---------------------------------------------------------------------------
# Follow-ups
# ---------------------------------------------------------------------------

@mcp.tool()
def log_follow_up(application_id: int, method: str = "email",
                  date_sent: str = "", notes: str = "") -> dict:
    """Log a follow-up attempt for an application.

    Args:
        application_id: Application ID (required).
        method: Contact method (email, linkedin, phone).
        date_sent: Date sent (YYYY-MM-DD). Defaults to today.
        notes: Notes about the follow-up.
    """
    last = db.query_one(
        "SELECT MAX(attempt_number) AS max_num FROM follow_ups WHERE application_id = %s",
        (application_id,),
    )
    next_num = (last["max_num"] or 0) + 1 if last else 1

    row = db.execute_returning(
        """
        INSERT INTO follow_ups (application_id, attempt_number, date_sent, method, notes)
        VALUES (%s, %s, COALESCE(%s::date, CURRENT_DATE), %s, %s)
        RETURNING *
        """,
        (application_id, next_num, date_sent or None, method, notes or None),
    )
    return row


@mcp.tool()
def get_stale_applications(days: int = 14) -> list:
    """Find applications with no activity for N days.

    Args:
        days: Number of days without activity to consider stale (default 14).
    """
    rows = db.query(
        """
        SELECT a.id, a.company_name, a.role, a.status, a.last_status_change,
               a.date_applied,
               EXTRACT(DAY FROM NOW() - COALESCE(a.last_status_change, a.date_applied::timestamp))::int AS days_stale,
               (SELECT COUNT(*) FROM follow_ups f WHERE f.application_id = a.id) AS follow_up_count
        FROM applications a
        WHERE a.status NOT IN ('Rejected', 'Ghosted', 'Withdrawn', 'Accepted', 'Rescinded')
          AND COALESCE(a.last_status_change, a.date_applied::timestamp) < NOW() - INTERVAL '%s days'
        ORDER BY days_stale DESC
        """,
        (days,),
    )
    return rows


# ---------------------------------------------------------------------------
# Interview Prep & Debrief
# ---------------------------------------------------------------------------

@mcp.tool()
def save_interview_prep(interview_id: int, company_dossier: str = "",
                        prepared_questions: str = "", talking_points: str = "",
                        star_stories_selected: str = "", questions_to_ask: str = "",
                        notes: str = "") -> dict:
    """Save interview prep materials. JSON fields should be passed as JSON strings.

    Args:
        interview_id: Interview ID (required).
        company_dossier: JSON string of company research snapshot.
        prepared_questions: JSON string of prepared Q&A items.
        talking_points: JSON string of talking points.
        star_stories_selected: JSON string of selected STAR stories.
        questions_to_ask: JSON string of questions to ask the interviewer.
        notes: Additional notes.
    """
    row = db.execute_returning(
        """
        INSERT INTO interview_prep (interview_id, company_dossier, prepared_questions,
            talking_points, star_stories_selected, questions_to_ask, notes)
        VALUES (%s,%s,%s,%s,%s,%s,%s)
        RETURNING id, interview_id, created_at
        """,
        (
            interview_id,
            company_dossier or None, prepared_questions or None,
            talking_points or None, star_stories_selected or None,
            questions_to_ask or None, notes or None,
        ),
    )
    return row


@mcp.tool()
def save_interview_debrief(interview_id: int, went_well: str = "", went_poorly: str = "",
                           questions_asked: str = "", next_steps: str = "",
                           overall_feeling: str = "", lessons_learned: str = "",
                           notes: str = "") -> dict:
    """Save a structured interview debrief. JSON fields as JSON strings.

    Args:
        interview_id: Interview ID (required).
        went_well: JSON string of things that went well.
        went_poorly: JSON string of things that went poorly.
        questions_asked: JSON string of questions asked and answers given.
        next_steps: Free-text next steps.
        overall_feeling: great, good, neutral, concerned, or poor.
        lessons_learned: Free-text lessons learned.
        notes: Additional notes.
    """
    row = db.execute_returning(
        """
        INSERT INTO interview_debriefs (interview_id, went_well, went_poorly,
            questions_asked, next_steps, overall_feeling, lessons_learned, notes)
        VALUES (%s,%s,%s,%s,%s,%s,%s,%s)
        RETURNING id, interview_id, overall_feeling, created_at
        """,
        (
            interview_id,
            went_well or None, went_poorly or None,
            questions_asked or None, next_steps or None,
            overall_feeling or None, lessons_learned or None,
            notes or None,
        ),
    )
    return row


# ---------------------------------------------------------------------------
# Profile Management
# ---------------------------------------------------------------------------

@mcp.tool()
def update_header(full_name: str = "", credentials: str = "", email: str = "",
                  phone: str = "", location: str = "", linkedin_url: str = "") -> dict:
    """Update resume header / candidate contact info.

    Args:
        full_name: Full name.
        credentials: Credentials string (e.g. "PhD, CSM, PMP, MBA").
        email: Email address.
        phone: Phone number.
        location: Location.
        linkedin_url: LinkedIn profile URL.
    """
    sets, params = [], []
    for field, val in [("full_name", full_name), ("credentials", credentials),
                       ("email", email), ("phone", phone), ("location", location),
                       ("linkedin_url", linkedin_url)]:
        if val:
            sets.append(f"{field} = %s")
            params.append(val)
    if not sets:
        return {"error": "No fields to update"}

    existing = db.query_one("SELECT id FROM resume_header LIMIT 1")
    if existing:
        params.append(existing["id"])
        row = db.execute_returning(
            f"UPDATE resume_header SET {', '.join(sets)} WHERE id = %s RETURNING *",
            params,
        )
    else:
        row = db.execute_returning(
            f"INSERT INTO resume_header ({', '.join(s.split(' = ')[0] for s in sets)}) VALUES ({', '.join(['%s'] * len(params))}) RETURNING *",
            params,
        )
    return row


# ---------------------------------------------------------------------------
# Document tools
# ---------------------------------------------------------------------------


@mcp.tool()
def mcp_read_docx(file_path: str) -> dict:
    """Extract text from a .docx file.

    Args:
        file_path: Path to the .docx file.

    Returns:
        {"text": str, "paragraphs": int}
    """
    sys.path.insert(0, str(Path(__file__).parent.parent / "utils"))
    from read_docx import read_full_text
    text = read_full_text(file_path)
    return {"text": text, "paragraphs": len([p for p in text.split("\n") if p.strip()])}


@mcp.tool()
def mcp_read_pdf(file_path: str, pages: str | None = None) -> dict:
    """Extract text from a .pdf file.

    Args:
        file_path: Path to the .pdf file.
        pages: Optional page range (e.g., "1-5"). Default reads all.

    Returns:
        {"text": str}
    """
    sys.path.insert(0, str(Path(__file__).parent.parent / "utils"))
    from read_pdf import read_pdf_text
    text = read_pdf_text(file_path, pages=pages)
    return {"text": text}


@mcp.tool()
def mcp_templatize_resume(file_path: str, output_dir: str = "/tmp", layout: str = "auto") -> dict:
    """Convert a .docx resume into a placeholder template.

    Args:
        file_path: Path to the .docx resume.
        output_dir: Directory for output files. Defaults to /tmp.
        layout: Template layout name. Default 'auto'.

    Returns:
        {"template_path": str, "map_path": str, "slots": int}
    """
    import json
    sys.path.insert(0, str(Path(__file__).parent.parent / "utils"))
    from templatize_resume import templatize

    stem = Path(file_path).stem
    out_docx = os.path.join(output_dir, f"{stem}_placeholder.docx")
    out_map = os.path.join(output_dir, f"{stem}_map.json")
    templatize(file_path, out_docx, out_map, layout_name=layout)

    with open(out_map) as f:
        tmap = json.load(f)
    return {"template_path": out_docx, "map_path": out_map, "slots": len(tmap)}


@mcp.tool()
def mcp_compare_docs(file_a: str, file_b: str) -> dict:
    """Compare two .docx documents and return a match score + diff.

    Args:
        file_a: Path to first .docx document.
        file_b: Path to second .docx document.

    Returns:
        {"match_percentage": float, "diff_count": int, "diff_text": str}
    """
    sys.path.insert(0, str(Path(__file__).parent.parent / "utils"))
    from compare_docs import extract_paragraphs, compare_text

    paras_a = extract_paragraphs(file_a)
    paras_b = extract_paragraphs(file_b)
    diff = compare_text(paras_a, paras_b)
    total = max(len(paras_a), len(paras_b), 1)
    matching = sum(1 for a, b in zip(paras_a, paras_b) if a.strip() == b.strip())
    return {
        "match_percentage": round((matching / total) * 100, 1),
        "diff_count": len([l for l in diff.split("\n") if l.startswith("+") or l.startswith("-")]),
        "diff_text": diff,
    }


@mcp.tool()
def mcp_docx_to_pdf(file_path: str, output_path: str | None = None) -> dict:
    """Convert a .docx file to .pdf.

    Args:
        file_path: Path to the .docx file.
        output_path: Optional output path. Defaults to same name with .pdf extension.

    Returns:
        {"pdf_path": str}
    """
    sys.path.insert(0, str(Path(__file__).parent.parent / "utils"))
    from docx_to_pdf import docx_to_pdf as _docx_to_pdf
    pdf_path = _docx_to_pdf(file_path, output_path=output_path)
    return {"pdf_path": pdf_path}


@mcp.tool()
def mcp_edit_docx(file_path: str, find_text: str, replace_text: str, output_path: str | None = None, replace_all: bool = False) -> dict:
    """Find and replace text in a .docx file.

    Args:
        file_path: Path to the .docx file.
        find_text: Text to find.
        replace_text: Replacement text.
        output_path: Optional output path. Defaults to overwriting original.
        replace_all: Replace all occurrences. Default False.

    Returns:
        {"replacements": int}
    """
    sys.path.insert(0, str(Path(__file__).parent.parent / "utils"))
    from edit_docx import find_replace
    count = find_replace(file_path, find_text, replace_text, output_path=output_path, replace_all=replace_all)
    return {"replacements": count}


@mcp.tool()
def onboard_resume(file_path: str) -> dict:
    """Run the full onboarding pipeline on a resume file.

    Parses resume into career data, creates template + recipe, verifies reconstruction.

    Args:
        file_path: Path to .docx or .pdf resume file.

    Returns:
        Full pipeline report with inserted row counts, template/recipe IDs, match score.
    """
    from pathlib import PurePath

    with open(file_path, "rb") as f:
        file_bytes = f.read()

    filename = PurePath(file_path).name
    file_ext = PurePath(file_path).suffix.lower()

    from routes.onboard import _process_file
    return _process_file(filename, file_bytes, file_ext)


# ---------------------------------------------------------------------------
# Notifications tools
# ---------------------------------------------------------------------------


@mcp.tool()
def get_notifications(
    status: str | None = None,
    type: str | None = None,
    severity: str | None = None,
    limit: int = 20,
) -> dict:
    """Get notifications, optionally filtered.

    Args:
        status: Filter by status - 'unread', 'read', 'dismissed', or None for all
        type: Filter by notification type (new_job, status_change, follow_up_due,
              stale_warning, interview_reminder, contact_follow_up, digest_ready, email_matched)
        severity: Filter by severity level (info, action_needed, urgent)
        limit: Max results (default 20)

    Returns:
        dict with count and notifications list
    """
    clauses, params = [], []

    if status == "unread":
        clauses.append("read = FALSE")
        clauses.append("dismissed = FALSE")
    elif status == "read":
        clauses.append("read = TRUE")
    elif status == "dismissed":
        clauses.append("dismissed = TRUE")
    else:
        # Default: hide dismissed
        clauses.append("dismissed = FALSE")

    if type:
        clauses.append("type = %s")
        params.append(type)

    if severity:
        clauses.append("severity = %s")
        params.append(severity)

    where = f"WHERE {' AND '.join(clauses)}" if clauses else ""
    rows = db.query(
        f"""
        SELECT *
        FROM notifications
        {where}
        ORDER BY created_at DESC
        LIMIT %s
        """,
        params + [limit],
    )
    return {"count": len(rows), "notifications": rows}


@mcp.tool()
def dismiss_notification(notification_id: int) -> dict:
    """Dismiss a notification by ID.

    Args:
        notification_id: The notification ID to dismiss

    Returns:
        dict with success status and notification ID
    """
    count = db.execute(
        "UPDATE notifications SET dismissed = TRUE, read = TRUE WHERE id = %s",
        (notification_id,),
    )
    if count == 0:
        return {"success": False, "error": f"Notification {notification_id} not found"}
    return {"success": True, "id": notification_id, "dismissed": True}


@mcp.tool()
def create_notification(
    type: str,
    title: str,
    severity: str = "info",
    body: str | None = None,
    link: str | None = None,
    entity_type: str | None = None,
    entity_id: int | None = None,
) -> dict:
    """Create a new notification.

    Args:
        type: Notification type (new_job, status_change, follow_up_due,
              stale_warning, interview_reminder, contact_follow_up,
              digest_ready, email_matched)
        title: Notification title
        severity: info, action_needed, or urgent (default: info)
        body: Optional detailed body text
        link: Optional frontend route path (e.g., /applications/42)
        entity_type: Optional entity type (application, saved_job, contact, fresh_job)
        entity_id: Optional entity ID

    Returns:
        dict with created notification record
    """
    row = db.execute_returning(
        """
        INSERT INTO notifications
            (type, severity, title, body, link, entity_type, entity_id)
        VALUES (%s, %s, %s, %s, %s, %s, %s)
        RETURNING *
        """,
        (type, severity, title, body, link, entity_type, entity_id),
    )
    return row or {"error": "Insert failed"}


# ---------------------------------------------------------------------------
# Fresh Jobs Inbox tools
# ---------------------------------------------------------------------------


@mcp.tool()
def get_fresh_jobs(status: str = "new", source_type: str | None = None, limit: int = 20) -> dict:
    """Get fresh jobs from the inbox, optionally filtered.

    Args:
        status: Filter by status (new, reviewed, saved, passed, expired, snoozed). Default: new
        source_type: Filter by source (api_search, plugin_capture, email_parsed, social_scan, rss_feed, manual)
        limit: Max results (default 20)

    Returns:
        dict with count and fresh_jobs list
    """
    clauses, params = [], []

    if status:
        clauses.append("status = %s")
        params.append(status)
    if source_type:
        clauses.append("source_type = %s")
        params.append(source_type)

    where = f"WHERE {' AND '.join(clauses)}" if clauses else ""
    rows = db.query(
        f"""
        SELECT id, source_type, source_url, title, company, location,
               salary_range, jd_snippet, auto_score, status, discovered_at, saved_job_id
        FROM fresh_jobs
        {where}
        ORDER BY discovered_at DESC
        LIMIT %s
        """,
        params + [limit],
    )
    return {"count": len(rows), "fresh_jobs": rows}


@mcp.tool()
def triage_job(job_id: int, action: str, notes: str | None = None) -> dict:
    """Triage a fresh job: save to pipeline, pass, or snooze.

    Args:
        job_id: Fresh job ID
        action: save, pass, snooze, or review
        notes: Optional notes about the triage decision (stored if action=save)

    Returns:
        dict with result (if saved, includes saved_job_id)
    """
    valid_actions = {"save", "pass", "snooze", "review"}
    if action not in valid_actions:
        return {"error": f"action must be one of: {', '.join(sorted(valid_actions))}"}

    job = db.query_one("SELECT * FROM fresh_jobs WHERE id = %s", (job_id,))
    if not job:
        return {"error": f"Fresh job {job_id} not found"}

    if action == "save":
        if job.get("saved_job_id"):
            return {"error": "Already saved", "saved_job_id": job["saved_job_id"]}
        saved = db.execute_returning(
            """
            INSERT INTO saved_jobs (title, company, url, location, salary_range, jd_text, notes, source, status)
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, 'saved')
            RETURNING id, title, company
            """,
            (
                job.get("title"),
                job.get("company"),
                job.get("source_url"),
                job.get("location"),
                job.get("salary_range"),
                job.get("jd_full") or job.get("jd_snippet"),
                notes,
                job.get("source_type"),
            ),
        )
        saved_job_id = saved["id"]
        db.execute(
            "UPDATE fresh_jobs SET status = 'saved', saved_job_id = %s WHERE id = %s",
            (saved_job_id, job_id),
        )
        return {
            "result": "saved",
            "fresh_job_id": job_id,
            "saved_job_id": saved_job_id,
            "title": saved.get("title"),
            "company": saved.get("company"),
        }

    status_map = {"pass": "passed", "snooze": "snoozed", "review": "reviewed"}
    new_status = status_map[action]
    db.execute("UPDATE fresh_jobs SET status = %s WHERE id = %s", (new_status, job_id))
    return {"result": action, "fresh_job_id": job_id, "status": new_status}


@mcp.tool()
def batch_triage(actions: str) -> dict:
    """Batch triage multiple fresh jobs.

    Pass actions as JSON string: [{"id": 1, "action": "save"}, {"id": 2, "action": "pass"}]

    Args:
        actions: JSON string of array with id and action pairs

    Returns:
        dict with results for each job
    """
    import json

    try:
        items = json.loads(actions)
    except (json.JSONDecodeError, TypeError) as e:
        return {"error": f"Invalid JSON: {e}"}

    if not isinstance(items, list):
        return {"error": "actions must be a JSON array"}

    valid_actions = {"save", "pass", "snooze", "review"}
    status_map = {"pass": "passed", "snooze": "snoozed", "review": "reviewed"}
    results = []

    for item in items:
        job_id = item.get("id")
        action = item.get("action")

        if not job_id or action not in valid_actions:
            results.append({"id": job_id, "error": f"invalid id or action '{action}'"})
            continue

        job = db.query_one("SELECT * FROM fresh_jobs WHERE id = %s", (job_id,))
        if not job:
            results.append({"id": job_id, "error": "not found"})
            continue

        if action == "save":
            if job.get("saved_job_id"):
                results.append({"id": job_id, "error": "already saved", "saved_job_id": job["saved_job_id"]})
                continue
            saved = db.execute_returning(
                """
                INSERT INTO saved_jobs (title, company, url, location, salary_range, jd_text, source, status)
                VALUES (%s, %s, %s, %s, %s, %s, %s, 'saved')
                RETURNING id
                """,
                (
                    job.get("title"),
                    job.get("company"),
                    job.get("source_url"),
                    job.get("location"),
                    job.get("salary_range"),
                    job.get("jd_full") or job.get("jd_snippet"),
                    job.get("source_type"),
                ),
            )
            saved_job_id = saved["id"]
            db.execute(
                "UPDATE fresh_jobs SET status = 'saved', saved_job_id = %s WHERE id = %s",
                (saved_job_id, job_id),
            )
            results.append({"id": job_id, "action": "save", "saved_job_id": saved_job_id, "status": "saved"})
        else:
            new_status = status_map[action]
            db.execute("UPDATE fresh_jobs SET status = %s WHERE id = %s", (new_status, job_id))
            results.append({"id": job_id, "action": action, "status": new_status})

    saved_count = sum(1 for r in results if r.get("action") == "save")
    passed_count = sum(1 for r in results if r.get("action") == "pass")
    error_count = sum(1 for r in results if "error" in r)

    return {
        "total": len(items),
        "saved": saved_count,
        "passed": passed_count,
        "errors": error_count,
        "results": results,
    }


# ---------------------------------------------------------------------------
# Application Aging & Link Monitoring tools
# ---------------------------------------------------------------------------


@mcp.tool()
def get_stale_applications(days: int = 14, status: str | None = None) -> dict:
    """Get applications that haven't had a status change in X days.

    Args:
        days: Number of days without activity to be considered stale (default 14)
        status: Optional status filter (e.g. 'Applied', 'Phone Screen')

    Returns:
        dict with count and stale applications list
    """
    clauses = ["a.last_status_change < NOW() - INTERVAL '%s days'"]
    params: list = [days]

    if status:
        clauses.append("a.status = %s")
        params.append(status)

    where = "WHERE " + " AND ".join(clauses)

    rows = db.query(
        f"""
        SELECT a.id, a.company_name, a.role, a.status, a.date_applied,
               a.last_status_change, a.link_status, a.posting_closed,
               a.jd_url,
               EXTRACT(DAY FROM NOW() - a.last_status_change)::int AS days_stale
        FROM applications a
        {where}
        ORDER BY a.last_status_change ASC
        """,
        params,
    )
    return {"count": len(rows), "applications": rows}


@mcp.tool()
def check_posting_status(entity_type: str, entity_id: int) -> dict:
    """Check and update the posting/link status of a saved job or application.

    Note: This marks the record with the current check timestamp. Actual HTTP
    link checking is deferred to a cron job or manual trigger -- this tool
    reads and returns the stored status, and stamps last_link_check_at = NOW().

    Args:
        entity_type: 'saved_job' or 'application'
        entity_id: The record ID

    Returns:
        dict with current link status info
    """
    if entity_type == "saved_job":
        row = db.execute_returning(
            """
            UPDATE saved_jobs
            SET last_link_check_at = NOW()
            WHERE id = %s
            RETURNING id, title, company, url AS jd_url,
                      link_status, posting_closed, posting_closed_at,
                      last_link_check_at
            """,
            (entity_id,),
        )
    elif entity_type == "application":
        row = db.execute_returning(
            """
            UPDATE applications
            SET last_link_check_at = NOW()
            WHERE id = %s
            RETURNING id, role, company_name AS company, jd_url,
                      link_status, posting_closed, posting_closed_at,
                      last_link_check_at
            """,
            (entity_id,),
        )
    else:
        return {"error": "entity_type must be 'saved_job' or 'application'"}

    if not row:
        return {"error": f"{entity_type} {entity_id} not found"}

    return {"entity_type": entity_type, **row}


@mcp.tool()
def get_aging_summary() -> dict:
    """Get a summary of application aging across the pipeline.

    Returns:
        dict with counts: stale_applications, stale_saved_jobs,
        closed_postings, needs_link_check
    """
    stale_apps = db.query_one(
        """
        SELECT COUNT(*)::int AS count
        FROM applications
        WHERE last_status_change < NOW() - INTERVAL '14 days'
        """
    )
    stale_jobs = db.query_one(
        """
        SELECT COUNT(*)::int AS count
        FROM saved_jobs
        WHERE updated_at < NOW() - INTERVAL '30 days'
          AND status NOT IN ('applied', 'archived')
        """
    )
    closed_apps = db.query_one(
        "SELECT COUNT(*)::int AS count FROM applications WHERE posting_closed = TRUE"
    )
    closed_jobs = db.query_one(
        "SELECT COUNT(*)::int AS count FROM saved_jobs WHERE posting_closed = TRUE"
    )
    unknown_apps = db.query_one(
        "SELECT COUNT(*)::int AS count FROM applications WHERE link_status = 'unknown'"
    )
    unknown_jobs = db.query_one(
        "SELECT COUNT(*)::int AS count FROM saved_jobs WHERE link_status = 'unknown'"
    )

    return {
        "stale_applications": stale_apps["count"] if stale_apps else 0,
        "stale_saved_jobs": stale_jobs["count"] if stale_jobs else 0,
        "closed_postings": (
            (closed_apps["count"] if closed_apps else 0)
            + (closed_jobs["count"] if closed_jobs else 0)
        ),
        "needs_link_check": (
            (unknown_apps["count"] if unknown_apps else 0)
            + (unknown_jobs["count"] if unknown_jobs else 0)
        ),
    }


# ---------------------------------------------------------------------------
# CRM tools
# ---------------------------------------------------------------------------


@mcp.tool()
def update_relationship_stage(contact_id: int, stage: str) -> dict:
    """Update a contact's relationship stage.

    Args:
        contact_id: Contact ID
        stage: cold, warm, active, close, or dormant

    Returns:
        dict with updated contact info
    """
    valid_stages = {"cold", "warm", "active", "close", "dormant"}
    if stage not in valid_stages:
        return {"error": f"stage must be one of: {', '.join(sorted(valid_stages))}"}

    row = db.execute_returning(
        """
        UPDATE contacts
        SET relationship_stage = %s, updated_at = NOW()
        WHERE id = %s
        RETURNING id, name, company, title, relationship_stage, health_score, updated_at
        """,
        (stage, contact_id),
    )
    if not row:
        return {"error": f"Contact {contact_id} not found"}
    return {"contact": row}


@mcp.tool()
def get_relationship_health(contact_id: int | None = None, limit: int = 20) -> dict:
    """Get contacts ranked by relationship health score (lowest = needs attention).

    Args:
        contact_id: Optional specific contact ID. If None, returns lowest-health contacts.
        limit: Max results (default 20)

    Returns:
        dict with contacts and their health scores
    """
    if contact_id is not None:
        row = db.query_one(
            """
            SELECT id, name, company, title, relationship_stage, health_score,
                   last_touchpoint_at, last_contact, tags
            FROM contacts
            WHERE id = %s
            """,
            (contact_id,),
        )
        if not row:
            return {"error": f"Contact {contact_id} not found"}
        return {"contact": row}

    rows = db.query(
        """
        SELECT id, name, company, title, relationship_stage, health_score,
               last_touchpoint_at, last_contact, tags
        FROM contacts
        ORDER BY health_score ASC NULLS FIRST
        LIMIT %s
        """,
        (limit,),
    )
    return {"contacts": rows, "count": len(rows)}


@mcp.tool()
def log_touchpoint(
    contact_id: int,
    type: str,
    channel: str = "email",
    direction: str = "outbound",
    notes: str | None = None,
) -> dict:
    """Log a touchpoint (interaction) with a contact. Updates last_touchpoint_at automatically.

    Args:
        contact_id: Contact ID
        type: email, linkedin_message, phone_call, coffee, meeting, event, referral
        channel: linkedin, email, phone, in_person, slack, other
        direction: inbound or outbound
        notes: Optional notes about the interaction

    Returns:
        dict with created touchpoint
    """
    contact = db.query_one("SELECT id FROM contacts WHERE id = %s", (contact_id,))
    if not contact:
        return {"error": f"Contact {contact_id} not found"}

    row = db.execute_returning(
        """
        INSERT INTO touchpoints (contact_id, type, channel, direction, notes)
        VALUES (%s, %s, %s, %s, %s)
        RETURNING *
        """,
        (contact_id, type, channel, direction, notes),
    )

    db.execute(
        """
        UPDATE contacts
        SET last_touchpoint_at = NOW(),
            last_contact = CURRENT_DATE,
            updated_at = NOW()
        WHERE id = %s
        """,
        (contact_id,),
    )

    return {"touchpoint": row}


@mcp.tool()
def get_networking_tasks(status: str = "pending", days: int = 7) -> dict:
    """Get upcoming networking tasks.

    Args:
        status: 'pending' (incomplete), 'completed', or 'overdue'
        days: For pending, show tasks due within this many days (default 7)

    Returns:
        dict with tasks list
    """
    if status == "completed":
        rows = db.query(
            """
            SELECT nt.id, nt.contact_id, nt.task_type, nt.title, nt.due_date,
                   nt.completed_at, nt.notes, nt.created_at,
                   c.name AS contact_name, c.company AS contact_company
            FROM networking_tasks nt
            JOIN contacts c ON c.id = nt.contact_id
            WHERE nt.completed = TRUE
            ORDER BY nt.completed_at DESC
            LIMIT 50
            """,
        )
    elif status == "overdue":
        today = date.today()
        rows = db.query(
            """
            SELECT nt.id, nt.contact_id, nt.task_type, nt.title, nt.due_date,
                   nt.notes, nt.created_at,
                   c.name AS contact_name, c.company AS contact_company
            FROM networking_tasks nt
            JOIN contacts c ON c.id = nt.contact_id
            WHERE nt.completed = FALSE
              AND nt.due_date < %s
            ORDER BY nt.due_date ASC
            """,
            (today,),
        )
    else:
        cutoff = date.today() + timedelta(days=days)
        rows = db.query(
            """
            SELECT nt.id, nt.contact_id, nt.task_type, nt.title, nt.due_date,
                   nt.notes, nt.created_at,
                   c.name AS contact_name, c.company AS contact_company
            FROM networking_tasks nt
            JOIN contacts c ON c.id = nt.contact_id
            WHERE nt.completed = FALSE
              AND (nt.due_date IS NULL OR nt.due_date <= %s)
            ORDER BY nt.due_date ASC NULLS LAST, nt.created_at ASC
            """,
            (cutoff,),
        )

    return {"tasks": rows, "count": len(rows), "status": status}


# ---------------------------------------------------------------------------
# Workflow Automation tools
# ---------------------------------------------------------------------------


@mcp.tool()
def get_workflows(enabled: bool | None = None, trigger_type: str | None = None) -> dict:
    """List all workflows, optionally filtered.

    Args:
        enabled: Filter by enabled status
        trigger_type: Filter by trigger type (schedule, event)

    Returns:
        dict with workflows list
    """
    clauses, params = [], []
    if enabled is not None:
        clauses.append("enabled = %s")
        params.append(enabled)
    if trigger_type:
        clauses.append("trigger_type = %s")
        params.append(trigger_type)

    where = f"WHERE {' AND '.join(clauses)}" if clauses else ""
    rows = db.query(
        f"SELECT * FROM workflows {where} ORDER BY created_at DESC",
        params or None,
    )
    return {"workflows": rows, "count": len(rows)}


@mcp.tool()
def trigger_workflow(workflow_id: int, trigger_data: str | None = None) -> dict:
    """Manually trigger a workflow and execute its action.

    Args:
        workflow_id: Workflow ID to trigger
        trigger_data: Optional JSON string with trigger context data

    Returns:
        dict with execution result
    """
    import json as _json

    workflow = db.query_one("SELECT * FROM workflows WHERE id = %s", (workflow_id,))
    if not workflow:
        return {"error": f"Workflow {workflow_id} not found"}

    parsed_trigger_data = None
    if trigger_data:
        try:
            parsed_trigger_data = _json.loads(trigger_data)
        except _json.JSONDecodeError:
            return {"error": "trigger_data is not valid JSON"}

    action_type = workflow["action_type"]
    action_config = workflow["action_config"] or {}
    action_result = {}
    success = True
    error_message = None

    try:
        if action_type == "create_notification":
            notif = db.execute_returning(
                """
                INSERT INTO notifications (type, severity, title, body)
                VALUES (%s, %s, %s, %s)
                RETURNING id, type, severity, title, body, created_at
                """,
                (
                    action_config.get("type", "workflow"),
                    action_config.get("severity", "info"),
                    action_config.get("title", f"Workflow: {workflow['name']}"),
                    action_config.get("body", ""),
                ),
            )
            action_result = {"action": "create_notification", "notification": notif}

        elif action_type == "log_activity":
            try:
                db.execute(
                    """
                    INSERT INTO activity_log (entity_type, entity_id, action, details)
                    VALUES (%s, %s, %s, %s::jsonb)
                    """,
                    (
                        action_config.get("entity_type", "workflow"),
                        action_config.get("entity_id"),
                        action_config.get("action", workflow["name"]),
                        _json.dumps(action_config.get("details", {})),
                    ),
                )
                action_result = {"action": "log_activity", "logged": True}
            except Exception as inner_err:
                action_result = {
                    "action": "log_activity",
                    "logged": False,
                    "note": "activity_log table unavailable; logged to workflow_log only",
                    "detail": str(inner_err),
                }

        elif action_type == "update_field":
            action_result = {
                "action": "update_field",
                "simulated": True,
                "would_update": {
                    "table": action_config.get("table"),
                    "field": action_config.get("field"),
                    "value": action_config.get("value"),
                    "where": action_config.get("where"),
                },
                "note": "update_field is simulated in MVP — no DB write performed",
            }

        else:
            action_result = {"action": action_type, "error": f"Unknown action_type: {action_type}"}
            success = False
            error_message = f"Unknown action_type: {action_type}"

    except Exception as exc:
        action_result = {"action": action_type, "error": str(exc)}
        success = False
        error_message = str(exc)

    log_row = db.execute_returning(
        """
        INSERT INTO workflow_log (workflow_id, trigger_data, action_result, success, error_message)
        VALUES (%s, %s::jsonb, %s::jsonb, %s, %s)
        RETURNING *
        """,
        (
            workflow_id,
            _json.dumps(parsed_trigger_data) if parsed_trigger_data is not None else None,
            _json.dumps(action_result),
            success,
            error_message,
        ),
    )

    db.execute("UPDATE workflows SET last_run_at = NOW() WHERE id = %s", (workflow_id,))

    return {
        "workflow_id": workflow_id,
        "workflow_name": workflow["name"],
        "success": success,
        "result": action_result,
        "log_id": log_row["id"] if log_row else None,
        "error": error_message,
    }


@mcp.tool()
def create_workflow(
    name: str,
    trigger_type: str,
    trigger_config: str,
    action_type: str,
    action_config: str,
    conditions: str | None = None,
) -> dict:
    """Create a new workflow automation.

    Args:
        name: Workflow name
        trigger_type: schedule or event
        trigger_config: JSON string with trigger configuration
        action_type: create_notification, update_field, or log_activity
        action_config: JSON string with action configuration
        conditions: Optional JSON string with conditions

    Returns:
        dict with created workflow
    """
    import json as _json

    try:
        parsed_trigger_config = _json.loads(trigger_config)
    except _json.JSONDecodeError:
        return {"error": "trigger_config is not valid JSON"}

    try:
        parsed_action_config = _json.loads(action_config)
    except _json.JSONDecodeError:
        return {"error": "action_config is not valid JSON"}

    parsed_conditions = None
    if conditions:
        try:
            parsed_conditions = _json.loads(conditions)
        except _json.JSONDecodeError:
            return {"error": "conditions is not valid JSON"}

    valid_trigger_types = ("schedule", "event")
    if trigger_type not in valid_trigger_types:
        return {"error": f"trigger_type must be one of: {', '.join(valid_trigger_types)}"}

    valid_action_types = ("create_notification", "update_field", "log_activity")
    if action_type not in valid_action_types:
        return {"error": f"action_type must be one of: {', '.join(valid_action_types)}"}

    row = db.execute_returning(
        """
        INSERT INTO workflows (name, trigger_type, trigger_config, conditions, action_type, action_config, enabled)
        VALUES (%s, %s, %s::jsonb, %s::jsonb, %s, %s::jsonb, TRUE)
        RETURNING *
        """,
        (
            name,
            trigger_type,
            _json.dumps(parsed_trigger_config),
            _json.dumps(parsed_conditions) if parsed_conditions is not None else None,
            action_type,
            _json.dumps(parsed_action_config),
        ),
    )
    return {"workflow": row, "created": True}


# ---------------------------------------------------------------------------
# Market Intelligence tools
# ---------------------------------------------------------------------------


@mcp.tool()
def get_market_signals(
    source: str | None = None,
    signal_type: str | None = None,
    severity: str | None = None,
    industry: str | None = None,
    region: str | None = None,
    limit: int = 20,
) -> dict:
    """Get market intelligence signals, optionally filtered.

    Args:
        source: Filter by source (bls_jolts, lisep_tru, warn_act, hn_hiring, news, manual)
        signal_type: Filter by type (layoff, hiring_freeze, market_trend, job_openings,
                     separation_rate, quit_rate, wage_data)
        severity: Filter by severity (positive, neutral, negative, critical)
        industry: Filter by industry sector
        region: Filter by region
        limit: Max results (default 20)

    Returns:
        dict with count and signals list
    """
    clauses, params = [], []
    if source:
        clauses.append("source = %s")
        params.append(source)
    if signal_type:
        clauses.append("signal_type = %s")
        params.append(signal_type)
    if severity:
        clauses.append("severity = %s")
        params.append(severity)
    if industry:
        clauses.append("industry ILIKE %s")
        params.append(f"%{industry}%")
    if region:
        clauses.append("region ILIKE %s")
        params.append(f"%{region}%")

    where = f"WHERE {' AND '.join(clauses)}" if clauses else ""
    rows = db.query(
        f"""
        SELECT *
        FROM market_signals
        {where}
        ORDER BY captured_at DESC
        LIMIT %s
        """,
        params + [limit],
    )
    return {"count": len(rows), "signals": rows}


@mcp.tool()
def add_market_signal(
    source: str,
    signal_type: str,
    title: str,
    severity: str = "neutral",
    body: str | None = None,
    data_json: str | None = None,
    region: str | None = None,
    industry: str | None = None,
    source_url: str | None = None,
) -> dict:
    """Add a market intelligence signal.

    Args:
        source: Signal source (bls_jolts, lisep_tru, warn_act, hn_hiring, news, manual)
        signal_type: Type of signal (layoff, hiring_freeze, market_trend, job_openings,
                     separation_rate, quit_rate, wage_data)
        title: Signal title/headline
        severity: positive, neutral, negative, or critical (default: neutral)
        body: Detailed description
        data_json: Optional JSON string with structured data specific to the signal type
        region: Geographic region
        industry: Industry sector
        source_url: URL to the original source

    Returns:
        dict with created signal
    """
    import json as _json
    if data_json is not None:
        try:
            _json.loads(data_json)
        except (_json.JSONDecodeError, TypeError):
            return {"error": "data_json must be a valid JSON string"}

    row = db.execute_returning(
        """
        INSERT INTO market_signals
            (source, signal_type, title, body, data_json, region, industry,
             severity, source_url)
        VALUES (%s, %s, %s, %s, %s::jsonb, %s, %s, %s, %s)
        RETURNING *
        """,
        (source, signal_type, title, body, data_json, region, industry, severity, source_url),
    )
    return row or {"error": "Insert failed"}


@mcp.tool()
def get_market_summary() -> dict:
    """Get a summary of current market intelligence.

    Returns:
        dict with counts by source/type/severity, plus recent highlights (last 7 days)
    """
    by_source = db.query(
        "SELECT source, COUNT(*) AS count FROM market_signals GROUP BY source ORDER BY count DESC"
    )
    by_type = db.query(
        "SELECT signal_type, COUNT(*) AS count FROM market_signals GROUP BY signal_type ORDER BY count DESC"
    )
    by_severity = db.query(
        "SELECT severity, COUNT(*) AS count FROM market_signals GROUP BY severity ORDER BY count DESC"
    )
    recent_highlights = db.query(
        """
        SELECT id, source, signal_type, title, severity, region, industry, captured_at
        FROM market_signals
        WHERE captured_at >= NOW() - INTERVAL '7 days'
        ORDER BY captured_at DESC
        LIMIT 20
        """
    )
    total = db.query_one("SELECT COUNT(*) AS total FROM market_signals")

    return {
        "total": total["total"] if total else 0,
        "by_source": by_source,
        "by_signal_type": by_type,
        "by_severity": by_severity,
        "recent_highlights": recent_highlights,
    }


# ---------------------------------------------------------------------------
# Mock Interviews (0_APP 8.5)
# ---------------------------------------------------------------------------

@mcp.tool()
def create_mock_interview(
    job_title: str,
    company: str,
    interview_type: str = "behavioral",
    difficulty: str = "medium",
    application_id: int | None = None,
) -> dict:
    """Create a mock interview session with generated questions.

    Args:
        job_title: Job title to tailor questions for
        company: Target company name
        interview_type: behavioral, technical, situational, case, or mixed
        difficulty: easy, medium, or hard
        application_id: Optional linked application ID
    """
    from mcp_tools_mock_interviews import create_mock_interview as _impl
    return _impl(job_title, company, interview_type, difficulty, application_id)


@mcp.tool()
def get_mock_interview(interview_id: int) -> dict:
    """Get a mock interview session with all questions and scores.

    Args:
        interview_id: ID of the mock interview to retrieve
    """
    from mcp_tools_mock_interviews import get_mock_interview as _impl
    return _impl(interview_id)


@mcp.tool()
def evaluate_mock_interview(interview_id: int, answers: dict) -> dict:
    """Evaluate answers for a mock interview and generate scores/feedback.

    Args:
        interview_id: ID of the mock interview
        answers: Dict mapping question_id to answer text, e.g. {"1": "My answer..."}
    """
    from mcp_tools_mock_interviews import evaluate_mock_interview as _impl
    return _impl(interview_id, answers)


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser()
    parser.add_argument("--sse", action="store_true", help="Run as SSE server (for Docker)")
    parser.add_argument("--port", type=int, default=8056, help="SSE server port")
    args = parser.parse_args()

    if args.sse:
        mcp.settings.port = args.port
        mcp.run(transport="sse")
    else:
        mcp.run(transport="stdio")
