"""Generate a .docx resume from a placeholder template + DB content.

Loads the placeholder template (with {{SLOT}} markers), builds a content
map from the DB spec/header/education/certifications, fills each slot,
and applies bold_label formatting where needed.

Usage:
    # Generate V32 base from DB
    python generate_resume.py --version v32 --output Output/V32_generated.docx

    # Generate and compare against original
    python generate_resume.py --version v32 --output Output/V32_generated.docx \
        --compare Originals/Stephen_Salaka_Resume_v32.docx

    # With tailoring overrides
    python generate_resume.py --version v32 --spec overrides.json \
        --output Output/tailored.docx

Dependencies:
    pip install python-docx psycopg2-binary
"""

import argparse
import io
import json
import os
import re
import sys
from copy import deepcopy
from pathlib import Path
from typing import Any, Optional

import psycopg2
from docx import Document
from docx.text.paragraph import Paragraph


# -- DB connection -----------------------------------------------------------

def get_db_connection():
    """Connect to SuperTroopers PostgreSQL."""
    return psycopg2.connect(
        host=os.environ.get("DB_HOST", "localhost"),
        port=int(os.environ.get("DB_PORT", "5555")),
        dbname=os.environ.get("DB_NAME", "supertroopers"),
        user=os.environ.get("DB_USER", "supertroopers"),
        password=os.environ.get("DB_PASSWORD", "WUHD8fBisb57FS4Q3bdvfuvgnim9fL1c"),
    )


# -- Data fetchers -----------------------------------------------------------

def fetch_template_blob(conn, template_name: str = "V32 Base") -> bytes:
    """Fetch the placeholder template .docx blob from resume_templates."""
    with conn.cursor() as cur:
        cur.execute(
            "SELECT template_blob FROM resume_templates WHERE name = %s AND is_active = true",
            (template_name,),
        )
        row = cur.fetchone()
        if not row:
            raise ValueError(f"Template '{template_name}' not found in DB")
        return bytes(row[0])


def fetch_template_map(conn, template_name: str = "V32 Base") -> dict:
    """Fetch the template map JSONB from resume_templates."""
    with conn.cursor() as cur:
        cur.execute(
            "SELECT template_map FROM resume_templates WHERE name = %s AND is_active = true",
            (template_name,),
        )
        row = cur.fetchone()
        if not row or not row[0]:
            return {}
        return row[0] if isinstance(row[0], dict) else json.loads(row[0])


def fetch_resume_spec(conn, version: str = "v32", variant: str = "base") -> dict:
    """Fetch the resume spec JSONB from resume_versions."""
    with conn.cursor() as cur:
        cur.execute(
            "SELECT spec FROM resume_versions WHERE version = %s AND variant = %s AND spec IS NOT NULL",
            (version, variant),
        )
        row = cur.fetchone()
        if not row:
            raise ValueError(f"Spec not found for {version}/{variant}")
        return row[0] if isinstance(row[0], dict) else json.loads(row[0])


def fetch_header(conn) -> dict:
    """Fetch candidate header info."""
    with conn.cursor() as cur:
        cur.execute(
            "SELECT full_name, credentials, location, location_note, "
            "email, phone, linkedin_url FROM resume_header LIMIT 1"
        )
        row = cur.fetchone()
        if not row:
            return {}
        cols = ["full_name", "credentials", "location", "location_note",
                "email", "phone", "linkedin_url"]
        return dict(zip(cols, row))


def fetch_education(conn) -> list[dict]:
    """Fetch education entries in order."""
    with conn.cursor() as cur:
        cur.execute(
            "SELECT degree, field, institution, location "
            "FROM education ORDER BY sort_order"
        )
        return [
            {"degree": r[0], "field": r[1], "institution": r[2], "location": r[3]}
            for r in cur.fetchall()
        ]


def fetch_certifications(conn) -> list[dict]:
    """Fetch certifications in order."""
    with conn.cursor() as cur:
        cur.execute(
            "SELECT name, issuer FROM certifications "
            "WHERE is_active = true ORDER BY sort_order"
        )
        return [{"name": r[0], "issuer": r[1]} for r in cur.fetchall()]


def fetch_career_history(conn, employers: list[str]) -> dict[str, dict]:
    """Fetch career_history rows for the given employers, keyed by employer name."""
    result = {}
    with conn.cursor() as cur:
        for emp in employers:
            cur.execute(
                "SELECT employer, title, start_date, end_date, location, "
                "industry, is_current, intro_text "
                "FROM career_history WHERE employer ILIKE %s ORDER BY start_date DESC LIMIT 1",
                (f"%{emp}%",),
            )
            row = cur.fetchone()
            if row:
                cols = ["employer", "title", "start_date", "end_date", "location",
                        "industry", "is_current", "intro_text"]
                result[emp] = dict(zip(cols, row))
    return result


# -- Recipe resolution -------------------------------------------------------

ALLOWED_TABLES = {
    "bullets", "career_history", "skills", "summary_variants",
    "education", "certifications", "resume_header",
}


def fetch_recipe(conn, recipe_id: int) -> dict:
    """Fetch a recipe row from resume_recipes."""
    with conn.cursor() as cur:
        cur.execute(
            "SELECT id, name, headline, template_id, recipe FROM resume_recipes WHERE id = %s",
            (recipe_id,),
        )
        row = cur.fetchone()
        if not row:
            raise ValueError(f"Recipe id={recipe_id} not found")
        cols = ["id", "name", "headline", "template_id", "recipe"]
        return dict(zip(cols, row))


def _resolve_single(cur, table: str, row_id: int, column: str | None, slot_name: str) -> str:
    """Resolve a single {table, id, column} reference to text."""
    if table not in ALLOWED_TABLES:
        raise ValueError(f"Table '{table}' not in allowed list for slot {slot_name}")

    # Assembly patterns (no column specified)
    if column is None or column == "":
        if table == "career_history":
            cur.execute(
                "SELECT employer, location, industry FROM career_history WHERE id = %s",
                (row_id,),
            )
            row = cur.fetchone()
            if not row:
                return ""
            parts = [row[0]]
            if row[1]:
                parts.append(f", {row[1]}")
            if row[2]:
                parts.append(f" {{{row[2]}}}")
            return "".join(parts)

        elif table == "education":
            cur.execute(
                "SELECT degree, field, institution, location FROM education WHERE id = %s",
                (row_id,),
            )
            row = cur.fetchone()
            if not row:
                return ""
            parts = []
            if row[0]:
                parts.append(row[0])
            if row[1]:
                parts.append(row[1])
            result = ", ".join(parts)
            if row[2]:
                result += f" | {row[2]}"
            if row[3]:
                result += f" \u2014 {row[3]}"  # em dash before location
            return result

        elif table == "certifications":
            cur.execute(
                "SELECT name, issuer FROM certifications WHERE id = %s",
                (row_id,),
            )
            row = cur.fetchone()
            if not row:
                return ""
            if row[1]:
                return f"{row[0]} | {row[1]}"
            return row[0]

        else:
            raise ValueError(f"No assembly rule for table '{table}' without column")

    # Special slots for resume_header
    if table == "resume_header":
        cur.execute(
            "SELECT full_name, credentials, location, location_note, "
            "email, phone, linkedin_url FROM resume_header WHERE id = %s",
            (row_id,),
        )
        row = cur.fetchone()
        if not row:
            return ""
        header = dict(zip(
            ["full_name", "credentials", "location", "location_note",
             "email", "phone", "linkedin_url"], row
        ))
        if column == "name" or slot_name == "HEADER_NAME":
            return f"{header['full_name']}, {header['credentials']}"
        elif column == "contact" or slot_name == "HEADER_CONTACT":
            parts = [header["location"]]
            if header.get("location_note"):
                parts[0] += f" ({header['location_note']})"
            parts.append(header["email"])
            parts.append(header["phone"])
            if header.get("linkedin_url"):
                parts.append(header["linkedin_url"])
            return " \u2022 ".join(parts)
        # Fall through to generic column fetch

    # Generic column fetch
    cur.execute(f"SELECT {column} FROM {table} WHERE id = %s", (row_id,))
    row = cur.fetchone()
    return row[0] if row and row[0] else ""


def resolve_recipe(conn, recipe_json: dict) -> dict[str, str]:
    """Resolve a recipe JSON into a content_map (placeholder -> text).

    Each entry in recipe_json maps a slot name to one of:
      - {table, id, column}   -> single value lookup
      - {table, ids, column}  -> array lookup, join with ' | '
      - {table, id}           -> assembly (no column)
      - {literal: "text"}     -> text as-is
      - {table: "resume_header", id, slot: "name"|"contact"} -> header assembly
    """
    content = {}
    with conn.cursor() as cur:
        for slot_name, ref in recipe_json.items():
            if "literal" in ref:
                content[slot_name] = ref["literal"]

            elif "ids" in ref:
                # Array reference
                table = ref["table"]
                ids = ref["ids"]
                column = ref.get("column", "name")
                if table not in ALLOWED_TABLES:
                    raise ValueError(f"Table '{table}' not allowed for slot {slot_name}")
                placeholders = ",".join(["%s"] * len(ids))
                # Preserve order by using array position
                cur.execute(
                    f"SELECT id, {column} FROM {table} WHERE id IN ({placeholders})",
                    ids,
                )
                rows = {r[0]: r[1] for r in cur.fetchall()}
                values = [rows.get(i, "") for i in ids]
                content[slot_name] = " | ".join(v for v in values if v)

            elif "table" in ref:
                # Single value or assembly
                table = ref["table"]
                row_id = ref.get("id", 1)
                column = ref.get("column") or ref.get("slot")
                content[slot_name] = _resolve_single(cur, table, row_id, column, slot_name)

    return content


def validate_recipe(conn, recipe_json: dict) -> dict:
    """Validate a recipe JSON by checking all referenced IDs exist in DB.

    Returns a dict with:
      - valid: bool
      - errors: list of {slot, table, id, error} for missing references
      - warnings: list of {slot, message} for non-critical issues
      - stats: {total_slots, db_refs, literals, valid_refs, missing_refs}
    """
    errors = []
    warnings = []
    db_refs = 0
    literals = 0
    valid_refs = 0

    with conn.cursor() as cur:
        for slot_name, ref in recipe_json.items():
            if "literal" in ref:
                literals += 1
                continue

            if "ids" in ref:
                # Array reference
                table = ref["table"]
                if table not in ALLOWED_TABLES:
                    errors.append({"slot": slot_name, "table": table, "error": f"Table '{table}' not in allowed list"})
                    continue
                ids = ref["ids"]
                db_refs += len(ids)
                cur.execute(f"SELECT id FROM {table} WHERE id = ANY(%s)", (ids,))
                found = {r[0] for r in cur.fetchall()}
                missing = [i for i in ids if i not in found]
                if missing:
                    errors.append({"slot": slot_name, "table": table, "ids": missing, "error": f"IDs not found: {missing}"})
                else:
                    valid_refs += len(ids)

            elif "table" in ref:
                table = ref["table"]
                if table not in ALLOWED_TABLES:
                    errors.append({"slot": slot_name, "table": table, "error": f"Table '{table}' not in allowed list"})
                    continue
                row_id = ref.get("id", 1)
                db_refs += 1
                cur.execute(f"SELECT id FROM {table} WHERE id = %s", (row_id,))
                if cur.fetchone():
                    valid_refs += 1
                else:
                    errors.append({"slot": slot_name, "table": table, "id": row_id, "error": f"Row not found"})

    return {
        "valid": len(errors) == 0,
        "errors": errors,
        "warnings": warnings,
        "stats": {
            "total_slots": len(recipe_json),
            "db_refs": db_refs,
            "literals": literals,
            "valid_refs": valid_refs,
            "missing_refs": len(errors),
        },
    }


# -- Paragraph formatting ---------------------------------------------------

PLACEHOLDER_RE = re.compile(r"\{\{([A-Z0-9_]+)\}\}")


def _run_has_shape(run) -> bool:
    """Check if a run contains VML shapes or drawings that must be preserved."""
    mc_ns = "http://schemas.openxmlformats.org/markup-compatibility/2006"
    w_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    return (
        run._element.find(f"{{{mc_ns}}}AlternateContent") is not None
        or run._element.find(f"{{{w_ns}}}pict") is not None
        or run._element.find(f"{{{w_ns}}}drawing") is not None
    )


def _fill_simple(paragraph: Paragraph, text: str, clear_bold: bool = False) -> None:
    """Replace placeholder with text. Preserves runs containing shapes/drawings.

    Args:
        clear_bold: If True, reset bold to None on all text runs (used when
                    bold_label separator wasn't found and text should not be bold).
    """
    if paragraph.runs:
        text_set = False
        for run in paragraph.runs:
            if _run_has_shape(run):
                continue
            if not text_set:
                run.text = text
                if clear_bold:
                    run.bold = None
                text_set = True
            else:
                run.text = ""
                if clear_bold:
                    run.bold = None
        if not text_set and paragraph.runs:
            paragraph.runs[0].text = text
    else:
        paragraph.text = text


def _find_transition_split(text: str) -> int:
    """Find the split point for V31-style bold bullets (action phrase + transition word).

    Returns the index where the non-bold text starts, or -1 if not found.
    """
    transitions = [
        ", leading to ", ", resulting in ", ", enabling ", ", achieving ",
        ", supporting ", ", accelerating ", ", expanding ", ", which ",
        " by ", " through ", " while ", " with ",
    ]
    search_limit = min(len(text), int(len(text) * 0.6))
    best_idx = -1
    for t in transitions:
        idx = text.find(t, 20, search_limit)
        if idx >= 0 and (best_idx < 0 or idx < best_idx):
            best_idx = idx
            if t.startswith(","):
                best_idx = idx + 1
    return best_idx


def _fill_bold_label(paragraph: Paragraph, text: str, separator: str = ": ") -> None:
    """Replace placeholder with bold label + non-bold body.

    If separator is "transition", uses transition-word detection (V31 style).
    Otherwise splits at the literal separator string.
    Preserves runs containing shapes/drawings.
    """
    if separator == "transition":
        idx = _find_transition_split(text)
    else:
        idx = text.find(separator)

    if idx < 0 or not paragraph.runs:
        _fill_simple(paragraph, text, clear_bold=True)
        return

    bold_part = text[:idx]
    non_bold_part = text[idx:]

    text_runs = [r for r in paragraph.runs if not _run_has_shape(r)]
    if not text_runs:
        _fill_simple(paragraph, text)
        return

    text_runs[0].text = bold_part
    text_runs[0].bold = True
    if len(text_runs) > 1:
        text_runs[1].text = non_bold_part
        text_runs[1].bold = None
        for run in text_runs[2:]:
            run.text = ""
            run.bold = None
    else:
        text_runs[0].text = text


# -- Bold separator by slot type --------------------------------------------
# V32 uses colon separators, V31 uses transition words for bullets

BOLD_SEPARATORS = {
    "highlight": ": ",
    "job_bullet": ": ",
    "education": " | ",
    "certification": " | ",
    "additional_exp": " | ",
    "ref_link": " | ",
    "job_header": ", ",
}


# -- Content map builder ----------------------------------------------------

def build_content_map(
    spec: dict,
    header: dict,
    education: list[dict],
    certifications: list[dict],
    career: dict[str, dict],
    template_map: dict,
) -> dict[str, str]:
    """Build placeholder_name → content_text mapping from DB data.

    Strategy: start with template_map original_text as the base (guarantees
    exact reproduction for unchanged slots), then overlay spec-derived content
    for slots that should be dynamic (summary, bullets, keywords, etc.).

    For tailoring, callers can further override specific slots after this
    function returns.
    """
    content = {}

    # Step 1: Seed from template_map original_text (exact base reconstruction)
    for slot in template_map.get("slots", []):
        placeholder = slot.get("placeholder")
        original = slot.get("original_text")
        if placeholder and original:
            content[placeholder] = original

    # Step 2: Overlay with spec-derived content for dynamic slots

    # -- Header --
    if header:
        content["HEADER_NAME"] = f"{header['full_name']}, {header['credentials']}"
        parts = [header["location"]]
        if header.get("location_note"):
            parts[0] += f" ({header['location_note']})"
        parts.append(header["email"])
        parts.append(header["phone"])
        if header.get("linkedin_url"):
            parts.append(header["linkedin_url"])
        content["HEADER_CONTACT"] = " \u2022 ".join(parts)

    # -- Headline & Summary --
    if "headline" in spec:
        content["HEADLINE"] = spec["headline"]
    if "summary_text" in spec:
        content["SUMMARY"] = spec["summary_text"]

    # -- Highlight bullets --
    for i, bullet in enumerate(spec.get("highlight_bullets", []), 1):
        content[f"HIGHLIGHT_{i}"] = bullet

    # -- Top keywords --
    if "keywords" in spec:
        content["KEYWORDS"] = " | ".join(spec["keywords"])

    # -- Experience blocks --
    employers = spec.get("experience_employers", [])
    exp_bullets = spec.get("experience_bullets", {})

    for job_n, emp_name in enumerate(employers, 1):
        emp_data = career.get(emp_name, {})
        bullets_raw = exp_bullets.get(emp_name, [])

        # Job intro — from career_history.intro_text or first long entry in spec bullets
        if emp_data.get("intro_text"):
            content[f"JOB_{job_n}_INTRO"] = emp_data["intro_text"]
        elif bullets_raw and len(bullets_raw[0]) > 200 and ": " not in bullets_raw[0][:80]:
            content[f"JOB_{job_n}_INTRO"] = bullets_raw[0]

        # Job bullets — filter out intro text and subtitles
        bullet_texts = []
        subtitle_texts = set()
        for sub_key in [f"JOB_{job_n}_SUBTITLE_1", f"JOB_{job_n}_SUBTITLE_2"]:
            if sub_key in content:
                subtitle_texts.add(content[sub_key])

        for b in bullets_raw:
            if b == content.get(f"JOB_{job_n}_INTRO"):
                continue
            if b in subtitle_texts:
                continue
            bullet_texts.append(b)

        for i, bt in enumerate(bullet_texts, 1):
            content[f"JOB_{job_n}_BULLET_{i}"] = bt

    # -- Executive & Technical keywords --
    if "executive_keywords" in spec:
        content["EXEC_KEYWORDS"] = " | ".join(spec["executive_keywords"])
    if "technical_keywords" in spec:
        content["TECH_KEYWORDS"] = " | ".join(spec["technical_keywords"])

    # -- References (from spec, preserving original_text numbering) --
    for i, ref_section in enumerate(spec.get("references", []), 1):
        # Only override if spec has different content than original
        for j, link in enumerate(ref_section.get("links", []), 1):
            content[f"REF_{i}_LINK_{j}"] = f"{link['text']} | {link['desc']}"

    return content


def _get_original_texts(template_map: dict, job_n: int) -> dict[str, str]:
    """Extract original_text values from template_map for a given job block."""
    result = {}
    prefix = f"JOB_{job_n}_"
    for slot in template_map.get("slots", []):
        placeholder = slot.get("placeholder", "")
        if placeholder and placeholder.startswith(prefix) and "original_text" in slot:
            result[placeholder] = slot["original_text"]
    return result


def _format_job_header(emp_data: dict) -> str:
    """Format a job header line from career_history data."""
    parts = [emp_data["employer"]]
    if emp_data.get("location"):
        parts.append(f", {emp_data['location']}")
    if emp_data.get("industry"):
        parts.append(f" {{{emp_data['industry']}}}")
    return "".join(parts)


# -- Generator --------------------------------------------------------------

def generate_resume(
    template_blob: bytes,
    content_map: dict[str, str],
    template_map: dict,
) -> Document:
    """Fill a placeholder template with content.

    Args:
        template_blob: The placeholder .docx template as bytes.
        content_map: Mapping of placeholder name to content text.
        template_map: Template map with slot types and formatting rules.

    Returns:
        A python-docx Document ready to save.
    """
    doc = Document(io.BytesIO(template_blob))
    paras = doc.paragraphs

    # Build a lookup: placeholder_name → (slot_type, formatting)
    slot_info = {}
    for slot in template_map.get("slots", []):
        if slot.get("placeholder"):
            slot_info[slot["placeholder"]] = {
                "slot_type": slot.get("slot_type", ""),
                "formatting": slot.get("formatting", {}),
            }

    # Fill each paragraph that has a placeholder
    filled = 0
    skipped = []
    nsmap_w = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

    for para in paras:
        match = PLACEHOLDER_RE.search(para.text)
        if not match:
            continue

        placeholder = match.group(1)
        if placeholder not in content_map:
            skipped.append(placeholder)
            _fill_simple(para, "")  # Clear unfilled placeholder
            continue

        text = content_map[placeholder]
        info = slot_info.get(placeholder, {})
        slot_type = info.get("slot_type", "")
        formatting = info.get("formatting", {})

        if slot_type == "ref_link":
            # Reference links: bold part in regular runs, description in hyperlink runs
            sep = " | "
            idx = text.find(sep)
            if idx >= 0:
                bold_part = text[:idx]
                desc_part = text[idx + len(sep):]
            else:
                bold_part = text
                desc_part = ""

            # Fill regular runs with bold part
            if para.runs:
                para.runs[0].text = bold_part
                para.runs[0].bold = True
                # Clear other regular runs but leave separator
                for run in para.runs[1:]:
                    run.text = ""

            # Fill hyperlink runs with description
            hlinks = para._element.findall(".//w:hyperlink", nsmap_w)
            if hlinks and desc_part:
                h_runs = hlinks[0].findall(".//w:r", nsmap_w)
                if h_runs:
                    t_elem = h_runs[0].find("w:t", nsmap_w)
                    if t_elem is not None:
                        t_elem.text = desc_part
                # Add separator between bold part and hyperlink
                if para.runs and len(para.runs) > 1:
                    para.runs[1].text = " | "
                    para.runs[1].bold = None
        elif formatting.get("bold_label") and slot_type in BOLD_SEPARATORS:
            sep = BOLD_SEPARATORS[slot_type]
            _fill_bold_label(para, text, sep)
        else:
            _fill_simple(para, text)

        filled += 1

    # Add hyperlinks to contact line (email + LinkedIn)
    _add_contact_hyperlinks(doc, content_map)

    return doc


def _add_contact_hyperlinks(doc: Document, content_map: dict) -> None:
    """Add mailto: and https: hyperlinks to the contact line paragraph.

    Finds email and LinkedIn URL in the HEADER_CONTACT text and wraps them
    in clickable hyperlinks.
    """
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    from lxml import etree

    contact_text = content_map.get("HEADER_CONTACT", "")
    if not contact_text:
        return

    # Find the contact paragraph (para 1)
    para = doc.paragraphs[1]
    full_text = para.text

    # Extract email and LinkedIn URL from the text
    import re as _re
    email_match = _re.search(r'[\w.+-]+@[\w.-]+\.\w+', full_text)
    linkedin_match = _re.search(r'https?://[^\s]+linkedin[^\s]*', full_text)

    nsmap_w = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    nsmap_r = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"

    for match, url_prefix in [(email_match, "mailto:"), (linkedin_match, "")]:
        if not match:
            continue
        matched_text = match.group(0)
        url = f"{url_prefix}{matched_text}" if url_prefix else matched_text

        # Add relationship
        rel = doc.part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

        # Find and wrap the text in a hyperlink element
        for run in para.runs:
            if matched_text in run.text:
                # Split the run text around the match
                before = run.text[:run.text.index(matched_text)]
                after = run.text[run.text.index(matched_text) + len(matched_text):]

                # Set the run to just the before text
                run.text = before

                # Create hyperlink element with a run inside
                hyperlink = etree.SubElement(para._element, f"{{{nsmap_w}}}hyperlink")
                hyperlink.set(f"{{{nsmap_r}}}id", rel)

                # Create the hyperlink run (copy formatting from original run)
                h_run = etree.SubElement(hyperlink, f"{{{nsmap_w}}}r")
                # Copy run properties
                rPr = run._element.find(f"{{{nsmap_w}}}rPr")
                if rPr is not None:
                    h_rPr = deepcopy(rPr)
                    # Add hyperlink color style
                    h_run.insert(0, h_rPr)
                h_t = etree.SubElement(h_run, f"{{{nsmap_w}}}t")
                h_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
                h_t.text = matched_text

                # If there's text after the match, add it as a new run
                if after:
                    after_run = etree.SubElement(para._element, f"{{{nsmap_w}}}r")
                    if rPr is not None:
                        after_run.insert(0, deepcopy(rPr))
                    after_t = etree.SubElement(after_run, f"{{{nsmap_w}}}t")
                    after_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
                    after_t.text = after

                break  # Only wrap first occurrence


# -- Comparison --------------------------------------------------------------

def extract_text(doc: Document) -> str:
    """Extract all text from a Document for comparison."""
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def compare_texts(text_a: str, text_b: str) -> dict:
    """Compare two text strings line by line."""
    lines_a = text_a.strip().splitlines()
    lines_b = text_b.strip().splitlines()

    matches = 0
    diffs = []
    max_lines = max(len(lines_a), len(lines_b))

    for i in range(max_lines):
        la = lines_a[i].strip() if i < len(lines_a) else "(missing)"
        lb = lines_b[i].strip() if i < len(lines_b) else "(missing)"
        if la == lb:
            matches += 1
        else:
            diffs.append({
                "line": i + 1,
                "original": la[:120],
                "generated": lb[:120],
            })

    return {
        "total_lines": max_lines,
        "matching": matches,
        "differing": len(diffs),
        "match_pct": round(100 * matches / max_lines, 1) if max_lines else 100,
        "diffs": diffs[:20],
    }


# -- CLI ---------------------------------------------------------------------

def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="Generate a .docx resume from placeholder template + DB spec."
    )
    parser.add_argument("--version", default="v32", help="Resume version (default: v32)")
    parser.add_argument("--variant", default="base", help="Resume variant (default: base)")
    parser.add_argument("--template-name", default="V32 Placeholder",
                        help="Template name in DB (default: V32 Placeholder)")
    parser.add_argument("--template-file", help="Local template .docx file (overrides DB)")
    parser.add_argument("--map-file", help="Local template map JSON (overrides DB)")
    parser.add_argument("--spec", help="Optional JSON with content overrides for tailoring")
    parser.add_argument("--recipe-id", type=int, default=0,
                        help="Recipe ID from resume_recipes (overrides --version/--variant)")
    parser.add_argument("--validate", action="store_true",
                        help="Validate recipe references without generating (requires --recipe-id)")
    parser.add_argument("--dry-run", action="store_true",
                        help="Resolve recipe and print content_map without generating .docx")
    parser.add_argument("--output", required=True, help="Output .docx path")
    parser.add_argument("--compare", help="Path to original .docx for text comparison")
    return parser


def main(argv: Optional[list[str]] = None) -> int:
    parser = build_parser()
    args = parser.parse_args(argv)

    output_path = Path(args.output)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    try:
        conn = get_db_connection()
    except Exception as e:
        print(f"Error connecting to DB: {e}", file=sys.stderr)
        return 1

    try:
        # === RECIPE PATH ===
        if args.recipe_id > 0:
            recipe_row = fetch_recipe(conn, args.recipe_id)
            print(f"Recipe loaded: id={recipe_row['id']} name={recipe_row['name']}")

            recipe_json = recipe_row["recipe"]
            if isinstance(recipe_json, str):
                recipe_json = json.loads(recipe_json)

            # --validate: check references only, no generation
            if args.validate:
                result = validate_recipe(conn, recipe_json)
                print(f"\nValidation: {'PASS' if result['valid'] else 'FAIL'}")
                s = result["stats"]
                print(f"  Total slots: {s['total_slots']} | DB refs: {s['db_refs']} | Literals: {s['literals']}")
                print(f"  Valid refs: {s['valid_refs']} | Missing: {s['missing_refs']}")
                if result["errors"]:
                    print(f"\n  Errors:")
                    for e in result["errors"]:
                        print(f"    {e['slot']}: {e['error']}")
                conn.close()
                return 0 if result["valid"] else 1

            # Load template from recipe's template_id
            with conn.cursor() as cur:
                cur.execute(
                    "SELECT name FROM resume_templates WHERE id = %s",
                    (recipe_row["template_id"],),
                )
                tpl_name = cur.fetchone()[0]
            template_blob = fetch_template_blob(conn, tpl_name)
            template_map = fetch_template_map(conn, tpl_name)
            print(f"Template loaded: {tpl_name}")

            # Resolve recipe references to content_map
            content_map = resolve_recipe(conn, recipe_json)

            # Inject headline from recipe row
            if recipe_row.get("headline"):
                content_map["HEADLINE"] = recipe_row["headline"]

            print(f"Content map: {len(content_map)} slots from recipe")

            # --dry-run: print resolved content without generating
            if args.dry_run:
                print(f"\n=== DRY RUN (resolved content) ===")
                for slot, text in sorted(content_map.items()):
                    preview = text[:80].replace("\n", " ") if text else "(empty)"
                    print(f"  {slot}: {preview}")
                conn.close()
                return 0

            # Generate
            doc = generate_resume(template_blob, content_map, template_map)
            doc.save(str(output_path))
            print(f"Generated: {output_path}")

            # Compare if requested
            if args.compare:
                compare_path = Path(args.compare)
                if compare_path.exists():
                    original_doc = Document(str(compare_path))
                    generated_doc = Document(str(output_path))
                    result = compare_texts(extract_text(original_doc), extract_text(generated_doc))
                    print(f"\nText comparison vs {args.compare}:")
                    print(f"  Lines: {result['total_lines']}")
                    print(f"  Matching: {result['matching']} ({result['match_pct']}%)")
                    print(f"  Differing: {result['differing']}")
                    if result["diffs"]:
                        print(f"\n  Differences:")
                        for d in result["diffs"]:
                            print(f"    Line {d['line']}:")
                            print(f"      Original:  {d['original']}")
                            print(f"      Generated: {d['generated']}")

            conn.close()
            return 0

        # === LEGACY SPEC PATH ===
        # Load template
        if args.template_file:
            with open(args.template_file, "rb") as f:
                template_blob = f.read()
            print(f"Template loaded from file: {args.template_file}")
        else:
            template_blob = fetch_template_blob(conn, args.template_name)
            print(f"Template loaded from DB: {args.template_name}")

        # Load template map
        if args.map_file:
            with open(args.map_file, "r", encoding="utf-8") as f:
                template_map = json.load(f)
            print(f"Template map loaded from file: {args.map_file}")
        else:
            template_map = fetch_template_map(conn, args.template_name)
            if not template_map:
                # Try loading from local file
                local_map = Path("Output/template_v32_map.json")
                if local_map.exists():
                    with open(local_map, "r", encoding="utf-8") as f:
                        template_map = json.load(f)
                    print("Template map loaded from local file")
                else:
                    print("Warning: no template map found, formatting may be basic", file=sys.stderr)
                    template_map = {}

        # Load spec (optional — if no spec, use template_map original_text only)
        try:
            spec = fetch_resume_spec(conn, args.version, args.variant)
            print(f"Spec loaded: {args.version}/{args.variant}")
        except ValueError:
            spec = {}
            print(f"No spec for {args.version}/{args.variant} — using template_map original_text only")

        # Load supporting data
        header = fetch_header(conn)
        education = fetch_education(conn)
        certifications = fetch_certifications(conn)
        employers = spec.get("experience_employers", [])
        career = fetch_career_history(conn, employers) if employers else {}
        print(f"Data loaded: header, {len(education)} edu, {len(certifications)} certs, {len(career)} employers")

        # Build content map
        content_map = build_content_map(spec, header, education, certifications, career, template_map)

        # Apply overrides if provided
        if args.spec:
            with open(args.spec, "r", encoding="utf-8") as f:
                overrides = json.load(f)
            content_map.update(overrides)
            print(f"Overrides applied from: {args.spec}")

        print(f"Content map: {len(content_map)} slots")

        # Generate
        doc = generate_resume(template_blob, content_map, template_map)
        doc.save(str(output_path))
        print(f"Generated: {output_path}")

        # Compare if requested
        if args.compare:
            compare_path = Path(args.compare)
            if not compare_path.exists():
                print(f"Warning: comparison file not found: {args.compare}", file=sys.stderr)
            else:
                original_doc = Document(str(compare_path))
                generated_doc = Document(str(output_path))

                original_text = extract_text(original_doc)
                generated_text = extract_text(generated_doc)

                result = compare_texts(original_text, generated_text)
                print(f"\nText comparison vs {args.compare}:")
                print(f"  Lines: {result['total_lines']}")
                print(f"  Matching: {result['matching']} ({result['match_pct']}%)")
                print(f"  Differing: {result['differing']}")

                if result["diffs"]:
                    print(f"\n  Differences:")
                    for d in result["diffs"]:
                        print(f"    Line {d['line']}:")
                        print(f"      Original:  {d['original']}")
                        print(f"      Generated: {d['generated']}")

    except Exception as e:
        print(f"Error: {e}", file=sys.stderr)
        import traceback
        traceback.print_exc()
        return 1
    finally:
        conn.close()

    return 0


if __name__ == "__main__":
    sys.exit(main())
