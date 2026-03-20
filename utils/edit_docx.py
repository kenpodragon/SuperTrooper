"""Find and replace text in .docx files while preserving all formatting.

Handles text that is split across multiple runs by reconstructing the full
paragraph text, locating matches, and distributing replacement text back
across the original run boundaries.

Usage:
    python edit_docx.py <file> --find "old text" --replace "new text" [--output outfile.docx] [--all]

Dependencies:
    pip install python-docx
"""

import argparse
import sys
from pathlib import Path
from typing import Optional

from docx import Document
from docx.text.paragraph import Paragraph


def _replace_in_paragraph(paragraph: Paragraph, find_text: str, replace_text: str, replace_all: bool = False) -> int:
    """Replace text in a paragraph, handling text split across runs.

    This reconstructs the full paragraph text from all runs, finds the target
    string, then redistributes the modified text back into the runs so that
    formatting (bold, italic, font, color, etc.) is preserved as much as
    possible.

    Args:
        paragraph: The paragraph to modify.
        find_text: Text to find.
        replace_text: Text to replace with.
        replace_all: If True, replace all occurrences; otherwise just the first.

    Returns:
        Number of replacements made in this paragraph.
    """
    runs = paragraph.runs
    if not runs:
        return 0

    # Build the full text and a mapping of character index -> (run_index, char_index_in_run)
    full_text = ""
    char_map: list[tuple[int, int]] = []
    for run_idx, run in enumerate(runs):
        for char_idx, char in enumerate(run.text):
            char_map.append((run_idx, char_idx))
            full_text += char

    if find_text not in full_text:
        return 0

    # Perform replacement on the full text
    if replace_all:
        count = full_text.count(find_text)
        new_full_text = full_text.replace(find_text, replace_text)
    else:
        count = 1
        new_full_text = full_text.replace(find_text, replace_text, 1)

    # Strategy: clear all runs, put the new text in the first run
    # Then restore formatting by splitting text proportionally across runs.
    # Simpler and more reliable: put all text in first run, clear the rest.
    if len(runs) == 1:
        runs[0].text = new_full_text
    else:
        # Distribute text: give the new full text to the first run,
        # empty the rest. This preserves the first run's formatting.
        # For better formatting preservation, we try to keep run boundaries
        # close to where they originally were.
        _redistribute_text(runs, full_text, new_full_text, find_text, replace_text)

    return count


def _redistribute_text(runs, old_full: str, new_full: str, find_text: str, replace_text: str) -> None:
    """Redistribute modified text across existing runs to preserve formatting.

    Strategy: Walk through runs and the new text simultaneously. Each run gets
    roughly the same number of characters as it originally had, adjusted for
    the length difference at replacement sites.

    Args:
        runs: List of Run objects.
        old_full: Original concatenated text.
        new_full: New concatenated text after replacement.
        find_text: The text that was searched for.
        replace_text: The text that replaced it.
    """
    # Calculate original run lengths
    old_lengths = [len(r.text) for r in runs]
    len_diff = len(replace_text) - len(find_text)

    # Find where replacements occur in the original text
    # and adjust run lengths accordingly
    new_lengths = list(old_lengths)
    pos = 0
    remaining_diff = len(new_full) - len(old_full)

    # Simple approach: find which run(s) contain the replacement boundary
    # and adjust their lengths. Then assign text slices.
    offset = 0
    old_pos = 0
    adjustments_remaining = remaining_diff

    for i, run in enumerate(runs):
        run_start = old_pos
        run_end = old_pos + old_lengths[i]
        old_pos = run_end

        # Check if a replacement boundary falls in this run
        search_pos = 0
        while search_pos <= len(old_full) - len(find_text):
            idx = old_full.find(find_text, search_pos)
            if idx == -1:
                break
            # If the find_text starts in or overlaps this run
            if idx < run_end and idx + len(find_text) > run_start:
                # This run is affected by the replacement
                overlap_start = max(idx, run_start)
                overlap_end = min(idx + len(find_text), run_end)
                chars_in_this_run = overlap_end - overlap_start

                if idx >= run_start:
                    # Replacement starts in this run — give it the length diff
                    adjustment = min(adjustments_remaining, len_diff)
                    new_lengths[i] += adjustment
                    adjustments_remaining -= adjustment
            search_pos = idx + len(find_text)

    # Ensure total new lengths match new text length
    total_new = sum(new_lengths)
    if total_new != len(new_full):
        # Adjust the last non-empty run
        diff = len(new_full) - total_new
        new_lengths[-1] += diff

    # Ensure no negative lengths
    for i in range(len(new_lengths)):
        if new_lengths[i] < 0:
            surplus = -new_lengths[i]
            new_lengths[i] = 0
            # Push surplus to next run
            if i + 1 < len(new_lengths):
                new_lengths[i + 1] += surplus

    # Assign text slices to runs
    pos = 0
    for i, run in enumerate(runs):
        end = pos + max(0, new_lengths[i])
        run.text = new_full[pos:end]
        pos = end

    # If there's remaining text, append to last run
    if pos < len(new_full):
        runs[-1].text += new_full[pos:]


def find_replace(
    path: str,
    find_text: str,
    replace_text: str,
    output_path: Optional[str] = None,
    replace_all: bool = False,
) -> int:
    """Find and replace text in a .docx file, preserving formatting.

    Args:
        path: Path to the input .docx file.
        find_text: Text to find.
        replace_text: Text to replace with.
        output_path: Path to save the result. If None, overwrites the input file.
        replace_all: If True, replace all occurrences; otherwise just the first.

    Returns:
        Total number of replacements made.
    """
    doc = Document(path)
    total_replacements = 0

    # Process paragraphs in the document body
    for para in doc.paragraphs:
        total_replacements += _replace_in_paragraph(para, find_text, replace_text, replace_all)

    # Process paragraphs in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    total_replacements += _replace_in_paragraph(para, find_text, replace_text, replace_all)

    # Process headers and footers
    for section in doc.sections:
        for header_footer in [section.header, section.footer]:
            if header_footer is not None:
                for para in header_footer.paragraphs:
                    total_replacements += _replace_in_paragraph(para, find_text, replace_text, replace_all)

    save_path = output_path or path
    doc.save(save_path)
    return total_replacements


def build_parser() -> argparse.ArgumentParser:
    """Build the argument parser."""
    parser = argparse.ArgumentParser(
        description="Find and replace text in .docx files while preserving formatting."
    )
    parser.add_argument("file", help="Path to the .docx file")
    parser.add_argument("--find", required=True, help="Text to find")
    parser.add_argument("--replace", required=True, help="Replacement text")
    parser.add_argument(
        "--output", metavar="FILE", help="Output file path (default: overwrite input)"
    )
    parser.add_argument(
        "--all", action="store_true", dest="replace_all",
        help="Replace all occurrences (default: first only)"
    )
    return parser


def main(argv: Optional[list[str]] = None) -> int:
    """CLI entry point.

    Args:
        argv: Command-line arguments (defaults to sys.argv[1:]).

    Returns:
        Exit code (0 for success, 1 for error).
    """
    parser = build_parser()
    args = parser.parse_args(argv)

    path = Path(args.file)
    if not path.exists():
        print(f"Error: file not found: {path}", file=sys.stderr)
        return 1

    try:
        count = find_replace(
            str(path),
            args.find,
            args.replace,
            output_path=args.output,
            replace_all=args.replace_all,
        )
        save_target = args.output or str(path)
        if count == 0:
            print(f"No occurrences of '{args.find}' found.")
        else:
            print(f"Replaced {count} occurrence(s). Saved to: {save_target}")
    except Exception as e:
        print(f"Error: {e}", file=sys.stderr)
        return 1

    return 0


if __name__ == "__main__":
    sys.exit(main())
