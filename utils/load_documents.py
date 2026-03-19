"""
Documents ETL loader for SuperTroopers.

Scans document directories, extracts text content and metadata, then loads
into the documents and resume_versions tables. Idempotent... uses UPSERT on
file path so re-runs update existing rows. Skips unchanged files by SHA-256 hash.

Usage:
    python load_documents.py [--dry-run]

Source directories:
    Originals/              Original resumes, cover letters, reference docs
    Templates/              Base templates, tracker
    Imports/Organized/      Categorized archive files

Dependencies:
    psycopg2-binary, python-docx, pdfplumber, openpyxl
"""

import argparse
import hashlib
import json
import logging
import os
import re
import sys
import warnings
from datetime import date
from pathlib import Path

# Suppress noisy pdfplumber / pdfminer warnings
warnings.filterwarnings("ignore", message=".*FontBBox.*")
logging.getLogger("pdfminer").setLevel(logging.ERROR)

import psycopg2
import psycopg2.extras

# ---------------------------------------------------------------------------
# Project root (three levels up from code/utils/load_documents.py)
# ---------------------------------------------------------------------------
PROJECT_ROOT = Path(__file__).resolve().parent.parent.parent

SCAN_DIRS = [
    PROJECT_ROOT / "Originals",
    PROJECT_ROOT / "Templates",
    PROJECT_ROOT / "Imports" / "Organized",
]

SUPPORTED_EXTENSIONS = {".docx", ".pdf", ".xlsx", ".txt", ".md", ".doc"}

# ---------------------------------------------------------------------------
# Document type classification rules
# ---------------------------------------------------------------------------
# Order matters... first match wins.
TYPE_RULES = [
    (r"(?i)cover.?letter|cl[_\s]", "cover_letter"),
    (r"(?i)interview.?prep|interview.?q", "interview_prep"),
    (r"(?i)coach", "coaching"),
    (r"(?i)reference.?letter|letter.?of.?rec", "reference_letter"),
    (r"(?i)questionnaire|assessment", "questionnaire"),
    (r"(?i)transcript", "transcript"),
    (r"(?i)track", "tracker"),
    (r"(?i)template", "template"),
    (r"(?i)research|company.?research", "research"),
    (r"(?i)network", "networking"),
    (r"(?i)resume|cv", "resume"),
]

# Path-based classification for Imports/Organized subdirectories
PATH_TYPE_MAP = {
    "CoverLetters": "cover_letter",
    "Resumes": "resume",
    "References": "reference_letter",
    "InterviewPrep": "interview_prep",
    "CompanyResearch": "research",
    "Applied": "resume",
    "CareerCoaching": "coaching",
    "Networking": "networking",
}


# ---------------------------------------------------------------------------
# DB connection
# ---------------------------------------------------------------------------
from db_config import get_db_config


def get_connection():
    """Build a psycopg2 connection from env vars or .env file."""
    return psycopg2.connect(**get_db_config())


# ---------------------------------------------------------------------------
# Ensure unique constraint on documents.path for UPSERT
# ---------------------------------------------------------------------------
def ensure_unique_constraint(conn):
    """Make sure a unique constraint on documents.path exists for UPSERT."""
    with conn.cursor() as cur:
        cur.execute("""
            SELECT 1
            FROM pg_constraint
            WHERE conrelid = 'documents'::regclass
              AND contype = 'u'
              AND array_length(conkey, 1) = 1
              AND conkey[1] = (
                  SELECT attnum FROM pg_attribute
                  WHERE attrelid = 'documents'::regclass AND attname = 'path'
              )
        """)
        if cur.fetchone() is None:
            print("  Adding unique constraint on documents.path...")
            cur.execute(
                "ALTER TABLE documents ADD CONSTRAINT uq_documents_path UNIQUE (path)"
            )
            conn.commit()
            print("  Constraint added.")


# ---------------------------------------------------------------------------
# Text extraction
# ---------------------------------------------------------------------------
def extract_text_docx(filepath):
    """Extract text from a .docx file using python-docx."""
    from docx import Document

    doc = Document(str(filepath))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return "\n".join(paragraphs)


def extract_text_pdf(filepath):
    """Extract text from a .pdf file using pdfplumber. Returns (text, page_count)."""
    import pdfplumber

    pages = []
    page_count = 0
    with pdfplumber.open(str(filepath)) as pdf:
        page_count = len(pdf.pages)
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
    return "\n\n".join(pages), page_count


def extract_text_xlsx(filepath):
    """Extract text summary from an .xlsx file using openpyxl."""
    import openpyxl

    wb = openpyxl.load_workbook(str(filepath), read_only=True, data_only=True)
    parts = []
    for ws in wb.worksheets:
        parts.append(f"--- Sheet: {ws.title} ---")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            line = " | ".join(cells).strip()
            if line.replace("|", "").strip():
                parts.append(line)
    wb.close()
    return "\n".join(parts)


def extract_text_plain(filepath):
    """Read plain text files (.txt, .md)."""
    with open(filepath, "r", encoding="utf-8", errors="replace") as f:
        return f.read()


def extract_text_doc(filepath):
    """Attempt to extract text from a .doc file. Falls back to a notice."""
    try:
        from docx import Document

        doc = Document(str(filepath))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n".join(paragraphs)
    except Exception:
        return None


def extract_text(filepath):
    """
    Extract text from a file based on its extension.
    Returns (text, metadata_dict) or raises on failure.
    """
    ext = filepath.suffix.lower()
    metadata = {
        "file_size_bytes": filepath.stat().st_size,
        "extension": ext,
    }

    text = None

    if ext == ".docx":
        text = extract_text_docx(filepath)
    elif ext == ".pdf":
        text, page_count = extract_text_pdf(filepath)
        metadata["page_count"] = page_count
    elif ext == ".xlsx":
        text = extract_text_xlsx(filepath)
    elif ext in (".txt", ".md"):
        text = extract_text_plain(filepath)
    elif ext == ".doc":
        text = extract_text_doc(filepath)
        if text is None:
            metadata["extraction_note"] = "binary .doc - needs Word COM"
            text = ""

    if text is not None:
        # Strip null bytes that break PostgreSQL text columns
        text = text.replace("\x00", "")
        word_count = len(text.split()) if text.strip() else 0
        metadata["word_count"] = word_count

    return text, metadata


# ---------------------------------------------------------------------------
# Classification
# ---------------------------------------------------------------------------
def classify_document(filepath):
    """Classify a document by its path and filename. Returns a type string."""
    path_str = str(filepath)

    # Check path-based rules first (Imports/Organized subdirectories)
    for folder_name, doc_type in PATH_TYPE_MAP.items():
        if folder_name in path_str:
            return doc_type

    # Check filename-based rules
    filename = filepath.name
    for pattern, doc_type in TYPE_RULES:
        if re.search(pattern, filename):
            return doc_type

    # Check parent directory name
    parent = filepath.parent.name
    for pattern, doc_type in TYPE_RULES:
        if re.search(pattern, parent):
            return doc_type

    # Fallback: if in Templates/, call it template
    if "Templates" in path_str:
        return "template"

    return "other"


# ---------------------------------------------------------------------------
# Resume version extraction
# ---------------------------------------------------------------------------
def extract_resume_version_info(filepath):
    """
    Try to extract version number and variant from a resume filename.
    Examples:
        Stephen_Salaka_Resume_v32_AI_Architect.docx -> ("v32", "AI Architect")
        Resume_v31.pdf -> ("v31", "base")
        Stephen Salaka Resume.docx -> (None, "base")
    Returns (version, variant) or None if not a resume.
    """
    name = filepath.stem

    # Look for version pattern like v32, v31, V5, etc.
    version_match = re.search(r"(?i)\b(v\d+)\b", name)
    version = version_match.group(1).lower() if version_match else None

    # Look for variant after the version, or known keywords
    variant = "base"
    variant_patterns = {
        r"(?i)ai.?architect": "AI Architect",
        r"(?i)sw.?architect|software.?architect": "SW Architect",
        r"(?i)simplified|simple": "Simplified",
        r"(?i)\bpm\b|product.?manage": "PM",
        r"(?i)\bcto\b": "CTO",
        r"(?i)vp.?eng": "VP Eng",
        r"(?i)director": "Director",
        r"(?i)sr.?swe|senior.?swe|senior.?software": "Sr SWE",
    }
    for pattern, variant_name in variant_patterns.items():
        if re.search(pattern, name):
            variant = variant_name
            break

    return version, variant


# ---------------------------------------------------------------------------
# File scanning
# ---------------------------------------------------------------------------
def scan_directories(scan_dirs):
    """Walk directories and yield Path objects for supported files."""
    for base_dir in scan_dirs:
        if not base_dir.exists():
            print(f"  SKIP: Directory not found: {base_dir}")
            continue
        for root, _dirs, files in os.walk(base_dir):
            for fname in sorted(files):
                fpath = Path(root) / fname
                if fpath.suffix.lower() in SUPPORTED_EXTENSIONS:
                    yield fpath


def compute_hash(text):
    """Compute SHA-256 hash of text content."""
    return hashlib.sha256(text.encode("utf-8")).hexdigest()


# ---------------------------------------------------------------------------
# Database operations
# ---------------------------------------------------------------------------
UPSERT_DOC_SQL = """
INSERT INTO documents (path, filename, type, content_text, content_hash,
                       version, variant, extracted_date, metadata_json)
VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)
ON CONFLICT (path) DO UPDATE SET
    filename = EXCLUDED.filename,
    type = EXCLUDED.type,
    content_text = EXCLUDED.content_text,
    content_hash = EXCLUDED.content_hash,
    version = EXCLUDED.version,
    variant = EXCLUDED.variant,
    extracted_date = EXCLUDED.extracted_date,
    metadata_json = EXCLUDED.metadata_json
RETURNING id
"""

UPSERT_RESUME_SQL = """
INSERT INTO resume_versions (version, variant, docx_path, pdf_path,
                             target_role_type, document_id, is_current)
VALUES (%s, %s, %s, %s, %s, %s, %s)
ON CONFLICT (document_id) DO UPDATE SET
    version = EXCLUDED.version,
    variant = EXCLUDED.variant,
    docx_path = EXCLUDED.docx_path,
    pdf_path = EXCLUDED.pdf_path,
    target_role_type = EXCLUDED.target_role_type,
    is_current = EXCLUDED.is_current
RETURNING id
"""


def ensure_resume_versions_constraint(conn):
    """Ensure a unique constraint on resume_versions.document_id for UPSERT."""
    with conn.cursor() as cur:
        cur.execute("""
            SELECT 1
            FROM pg_constraint
            WHERE conrelid = 'resume_versions'::regclass
              AND contype = 'u'
              AND array_length(conkey, 1) = 1
              AND conkey[1] = (
                  SELECT attnum FROM pg_attribute
                  WHERE attrelid = 'resume_versions'::regclass
                    AND attname = 'document_id'
              )
        """)
        if cur.fetchone() is None:
            print("  Adding unique constraint on resume_versions.document_id...")
            cur.execute(
                "ALTER TABLE resume_versions "
                "ADD CONSTRAINT uq_resume_versions_document_id UNIQUE (document_id)"
            )
            conn.commit()
            print("  Constraint added.")


def load_file(filepath, conn, dry_run=False):
    """
    Process a single file: extract text, classify, hash, upsert.
    Returns a status string: 'inserted', 'updated', 'skipped', or 'error'.
    """
    rel_path = str(filepath)

    # Extract text + metadata
    try:
        text, metadata = extract_text(filepath)
    except Exception as e:
        print(f"  ERROR extracting {filepath.name}: {e}")
        return "error"

    if text is None:
        print(f"  ERROR: Unsupported format {filepath.suffix}")
        return "error"

    # Compute hash
    content_hash = compute_hash(text) if text else compute_hash("")

    # Classify
    doc_type = classify_document(filepath)

    # Extract version/variant for resumes
    version, variant = (None, None)
    if doc_type == "resume":
        version, variant = extract_resume_version_info(filepath)

    if dry_run:
        print(f"  [DRY RUN] {filepath.name} -> type={doc_type}, hash={content_hash[:12]}...")
        return "inserted"

    # Check if this exact hash already exists at this path
    with conn.cursor() as cur:
        cur.execute(
            "SELECT id, content_hash FROM documents WHERE path = %s",
            (rel_path,),
        )
        existing = cur.fetchone()

        if existing and existing[1] == content_hash:
            return "skipped"

        action = "updated" if existing else "inserted"

        # Upsert document
        cur.execute(
            UPSERT_DOC_SQL,
            (
                rel_path,
                filepath.name,
                doc_type,
                text,
                content_hash,
                version,
                variant,
                date.today(),
                json.dumps(metadata),
            ),
        )
        doc_id = cur.fetchone()[0]

        # For resume files, also populate resume_versions
        if doc_type == "resume" and (version or variant):
            ext = filepath.suffix.lower()
            docx_path = rel_path if ext == ".docx" else None
            pdf_path = rel_path if ext == ".pdf" else None

            # Map variant to target_role_type
            variant_to_role = {
                "AI Architect": "AI Architect",
                "SW Architect": "SW Architect",
                "PM": "PM",
                "CTO": "CTO",
                "VP Eng": "VP Eng",
                "Director": "Director",
                "Sr SWE": "Sr SWE",
                "Simplified": "Simplified",
                "base": None,
            }
            target_role = variant_to_role.get(variant)

            cur.execute(
                UPSERT_RESUME_SQL,
                (
                    version,
                    variant,
                    docx_path,
                    pdf_path,
                    target_role,
                    doc_id,
                    False,  # is_current... user sets manually
                ),
            )

    return action


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------
def build_parser(prog=None):
    parser = argparse.ArgumentParser(
        prog=prog,
        description="Scan documents and load metadata + text into SuperTroopers DB",
    )
    parser.add_argument(
        "--dry-run",
        action="store_true",
        help="Scan and classify without writing to DB",
    )
    parser.add_argument(
        "--dir",
        action="append",
        help="Additional directory to scan (can be repeated)",
    )
    return parser


def main(argv=None):
    parser = build_parser()
    args = parser.parse_args(argv)

    scan_dirs = list(SCAN_DIRS)
    if args.dir:
        for d in args.dir:
            scan_dirs.append(Path(d).resolve())

    print("Documents ETL Loader")
    print("=" * 60)
    print(f"Scanning {len(scan_dirs)} directories...")

    # Collect files first
    files = list(scan_directories(scan_dirs))
    print(f"  Found {len(files)} supported files")

    if not files:
        print("Nothing to load.")
        return

    if args.dry_run:
        print("\nDry run... no DB changes will be made.\n")
        counts = {"inserted": 0, "updated": 0, "skipped": 0, "error": 0}
        for f in files:
            try:
                text, metadata = extract_text(f)
                doc_type = classify_document(f)
                content_hash = compute_hash(text) if text else "empty"
                print(
                    f"  {f.name:50s} type={doc_type:16s} "
                    f"words={metadata.get('word_count', 0):>6d}  "
                    f"hash={content_hash[:12]}..."
                )
                counts["inserted"] += 1
            except Exception as e:
                print(f"  {f.name:50s} ERROR: {e}")
                counts["error"] += 1

        print(f"\n{counts['inserted']} files would be loaded, {counts['error']} errors.")
        return

    print("\nConnecting to SuperTroopers DB...")
    conn = get_connection()
    print("  Connected.")

    print("Checking constraints...")
    ensure_unique_constraint(conn)
    ensure_resume_versions_constraint(conn)

    print(f"\nProcessing {len(files)} files...\n")

    counts = {"inserted": 0, "updated": 0, "skipped": 0, "error": 0}

    for i, filepath in enumerate(files, 1):
        status = load_file(filepath, conn, dry_run=False)
        counts[status] += 1

        if status == "inserted":
            print(f"  [{i}/{len(files)}] NEW: {filepath.name}")
        elif status == "updated":
            print(f"  [{i}/{len(files)}] UPD: {filepath.name}")
        elif status == "error":
            pass  # already printed in load_file
        # skipped files stay silent

        # Commit in batches of 25
        if i % 25 == 0:
            conn.commit()

    # Final commit
    conn.commit()
    conn.close()

    print(f"\n{'=' * 60}")
    print(f"Done. Inserted: {counts['inserted']}, Updated: {counts['updated']}, "
          f"Skipped: {counts['skipped']}, Errors: {counts['error']}")


if __name__ == "__main__":
    main()
