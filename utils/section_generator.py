"""section_generator.py — .docx paragraph cloning engine for dynamic sections.

Takes a template .docx (bytes), a template_map describing section structure,
and a resolved recipe dict, and returns a filled .docx (bytes).

Section types:
  - Singular (repeating=False): fill prototype paragraph in-place
  - Repeating (repeating=True): clone prototype N times (one per item), fill each
  - Repeating with 0 items: remove prototype + section header paragraph
  - EXPERIENCE: compound section — clone per company header, job title, synopsis, bullets

Offset tracking: insertions/removals shift subsequent para_index values.
All section lookups use: section_def["para_index"] + running offset.
"""

import io
from copy import deepcopy
from typing import Optional

from docx import Document
from docx.text.paragraph import Paragraph


# ---------------------------------------------------------------------------
# Low-level paragraph operations
# ---------------------------------------------------------------------------

def clone_paragraph(doc: Document, prototype: Paragraph, after: Paragraph) -> Paragraph:
    """Deep-copy prototype paragraph XML and insert it after `after`.

    Returns the new Paragraph object. Preserves all run formatting via deepcopy.
    """
    new_elem = deepcopy(prototype._element)
    after._element.addnext(new_elem)
    return Paragraph(new_elem, after._parent)


def remove_paragraph(paragraph: Paragraph) -> None:
    """Remove paragraph element from its parent XML node."""
    p_elem = paragraph._element
    p_elem.getparent().remove(p_elem)


def fill_paragraph(paragraph: Paragraph, text: str, fmt: str, separator: Optional[str] = None) -> None:
    """Fill a paragraph with text using the specified format mode.

    Args:
        paragraph: Target paragraph to fill.
        text: Text content to write.
        fmt: 'simple' — preserve first run formatting, set text, clear others.
             'bold_label' — split at separator; first part bold, rest normal.
        separator: Required when fmt='bold_label'. Literal string to split on.
    """
    if fmt == "bold_label" and separator and separator in text:
        _fill_bold_label(paragraph, text, separator)
    else:
        _fill_simple(paragraph, text)


def _fill_simple(paragraph: Paragraph, text: str) -> None:
    """Preserve first run's formatting, set its text, clear other runs."""
    runs = paragraph.runs
    if runs:
        runs[0].text = text
        for run in runs[1:]:
            run.text = ""
    else:
        paragraph.text = text


def _fill_bold_label(paragraph: Paragraph, text: str, separator: str) -> None:
    """Split text at separator; first part bold, rest normal.

    Clears existing runs and adds new ones with correct formatting.
    """
    idx = text.find(separator)
    if idx < 0:
        _fill_simple(paragraph, text)
        return

    bold_part = text[:idx]
    rest_part = text[idx:]  # includes the separator itself

    runs = paragraph.runs
    if not runs:
        paragraph.text = text
        return

    # Clear all runs first
    for run in runs:
        run.text = ""

    # Re-use first run for bold part
    runs[0].text = bold_part
    runs[0].bold = True

    # Re-use second run for rest, or create a new run element
    if len(runs) > 1:
        runs[1].text = rest_part
        runs[1].bold = None
    else:
        # Clone the first run's properties but reset bold
        from copy import deepcopy
        from lxml import etree

        w_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        new_run = deepcopy(runs[0]._element)
        # Clear text content of cloned run
        for t_elem in new_run.findall(f"{{{w_ns}}}t"):
            new_run.remove(t_elem)
        # Clear bold in rPr
        rPr = new_run.find(f"{{{w_ns}}}rPr")
        if rPr is not None:
            b_elem = rPr.find(f"{{{w_ns}}}b")
            if b_elem is not None:
                rPr.remove(b_elem)
        # Set text
        t_elem = etree.SubElement(new_run, f"{{{w_ns}}}t")
        t_elem.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        t_elem.text = rest_part
        # Insert after first run
        runs[0]._element.addnext(new_run)


# ---------------------------------------------------------------------------
# Text assembly helpers
# ---------------------------------------------------------------------------

def _format_date_range(job: dict) -> str:
    """Format a date range string from a job dict.

    Returns "start – end" or "start – Present" if is_current.
    Dates may be date objects, strings, or None.
    """
    def _fmt(d) -> str:
        if d is None:
            return ""
        if hasattr(d, "strftime"):
            return d.strftime("%b %Y")
        s = str(d)
        # Handle "YYYY-MM-DD" strings
        if len(s) >= 7 and s[4] == "-":
            try:
                from datetime import datetime
                dt = datetime.strptime(s[:7], "%Y-%m")
                return dt.strftime("%b %Y")
            except ValueError:
                pass
        return s

    start = _fmt(job.get("start_date"))
    if job.get("is_current"):
        end = "Present"
    else:
        end = _fmt(job.get("end_date")) or "Present"

    if start:
        return f"{start} \u2013 {end}"
    return end


def _assemble_text(section_name: str, data: dict) -> str:
    """Assemble display text for singular sections.

    Args:
        section_name: e.g. 'HEADER', 'SUMMARY', 'HEADLINE'
        data: dict with section fields (from resolver)

    Returns assembled string.
    """
    if section_name == "HEADER":
        name = data.get("full_name", "")
        creds = data.get("credentials", "")
        if creds:
            return f"{name}, {creds}"
        return name

    if section_name == "HEADER_CONTACT":
        parts = []
        loc = data.get("location", "")
        loc_note = data.get("location_note", "")
        if loc:
            if loc_note:
                parts.append(f"{loc} ({loc_note})")
            else:
                parts.append(loc)
        if data.get("email"):
            parts.append(data["email"])
        if data.get("phone"):
            parts.append(data["phone"])
        if data.get("linkedin_url"):
            parts.append(data["linkedin_url"])
        return " \u2022 ".join(parts)

    if section_name in ("SUMMARY", "HEADLINE"):
        # summary_variants row — return text or headline
        if section_name == "HEADLINE":
            return data.get("headline", data.get("text", ""))
        return data.get("text", "")

    # Fallback: return string representation
    if isinstance(data, str):
        return data
    return str(data)


def _assemble_item_text(section_name: str, data: dict) -> str:
    """Assemble display text for a single item in a repeating section.

    Args:
        section_name: e.g. 'CERTIFICATIONS', 'EDUCATION', 'SKILLS', 'HIGHLIGHTS'
        data: dict with item fields (from resolver)

    Returns assembled string.
    """
    if section_name == "CERTIFICATIONS":
        name = data.get("name", "")
        issuer = data.get("issuer", "")
        if issuer:
            return f"{name} | {issuer}"
        return name

    if section_name == "EDUCATION":
        parts = []
        for field in ("degree", "field", "institution", "location"):
            val = data.get(field, "")
            if val:
                parts.append(val)
        return " | ".join(parts)

    if section_name == "SKILLS":
        return data.get("name", "")

    if section_name in ("HIGHLIGHTS", "HIGHLIGHT"):
        return data.get("text", data.get("content", ""))

    if section_name == "ADDITIONAL_EXP":
        employer = data.get("employer", "")
        title = data.get("title", "")
        if employer and title:
            return f"{employer} \u2014 {title}"
        return employer or title

    # Fallback
    if isinstance(data, str):
        return data
    return str(data)


# ---------------------------------------------------------------------------
# EXPERIENCE compound section
# ---------------------------------------------------------------------------

def _generate_experience(
    doc: Document,
    prototype: Paragraph,
    companies: list,
    offset: int,
) -> int:
    """Generate EXPERIENCE section by cloning prototype for each sub-element.

    For each company:
      - company header line (bold_label, separator=", ")
      - for each job:
          - title line with date range (bold_label, separator=", ")
          - synopsis/intro (simple) if present
          - bullets (bold_label, separator=": ")

    Removes the original prototype when done.

    Args:
        doc: Document object
        prototype: The {{EXPERIENCE}} prototype paragraph
        companies: Resolved list from section_resolver._resolve_experience
        offset: Current paragraph index offset (not used directly here,
                but returned delta is applied by caller)

    Returns:
        int: net change in paragraph count (cloned - 1 for removed prototype)
    """
    insert_after = prototype
    cloned_count = 0

    for company_entry in companies:
        company = company_entry.get("company") or {}
        jobs = company_entry.get("jobs", [])

        # Company header: "Employer, Location {Industry}"
        company_parts = []
        employer = company.get("employer", "")
        location = company.get("location", "")
        industry = company.get("industry", "")
        if employer:
            company_parts.append(employer)
        if location:
            company_parts.append(location)
        if industry:
            company_parts.append(f"{{{industry}}}")
        company_header_text = ", ".join(company_parts)

        para = clone_paragraph(doc, prototype, insert_after)
        fill_paragraph(para, company_header_text, "bold_label", ", ")
        insert_after = para
        cloned_count += 1

        for job_entry in jobs:
            job = job_entry.get("job") or {}
            bullets = job_entry.get("bullets", [])

            # Job title line: "Title, DateRange"
            title = job.get("title", "")
            date_range = _format_date_range(job)
            if title and date_range:
                title_text = f"{title}, {date_range}"
            elif title:
                title_text = title
            else:
                title_text = date_range

            para = clone_paragraph(doc, prototype, insert_after)
            fill_paragraph(para, title_text, "bold_label", ", ")
            insert_after = para
            cloned_count += 1

            # Synopsis / intro_text
            intro = job.get("intro_text")
            if intro:
                # Handle intro stored as dict (legacy import format)
                if isinstance(intro, dict):
                    intro = intro.get("text", "")
                if intro:
                    para = clone_paragraph(doc, prototype, insert_after)
                    fill_paragraph(para, intro, "simple")
                    insert_after = para
                    cloned_count += 1

            # Bullets
            for bullet in bullets:
                if isinstance(bullet, dict):
                    bullet_text = bullet.get("text", "")
                else:
                    bullet_text = str(bullet)
                if bullet_text:
                    para = clone_paragraph(doc, prototype, insert_after)
                    fill_paragraph(para, bullet_text, "bold_label", ": ")
                    insert_after = para
                    cloned_count += 1

    # Remove the original prototype
    remove_paragraph(prototype)
    # Net change: added cloned_count, removed 1
    return cloned_count - 1


# ---------------------------------------------------------------------------
# Main entry point
# ---------------------------------------------------------------------------

def generate_from_sections(
    template_blob: bytes,
    template_map: dict,
    resolved: dict,
) -> bytes:
    """Generate a filled .docx from a template, template_map, and resolved content.

    Args:
        template_blob: Raw bytes of the placeholder template .docx
        template_map: Dict mapping section names to their definition:
            {
              "CERTIFICATIONS": {
                "repeating": True,
                "para_index": 4,
                "format": "bold_label",
                "separator": " | ",
                "section_header_para": 3,   # optional, index of section header above prototype
              },
              "HEADER": {
                "repeating": False,
                "para_index": 0,
                "format": "simple",
              },
              ...
            }
        resolved: Dict mapping section names to resolved content:
            - singular: dict (single row) or str
            - repeating: list[dict] or list[str]
            - EXPERIENCE: list[{company, jobs: [{job, bullets}]}]

    Returns:
        bytes: The filled .docx as raw bytes.
    """
    doc = Document(io.BytesIO(template_blob))

    # Sort sections by para_index ascending (process top-to-bottom)
    sections_sorted = sorted(
        template_map.items(),
        key=lambda kv: kv[1].get("para_index", 0),
    )

    offset = 0  # running paragraph index adjustment due to insertions/removals

    for section_name, section_def in sections_sorted:
        raw_index = section_def.get("para_index")
        if raw_index is None:
            continue

        real_index = raw_index + offset
        paras = doc.paragraphs

        if real_index < 0 or real_index >= len(paras):
            # Index out of range after offset shifts — skip gracefully
            continue

        prototype = paras[real_index]
        fmt = section_def.get("format", "simple")
        separator = section_def.get("separator")
        repeating = section_def.get("repeating", False)
        content = resolved.get(section_name)

        # ----------------------------------------------------------------
        # SINGULAR section
        # ----------------------------------------------------------------
        if not repeating:
            if content is None:
                continue
            if isinstance(content, dict):
                text = _assemble_text(section_name, content)
            else:
                text = str(content)
            fill_paragraph(prototype, text, fmt, separator)
            # No offset change — paragraph count unchanged

        # ----------------------------------------------------------------
        # EXPERIENCE (compound repeating)
        # ----------------------------------------------------------------
        elif section_name == "EXPERIENCE":
            if not content:
                # 0 companies — remove prototype + section header
                section_header_idx = section_def.get("section_header_para")
                if section_header_idx is not None:
                    header_real = section_header_idx + offset
                    if 0 <= header_real < len(doc.paragraphs):
                        remove_paragraph(doc.paragraphs[header_real])
                        offset -= 1
                        # prototype index shifted up by 1 after header removal
                        real_index = raw_index + offset
                        if 0 <= real_index < len(doc.paragraphs):
                            prototype = doc.paragraphs[real_index]
                        else:
                            continue
                remove_paragraph(prototype)
                offset -= 1
            else:
                delta = _generate_experience(doc, prototype, content, offset)
                offset += delta

        # ----------------------------------------------------------------
        # REPEATING section
        # ----------------------------------------------------------------
        else:
            items = content if isinstance(content, list) else ([] if content is None else [content])

            if len(items) == 0:
                # Remove prototype + optional section header
                section_header_idx = section_def.get("section_header_para")
                if section_header_idx is not None:
                    header_real = section_header_idx + offset
                    if 0 <= header_real < len(doc.paragraphs):
                        remove_paragraph(doc.paragraphs[header_real])
                        offset -= 1
                        # Recalculate prototype position after header removal
                        real_index = raw_index + offset
                        if 0 <= real_index < len(doc.paragraphs):
                            prototype = doc.paragraphs[real_index]
                        else:
                            continue
                remove_paragraph(prototype)
                offset -= 1

            elif len(items) == 1:
                # Fill prototype in-place (no cloning needed)
                item = items[0]
                if isinstance(item, dict):
                    text = _assemble_item_text(section_name, item)
                else:
                    text = str(item)
                fill_paragraph(prototype, text, fmt, separator)
                # No offset change

            else:
                # Clone N times, fill each, remove original prototype
                insert_after = prototype
                for item in items:
                    if isinstance(item, dict):
                        text = _assemble_item_text(section_name, item)
                    else:
                        text = str(item)
                    para = clone_paragraph(doc, prototype, insert_after)
                    fill_paragraph(para, text, fmt, separator)
                    insert_after = para

                # Remove original prototype (it was the template marker)
                remove_paragraph(prototype)
                # Net: added N, removed 1
                offset += len(items) - 1

    # Save to bytes and return
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
