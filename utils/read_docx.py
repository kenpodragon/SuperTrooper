"""Extract text, sections, or search content from .docx files.

Usage:
    python read_docx.py <file> [--search "query"] [--section "heading name"]

Dependencies:
    pip install python-docx
"""

import argparse
import sys
from pathlib import Path
from typing import Optional

from docx import Document


def read_full_text(path: str) -> str:
    """Read all text from a .docx file, preserving paragraph breaks.

    Args:
        path: Path to the .docx file.

    Returns:
        Full document text with newlines between paragraphs.
    """
    doc = Document(path)
    paragraphs = [p.text for p in doc.paragraphs]
    return "\n".join(paragraphs)


def read_sections(path: str) -> dict[str, str]:
    """Split document text by headings into named sections.

    Args:
        path: Path to the .docx file.

    Returns:
        Dictionary mapping heading text to the content under that heading.
        Content before any heading is stored under the key "_preamble".
    """
    doc = Document(path)
    sections: dict[str, str] = {}
    current_heading = "_preamble"
    current_lines: list[str] = []

    for para in doc.paragraphs:
        style_name = (para.style.name or "").lower()
        if style_name.startswith("heading"):
            # Save the previous section
            sections[current_heading] = "\n".join(current_lines).strip()
            current_heading = para.text.strip()
            current_lines = []
        else:
            current_lines.append(para.text)

    # Save the last section
    sections[current_heading] = "\n".join(current_lines).strip()

    # Remove empty preamble
    if "_preamble" in sections and not sections["_preamble"]:
        del sections["_preamble"]

    return sections


def search_text(path: str, query: str) -> list[str]:
    """Search for paragraphs containing the query string (case-insensitive).

    Args:
        path: Path to the .docx file.
        query: Text to search for.

    Returns:
        List of paragraphs that contain the query.
    """
    doc = Document(path)
    query_lower = query.lower()
    results = []
    for para in doc.paragraphs:
        if query_lower in para.text.lower():
            results.append(para.text)
    return results


def build_parser() -> argparse.ArgumentParser:
    """Build the argument parser."""
    parser = argparse.ArgumentParser(
        description="Extract text, sections, or search content from .docx files."
    )
    parser.add_argument("file", help="Path to the .docx file")
    parser.add_argument(
        "--search", metavar="QUERY", help="Search for paragraphs containing this text"
    )
    parser.add_argument(
        "--section",
        metavar="HEADING",
        help="Print only the section under this heading name",
    )
    parser.add_argument(
        "--list-sections",
        action="store_true",
        help="List all section headings in the document",
    )
    return parser


def main(argv: Optional[list[str]] = None) -> int:
    """CLI entry point.

    Args:
        argv: Command-line arguments (defaults to sys.argv[1:]).

    Returns:
        Exit code (0 for success, 1 for error).
    """
    parser = build_parser()
    args = parser.parse_args(argv)

    path = Path(args.file)
    if not path.exists():
        print(f"Error: file not found: {path}", file=sys.stderr)
        return 1
    if not path.suffix.lower() == ".docx":
        print(f"Warning: file does not have .docx extension: {path}", file=sys.stderr)

    try:
        if args.list_sections:
            sections = read_sections(str(path))
            for heading in sections:
                print(heading)
        elif args.search:
            results = search_text(str(path), args.search)
            if not results:
                print(f"No paragraphs found matching: {args.search}")
            else:
                print(f"Found {len(results)} matching paragraph(s):\n")
                for i, text in enumerate(results, 1):
                    print(f"  [{i}] {text}\n")
        elif args.section:
            sections = read_sections(str(path))
            # Case-insensitive heading lookup
            target = args.section.lower()
            match = None
            for heading, content in sections.items():
                if heading.lower() == target:
                    match = (heading, content)
                    break
            if match:
                print(f"=== {match[0]} ===\n")
                print(match[1])
            else:
                print(f"Section not found: {args.section}", file=sys.stderr)
                print("Available sections:", file=sys.stderr)
                for heading in sections:
                    print(f"  - {heading}", file=sys.stderr)
                return 1
        else:
            print(read_full_text(str(path)))
    except Exception as e:
        print(f"Error reading document: {e}", file=sys.stderr)
        return 1

    return 0


if __name__ == "__main__":
    sys.exit(main())
