"""General-purpose .docx resume parser with section detection and formatting extraction.

Parses any resume layout, detecting sections like header, experience, education,
bullets, etc. based on formatting heuristics (font size, bold, alignment, patterns).

Also provides parse_resume_for_kb() which converts the flat paragraph list into
the nested career_history + bullets format expected by the onboard DB insert pipeline.
"""

import logging
import re
from datetime import date
from pathlib import Path
from docx import Document
from docx.shared import Pt, Emu

logger = logging.getLogger(__name__)


# --- Contact pattern regexes ---
EMAIL_RE = re.compile(r'[\w.+-]+@[\w-]+\.[\w.-]+')
PHONE_RE = re.compile(r'\(?\d{3}\)?[\s.\-]?\d{3}[\s.\-]?\d{4}')
LINKEDIN_RE = re.compile(r'linkedin\.com/in/', re.IGNORECASE)

# --- Date patterns for job headers ---
DATE_RE = re.compile(
    r'(?:'
    r'(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\w*[\s,.]*\d{4}'
    r'|\d{4}\s*[-–—]\s*(?:\d{4}|[Pp]resent|[Cc]urrent)'
    r'|(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\w*\s+\d{4}\s*[-–—]'
    r')',
    re.IGNORECASE
)

# --- Section header keywords (case-insensitive match) ---
SECTION_KEYWORDS = {
    'experience': 'experience',
    'education': 'education',
    'certification': 'certification',
    'skills': 'skills',
    'summary': 'summary',
    'highlights': 'highlights',
    'objective': 'summary',
    'profile': 'summary',
    'qualifications': 'highlights',
    'competencies': 'highlights',
    'expertise': 'skills',
    'technical': 'skills',
    'additional': 'additional',
    'references': 'references',
    'awards': 'additional',
    'honors': 'additional',
    'publications': 'additional',
    'projects': 'experience',
    'volunteer': 'additional',
    'interests': 'additional',
    'languages': 'skills',
    'professional development': 'education',
    'training': 'education',
    'keywords': 'keywords',
    'previous positions': 'additional',
    'previous experience': 'additional',
    'earlier career': 'additional',
    'other experience': 'additional',
    'organizations': 'additional',
    'patents': 'additional',
    'recommendations': 'additional',
}

# Bullet prefix characters
BULLET_CHARS = set('•●○▪▸►–—‣⁃◦◆◇■□▶')

# Action verbs that signal achievement bullets (past tense)
_ACHIEVEMENT_VERBS = {
    'achieved', 'administered', 'advanced', 'analyzed', 'applied', 'architected',
    'attracted', 'automated', 'boosted', 'built', 'catalyzed', 'centralized',
    'championed', 'coached', 'co-led', 'collaborated', 'completed', 'conceived',
    'conceptualized', 'consolidated', 'contributed', 'converted', 'coordinated',
    'created', 'cultivated', 'cut', 'decreased', 'defined', 'delivered',
    'deployed', 'designed', 'developed', 'devised', 'directed', 'doubled',
    'drove', 'earned', 'eliminated', 'enabled', 'engineered', 'enhanced',
    'ensured', 'established', 'evaluated', 'exceeded', 'executed', 'expanded',
    'expedited', 'facilitated', 'formulated', 'founded', 'generated', 'grew',
    'guided', 'headed', 'hired', 'identified', 'implemented', 'improved',
    'increased', 'influenced', 'initiated', 'innovated', 'instituted',
    'integrated', 'introduced', 'launched', 'led', 'leveraged', 'maintained',
    'managed', 'maximized', 'mentored', 'migrated', 'minimized', 'modernized',
    'negotiated', 'onboarded', 'operated', 'optimized', 'orchestrated',
    'organized', 'oversaw', 'partnered', 'performed', 'piloted', 'pioneered',
    'planned', 'presented', 'prevented', 'produced', 'promoted', 'proposed',
    'provided', 'published', 'raised', 'realized', 'rebuilt', 'recruited',
    'redesigned', 'reduced', 'reengineered', 'restructured', 'revamped',
    'reversed', 'revitalized', 'saved', 'scaled', 'secured', 'seamlessly',
    'simplified', 'slashed', 'solved', 'spearheaded', 'standardized',
    'streamlined', 'strengthened', 'structured', 'supervised', 'surpassed',
    'trained', 'transformed', 'transitioned', 'tripled', 'unified', 'upgraded',
    'utilized', 'won',
}

# Inline skill label patterns (e.g., "Key Skills: ...", "Core Competencies: ...")
_INLINE_SKILLS_RE = re.compile(
    r'^(?:Key\s+Skills|Core\s+Competencies|Technical\s+Skills|Areas?\s+of\s+Expertise)\s*:',
    re.IGNORECASE,
)


def _extract_formatting(paragraph) -> dict:
    """Extract formatting metadata from a paragraph."""
    fmt = {
        'font_size': None,
        'bold': None,
        'italic': None,
        'alignment': None,
        'font_name': None,
        'underline': None,
        'space_before': None,
        'space_after': None,
    }

    # Paragraph-level formatting
    pf = paragraph.paragraph_format
    if pf.alignment is not None:
        fmt['alignment'] = str(pf.alignment)
    if pf.space_before is not None:
        fmt['space_before'] = pf.space_before
    if pf.space_after is not None:
        fmt['space_after'] = pf.space_after

    # Run-level formatting (use first run as representative)
    if paragraph.runs:
        run = paragraph.runs[0]
        if run.font.size is not None:
            # Convert EMUs to points
            fmt['font_size'] = round(run.font.size / 12700, 1)
        fmt['bold'] = run.bold
        fmt['italic'] = run.italic
        fmt['font_name'] = run.font.name
        fmt['underline'] = run.underline

    # Try style-level font size if run didn't have one
    if fmt['font_size'] is None and paragraph.style and paragraph.style.font:
        if paragraph.style.font.size is not None:
            fmt['font_size'] = round(paragraph.style.font.size / 12700, 1)

    return fmt


def _has_contact_info(text: str) -> bool:
    """Check if text contains contact information patterns."""
    return bool(EMAIL_RE.search(text) or PHONE_RE.search(text) or LINKEDIN_RE.search(text))


def _has_date_pattern(text: str) -> bool:
    """Check if text contains date patterns typical of job entries."""
    return bool(DATE_RE.search(text))


def _is_bullet_text(text: str) -> bool:
    """Check if text starts with a bullet character or dash-space pattern."""
    if not text:
        return False
    if text[0] in BULLET_CHARS:
        return True
    # Dash followed by space (but not em-dash which is in BULLET_CHARS)
    if text.startswith('- '):
        return True
    return False


def _is_achievement_text(text: str) -> bool:
    """Check if text looks like an achievement bullet (starts with action verb or has strong metrics)."""
    words = text.split(None, 1)
    if not words:
        return False
    first_word = words[0].lower().rstrip('.,;:')

    if first_word in _ACHIEVEMENT_VERBS:
        return True

    # "Label: Action verb..." pattern (e.g., "Expansion Through Client Acquisition: Spearheaded...")
    if ':' in text:
        after_colon = text.split(':', 1)[1].strip()
        after_words = after_colon.split(None, 1)
        if after_words and after_words[0].lower().rstrip('.,;:') in _ACHIEVEMENT_VERBS:
            return True

    # Strong metrics signal: even without a recognized verb, numbers + context = achievement
    has_metric = bool(re.search(
        r'\d+%|\$[\d,.]+K?M?B?|\d+x\b|\d+\+?\s*(?:team|engineer|member|staff|people|user|client|customer)',
        text, re.IGNORECASE,
    ))
    if has_metric and len(text) > 60:
        return True

    return False


def _looks_like_company_name(text: str, prev_type: str | None) -> bool:
    """Check if short (<80 char) non-bold text is a standalone company name.

    Heuristics:
    - Appears after end of a previous job (bullet, job_intro) or section_header
    - Short, no colon (not a skill label), not starting with common description words
    - Title-case or all-caps typical of company names
    """
    # Only plausible between jobs
    if prev_type not in ('bullet', 'job_intro', 'section_header', None):
        return False
    # Skill labels like "Key Skills: ..."
    if ':' in text:
        return False
    # Very short text (under ~50 chars) that looks like a proper noun
    stripped = text.strip()
    if len(stripped) > 60:
        return False
    # Reject bracket-tagged metadata like "[Source: V32]", "[Listed under...]"
    if stripped.startswith('['):
        return False
    # Reject if it starts with a lowercase word (description, not a name)
    if stripped and stripped[0].islower():
        return False
    # Reject narrative sentences ("As the CTO...", "During my time...", "From 2012...")
    narrative_starts = {'as an', 'as the', 'at ', 'during', 'from ', 'i demo', 'i served',
                        'expert', 'well-vers', 'master', 'my toolkit', 'proficient',
                        'honored', 'salaka', 'in my role'}
    lower = stripped.lower()
    if any(lower.startswith(ns) for ns in narrative_starts):
        return False
    # Reject skill/tech lists (multiple tech words with no company structure)
    tech_words = {'python', 'java', 'aws', 'sql', 'docker', 'azure', 'react', 'node',
                  'c++', 'c#', 'kubernetes', 'terraform', 'atlassian', 'tensor', 'apache'}
    words_lower = [w.lower().rstrip('.,;') for w in stripped.split()]
    if len(words_lower) >= 3 and sum(1 for w in words_lower if w in tech_words) >= 2:
        return False
    # Reject pipe-delimited lines (role summaries, not company names)
    if '|' in stripped:
        return False
    # Reject if it reads like a description ("provider", "startup", "agency")
    desc_words = {'provider', 'serving', 'specializing', 'offering', 'delivering',
                  'connecting', 'helping', 'supporting', 'focused on', 'software development,',
                  'strategic partner', 'process st'}
    if any(dw in lower for dw in desc_words):
        return False
    return True


def _is_junk_employer(text: str) -> bool:
    """Reject text that should never become a career_history employer.

    Catches: narrative sentences, skill lists, template placeholders,
    section headers, metadata tags, single generic words.
    """
    stripped = text.strip()
    lower = stripped.lower()

    # Bracket-tagged metadata: [Source: V32], [Listed under...]
    if stripped.startswith('['):
        return True
    # Template placeholders: (A2) Company Name..., (C3) Location...
    if re.match(r'^\([A-Z]\d\)', stripped):
        return True
    # Narrative sentences starting with pronouns/prepositions
    narrative_re = re.compile(
        r'^(as an?|as the|at |during|from \d|i [a-z]|expert|well[- ]vers|'
        r'master|my |proficient|honored|salaka|in my |successfully |'
        r'kind regards|page \d)',
        re.IGNORECASE,
    )
    if narrative_re.match(stripped):
        return True
    # Skill/tech lists (3+ words with 2+ tech keywords)
    tech_words = {'python', 'java', 'aws', 'sql', 'docker', 'azure', 'react', 'node',
                  'c++', 'c#', 'kubernetes', 'terraform', 'atlassian', 'tensor', 'apache',
                  'tomcat', 'javascript', 'typescript', 'php', 'mysql', 'ruby', 'golang'}
    words_lower = [w.lower().rstrip('.,;') for w in stripped.split()]
    if len(words_lower) >= 3 and sum(1 for w in words_lower if w in tech_words) >= 2:
        return True
    # Single generic words that aren't companies
    generic_singles = {'unknown', 'publishing', 'ergebnisse', 'social impact', 'languages',
                       'training', 'organizations', 'patents', 'recruitment', 'manufacturing',
                       'c#', 'c++', 'java', 'python', 'previous positions'}
    if lower in generic_singles:
        return True
    # Title-as-employer: starts with a job title prefix (CTO/, Director,, VP , Senior )
    if re.match(r'^(CTO/|Director[,\s]|VP\s|Senior\s|Chief\s)', stripped) and ',' in stripped:
        return True
    # Too long to be a company name (>80 chars = likely a description)
    if len(stripped) > 80:
        return True
    # Pipe-delimited summary lines
    if '|' in stripped and stripped.count('|') >= 2:
        return True
    # German/non-English text patterns
    if any(w in lower for w in ['gründer', 'direktor', 'für', 'endergebnis', 'ergebnisse',
                                 'anwendungsentwick', 'dezember', 'januar', 'gegenwart',
                                 'jahre', 'monate']):
        return True
    # Job titles masquerading as employers (start with title words, no company after)
    title_as_emp = re.match(
        r'^(Chief|Director|VP|Vice President|Senior|Lead|Head)\s',
        stripped, re.IGNORECASE,
    )
    if title_as_emp and not _LOCATION_RE.search(stripped):
        # Has title words but no location → likely a title not a company
        # Exception: if it also contains a real company name indicator
        if not any(w in lower for w in ['inc', 'llc', 'corp', 'associates', 'solutions',
                                         'consulting', 'tsolutions', 'smtc', 'atex',
                                         'newscycle', 'granted', 'tutor', 'nova',
                                         'datavers', 'mealmatch', 'wall street']):
            return True
    # Questions/conversational text
    if lower.startswith('is there') or lower.startswith('what ') or lower.startswith('how '):
        return True
    # Personal names (not companies)
    if lower in ('stephen salaka', 'stephen a salaka', 'stephen a. salaka'):
        return True
    # Transcript/document references
    if 'transcript' in lower:
        return True
    # Generic terms, section headers, and non-company text
    generic_terms = {'cloud computing', 'lean six sigma master black belt trainer',
                     'city and state: melbourne, fl', 'eagle scout',
                     'rejection response', 'distributed development system and platform',
                     'thoughtchain distributed ai processing'}
    if lower in generic_terms or lower.rstrip(' ,') in generic_terms:
        return True
    # Skills lists masquerading as employers (comma-separated technical terms)
    if lower.count(',') >= 2 and any(w in lower for w in ['coding', 'migrations', 'quality',
                                                            'development,', 'management,']):
        return True
    # Contains literal placeholder text
    if 'location' == lower.split(',')[-1].strip().lower():
        return True
    # Embedded tab characters (malformed parse artifacts)
    if '\t' in stripped:
        return True
    # Parenthetical with no content: "... ()" or "... ( )" at end
    if re.search(r'\(\s*\)\s*$', stripped):
        return True
    # "Manager Foo, State" pattern — title + store, not a company
    if re.match(r'^(Manager|Customer Service)\s', stripped, re.IGNORECASE):
        return True
    return False


def _match_section_keyword(text: str) -> str | None:
    """Match text against known section header keywords. Returns section category or None."""
    lower = text.strip().lower()
    # Remove common decorators
    lower = lower.strip(':').strip()

    for keyword, category in SECTION_KEYWORDS.items():
        if keyword in lower:
            return category
    return None


def _classify_paragraph(
    text: str,
    formatting: dict,
    para_index: int,
    current_section: str,
    is_early: bool,
    prev_type: str | None,
) -> tuple[str, str]:
    """Classify a paragraph into a type and update the current section.

    Returns (paragraph_type, updated_current_section).
    """
    font_size = formatting.get('font_size')
    bold = formatting.get('bold')

    # --- Header detection (early paragraphs with large font or contact info) ---
    if is_early and font_size and font_size >= 14:
        return 'header', 'header'

    if is_early and _has_contact_info(text) and current_section in ('header', ''):
        return 'header', 'header'

    # --- Section header detection ---
    # Bold text that matches section keywords and is relatively short
    section_cat = _match_section_keyword(text)
    if section_cat and len(text) < 80:
        # Check if it looks like a standalone header (bold or larger font)
        if bold or (font_size and font_size >= 11):
            return 'section_header', section_cat

    # --- Headline (title/tagline after name, before content) ---
    if is_early and font_size and 12 < font_size < 20 and bold:
        return 'headline', 'header'

    # --- Within specific sections ---
    if current_section == 'experience' or current_section == 'additional':
        # 1. Bullet items (explicit bullet chars)
        if _is_bullet_text(text):
            return 'bullet', current_section

        # 2. Inline skills line ("Key Skills: ...", "Core Competencies: ...")
        if _INLINE_SKILLS_RE.match(text):
            return 'skills', current_section

        # 3. Bold text with colon = achievement bullet (check BEFORE date
        #    pattern so bullets mentioning dates like "October 2025" aren't misclassified)
        if bold and ':' in text and len(text) > 50:
            return 'bullet', current_section

        # 4. Job header: has date pattern (title or company line with dates)
        if _has_date_pattern(text):
            return 'job_header', current_section

        # 5. Bold short text after job_header/section_header = job title
        if bold and len(text) < 120 and prev_type in ('job_header', 'section_header'):
            return 'job_header', current_section

        # 6. Non-bold company line with location/industry markers
        if not bold and _is_company_line(text):
            return 'job_header', current_section

        # 7. Non-bold achievement text (starts with action verb or has strong metrics)
        if not bold and _is_achievement_text(text):
            return 'bullet', current_section

        # 8. Short non-bold text: standalone company name vs job intro
        if not bold and len(text) < 80:
            if _looks_like_company_name(text, prev_type):
                return 'job_header', current_section
            return 'job_intro', current_section

        # 9. Non-bold longer text without action verbs = job intro (description)
        if not bold and len(text) > 80:
            return 'job_intro', current_section

        # 10. Bold longer text = bullet
        if bold and len(text) > 50:
            return 'bullet', current_section

        return 'bullet', current_section

    if current_section == 'education':
        return 'education', current_section

    if current_section == 'certification':
        return 'certification', current_section

    if current_section == 'skills':
        return 'skills', current_section

    if current_section == 'keywords':
        return 'keywords', current_section

    if current_section == 'summary':
        # Bold text with colon = highlight bullet (common in executive resumes)
        if bold and ':' in text and len(text) > 50:
            return 'highlights', 'highlights'
        # Pipe-separated = keywords/skills line
        if '|' in text and text.count('|') >= 3:
            return 'keywords', 'keywords'
        return 'summary', current_section

    if current_section == 'highlights':
        # Bold text with colon = highlight bullet
        if bold and ':' in text:
            return 'highlights', current_section
        # Pipe-separated = keywords/skills line
        if '|' in text and text.count('|') >= 3:
            return 'keywords', 'keywords'
        return 'highlights', current_section

    if current_section == 'references':
        return 'reference', current_section

    if current_section == 'header':
        # Still in header area
        if _has_contact_info(text):
            return 'header', 'header'
        # Summary text after header
        if not bold and len(text) > 80:
            return 'summary', 'summary'
        # Highlight bullets with bold + colon
        if bold and ':' in text:
            return 'highlights', 'highlights'
        # Keywords line (pipe-separated)
        if '|' in text and text.count('|') >= 3:
            return 'keywords', 'keywords'
        return 'header', 'header'

    # --- Fallback ---
    return 'unknown', current_section


def parse_resume_structure(file_path: str) -> list[dict]:
    """Parse a .docx resume into structured sections with formatting metadata.

    Args:
        file_path: Path to a .docx resume file.

    Returns:
        Ordered list of dicts, each with:
            - type: str - paragraph classification
            - text: str - paragraph text content
            - formatting: dict - font_size, bold, italic, alignment, font_name, underline, space_before, space_after
            - paragraph_index: int - position in document
            - parent_section: str - which section this paragraph belongs to
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"Resume file not found: {file_path}")
    if path.suffix.lower() != '.docx':
        raise ValueError(f"Expected .docx file, got: {path.suffix}")

    doc = Document(str(path))

    results = []
    current_section = ''
    total_paragraphs = len(doc.paragraphs)
    prev_type = None

    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue

        formatting = _extract_formatting(paragraph)
        is_early = i < min(12, total_paragraphs * 0.15)

        para_type, current_section = _classify_paragraph(
            text=text,
            formatting=formatting,
            para_index=i,
            current_section=current_section,
            is_early=is_early,
            prev_type=prev_type,
        )

        results.append({
            'type': para_type,
            'text': text,
            'formatting': formatting,
            'paragraph_index': i,
            'parent_section': current_section,
        })

        prev_type = para_type

    return results


# ---------------------------------------------------------------------------
# KB conversion helpers
# ---------------------------------------------------------------------------

_MONTH_MAP = {
    'jan': 1, 'feb': 2, 'mar': 3, 'apr': 4, 'may': 5, 'jun': 6,
    'jul': 7, 'aug': 8, 'sep': 9, 'oct': 10, 'nov': 11, 'dec': 12,
    'january': 1, 'february': 2, 'march': 3, 'april': 4, 'june': 6,
    'july': 7, 'august': 8, 'september': 9, 'october': 10, 'november': 11, 'december': 12,
}


def _parse_date_str(s: str | None) -> date | None:
    """Convert a date string like 'June 2024', 'Aug 2021', '2012' to a date object."""
    if not s:
        return None
    s = s.strip()
    if s.lower() in ('present', 'current'):
        return None

    # Try "Month Year" format
    m = re.match(r'^([A-Za-z]+)\s*,?\s*(\d{4})$', s)
    if m:
        month_str = m.group(1).lower()
        year = int(m.group(2))
        month = _MONTH_MAP.get(month_str)
        if month:
            return date(year, month, 1)

    # Try bare year "2012"
    m = re.match(r'^(\d{4})$', s)
    if m:
        return date(int(m.group(1)), 1, 1)

    return None


_BRACES_RE = re.compile(r'\{([^})]+)[})]')
_PARENS_NOTE_RE = re.compile(r'\(([^)]*(?:Onsite|Remote|Hybrid)[^)]*)\)', re.IGNORECASE)
_LOCATION_RE = re.compile(r'([A-Z][a-z]+(?:\s[A-Z][a-z]+)*),\s*([A-Z]{2})\b')
_DATE_RANGE_RE = re.compile(
    r'((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\w*[\s,]*\d{4})'
    r'\s*[-–—]\s*'
    r'((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\w*[\s,]*\d{4}|[Pp]resent|[Cc]urrent)',
)
_YEAR_RANGE_RE = re.compile(r'(\d{4})\s*[-–—]\s*(\d{4}|[Pp]resent|[Cc]urrent)')
_TITLE_WORDS = re.compile(
    r'\b(?:Director|VP|Vice President|President|Chief|CTO|CEO|CIO|CFO|COO|CPO|CMO|'
    r'Manager|Engineer|Architect|Lead|Head|Senior|Principal|Staff|Founder|'
    r'Consultant|Analyst|Developer|Coordinator|Officer|Advisor)\b',
    re.IGNORECASE
)


def _is_company_line(text: str) -> bool:
    """Heuristic: line is a company header (vs a title line)."""
    if _BRACES_RE.search(text):
        return True
    if _PARENS_NOTE_RE.search(text):
        return True
    loc = _LOCATION_RE.search(text)
    if loc and not _TITLE_WORDS.match(text.split(',')[0].strip()):
        return True
    return False


def _parse_company_line(text: str) -> dict:
    """Extract employer, location, dates, industry from a company line."""
    industry = None
    m = _BRACES_RE.search(text)
    if m:
        industry = m.group(1).strip()
        text = text[:m.start()] + text[m.end():]

    work_mode = None
    m = _PARENS_NOTE_RE.search(text)
    if m:
        work_mode = m.group(1).strip()
        text = text[:m.start()] + text[m.end():]

    start_date, end_date = None, None
    m = _DATE_RANGE_RE.search(text)
    if m:
        start_date = m.group(1).strip()
        end_date = m.group(2).strip()
        text = text[:m.start()] + text[m.end():]
    else:
        m = _YEAR_RANGE_RE.search(text)
        if m:
            start_date = m.group(1).strip()
            end_date = m.group(2).strip()
            text = text[:m.start()] + text[m.end():]

    location = None
    m = _LOCATION_RE.search(text)
    if m:
        location = f"{m.group(1)}, {m.group(2)}"

    employer = re.sub(r'\s*,\s*$', '', text.strip()).strip()
    employer = re.sub(r'\s{2,}', ' ', employer).strip(' ,\t')
    # Remove empty trailing parens from parse artifacts
    employer = re.sub(r'\s*\(\s*\)\s*$', '', employer).strip()

    is_current = end_date is not None and end_date.lower() in ('present', 'current')

    return {
        'employer': employer or 'Unknown',
        'title': '',
        'start_date': start_date,
        'end_date': None if is_current else end_date,
        'location': location,
        'industry': industry,
        'intro_text': '',
        'is_current': is_current,
        'bullets': [],
        'notes': work_mode,
    }


def _parse_title_dates(text: str) -> tuple[str, str | None, str | None]:
    """Extract title and optional date range from a title line."""
    start, end = None, None

    # Try standard dash/en-dash separator first
    m = _DATE_RANGE_RE.search(text)
    if m:
        start = m.group(1).strip()
        end = m.group(2).strip()
        text = text[:m.start()].strip().rstrip(',').strip()
    else:
        m = _YEAR_RANGE_RE.search(text)
        if m:
            start = m.group(1).strip()
            end = m.group(2).strip()
            text = text[:m.start()].strip().rstrip(',').strip()

    # Also try pipe-separated dates: "Title  Aug 2021 | Jan 2024"
    if not start:
        pipe_date = re.search(
            r'((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\w*\s+\d{4})'
            r'\s*\|\s*'
            r'((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\w*\s+\d{4}|[Pp]resent|[Cc]urrent)',
            text,
        )
        if pipe_date:
            start = pipe_date.group(1).strip()
            end = pipe_date.group(2).strip()
            text = text[:pipe_date.start()].strip().rstrip(',|').strip()

    # Try tab-separated dates: "Title\tAug 2021 - Jan 2024"
    if not start and '\t' in text:
        parts = text.split('\t', 1)
        if len(parts) == 2:
            m2 = _DATE_RANGE_RE.search(parts[1])
            if not m2:
                m2 = _YEAR_RANGE_RE.search(parts[1])
            if m2:
                start = m2.group(1).strip()
                end = m2.group(2).strip()
                text = parts[0].strip()

    # Clean up title artifacts
    text = text.strip('\t ,|')
    # Remove dangling open parens with no close: "CTO/Co-Founder ("
    text = re.sub(r'\s*\(\s*$', '', text)
    # Remove empty parens: "Director ()"
    text = re.sub(r'\s*\(\s*\)\s*$', '', text)
    # Remove [Source: ...] tags
    text = re.sub(r'\[Source:[^\]]*\]', '', text)
    # Remove "Listed under..." prefix
    text = re.sub(r'^Listed under.*?:\s*', '', text)
    # Collapse multiple spaces
    text = re.sub(r'\s{2,}', ' ', text).strip()
    return text, start, end


def _parse_oneliner(text: str) -> dict | None:
    """Parse a one-liner 'Title | Company (Industry)  Dates' entry.

    Handles multi-pipe entries like:
        Senior Testing AI Manager | QA Engineering (Part Time) | Fact Finders Pro (NPO Startup)  May 2025 - Present
    The LAST pipe-separated segment is the company (possibly with industry in parens).
    Everything before is the title.
    """
    # Split off dates from end (tab or multi-space separated)
    parts = re.split(r'\t+|\s{3,}', text, maxsplit=1)
    main_part = parts[0].strip()
    date_part = parts[1].strip() if len(parts) > 1 else ''

    if '|' not in main_part:
        return None

    segments = [s.strip() for s in main_part.split('|')]
    if len(segments) < 2:
        return None

    # Last segment is the company (possibly with industry in parens)
    company_seg = segments[-1]
    title_seg = ' | '.join(segments[:-1])

    # Extract industry from nested parens at end of company segment
    industry = None
    m = re.search(r'\((.+)\)\s*$', company_seg)
    if m:
        industry = m.group(1).strip()
        company_seg = company_seg[:m.start()].strip()

    # Parse dates
    start_date, end_date = None, None
    dm = _DATE_RANGE_RE.search(date_part)
    if dm:
        start_date = dm.group(1).strip()
        end_date = dm.group(2).strip()
    else:
        dm = _YEAR_RANGE_RE.search(date_part)
        if dm:
            start_date = dm.group(1).strip()
            end_date = dm.group(2).strip()

    is_current = end_date is not None and end_date.lower() in ('present', 'current')
    return {
        'employer': company_seg or 'Unknown',
        'title': title_seg,
        'start_date': start_date,
        'end_date': None if is_current else end_date,
        'location': None,
        'industry': industry,
        'intro_text': '',
        'is_current': is_current,
        'bullets': [],
        'notes': None,
    }


def parse_resume_for_kb(file_path: str) -> dict:
    """Parse a .docx resume into the nested KB format for DB insert.

    Uses parse_resume_structure() to classify paragraphs, then groups them
    into career_history entries with nested bullets, plus skills, education, etc.

    Returns:
        dict with:
            career_history: list of dicts (employer, title, dates, location, intro_text, bullets)
            skills: list of skill name strings
            education: list of education line strings
            certifications: list of certification line strings
            highlights: list of highlight bullet strings
            summary: str - professional summary text
            confidence: float 0-1
    """
    paragraphs = parse_resume_structure(file_path)

    career_history = []
    current_job = None
    in_additional = False
    skills = []
    education = []
    certifications = []
    highlights = []
    summary_parts = []

    for para in paragraphs:
        ptype = para['type']
        text = para['text']
        section = para['parent_section']

        # -- Section header: "Additional Work Experience" flips one-liner mode --
        if ptype == 'section_header':
            if current_job:
                career_history.append(current_job)
                current_job = None
            if 'additional' in text.lower():
                in_additional = True
            else:
                in_additional = False
            continue

        # -- Highlights (top-of-resume achievement bullets) --
        if ptype == 'highlights':
            highlights.append(text)
            continue

        # -- Summary --
        if ptype == 'summary':
            summary_parts.append(text)
            continue

        # -- Skills / Keywords (pipe, comma, or tab separated) --
        if ptype in ('skills', 'keywords'):
            # Split on pipes first (most common in skill lines), then tabs
            for chunk in re.split(r'[|\t]', text):
                # Then split on commas, but protect numbers like "32,000"
                for part in re.split(r',\s+(?![0-9])', chunk):
                    skill = part.strip().strip('•–—·')
                    if not skill or len(skill) > 55 or len(skill) < 2:
                        continue
                    # Reject sentence-like fragments
                    words = skill.split()
                    if len(words) > 7:
                        continue
                    if re.match(r'^\d', skill) and not re.match(r'^\d+\+?\s*(years?|yrs?)', skill, re.I):
                        continue  # starts with number but isn't "5+ years"
                    # Reject if ends with period or closing paren (sentence fragment)
                    if skill.endswith('.') or skill.endswith(')'):
                        continue
                    # Reject if contains % or $ (metric, not a skill)
                    if '%' in skill or '$' in skill:
                        continue
                    # Reject conjunction fragments ("and AWS", "or Python")
                    if re.match(r'^(and|or|with|the|for|to|of|in|a|an)\s', skill, re.I):
                        continue
                    # Reject category headers with colon ("Database Platforms: MySQL")
                    if ':' in skill:
                        # Split and keep only the part after colon as the skill
                        after_colon = skill.split(':', 1)[1].strip()
                        if after_colon and len(after_colon) > 2 and len(after_colon) < 50:
                            skill = after_colon
                            words = skill.split()
                        else:
                            continue
                    # Reject single common words that aren't skills
                    if len(words) == 1 and len(skill) < 4:
                        continue
                    # Reject phrases that START with a verb (achievement language)
                    first_word_lower = words[0].lower().rstrip('.,;:')
                    if first_word_lower in _ACHIEVEMENT_VERBS:
                        continue
                    # Also reject requirement-like words
                    if first_word_lower in {'must', 'ability', 'required', 'need', 'needs',
                                            'requires', 'able', 'responsible'}:
                        continue
                    # Reject if it reads like a sentence (>4 words with common sentence words)
                    if len(words) > 4:
                        filler = {'with', 'that', 'this', 'from', 'into', 'have', 'been',
                                  'able', 'level', 'high', 'our', 'their', 'your'}
                        if sum(1 for w in words if w.lower() in filler) >= 2:
                            continue
                    skills.append(skill)
            continue

        # -- Education --
        if ptype == 'education':
            education.append(text)
            continue

        # -- Certifications (may be tab-separated on one line) --
        if ptype == 'certification':
            # Split on tabs first (some resumes put 2 certs per line)
            for chunk in re.split(r'\t+', text):
                cert = chunk.strip()
                if cert and len(cert) > 3:
                    certifications.append(cert)
            continue

        # -- Experience section content --
        if section in ('experience', 'additional') or in_additional:

            if ptype == 'job_header':
                # In additional section, try one-liner parse
                if in_additional:
                    entry = _parse_oneliner(text)
                    if entry and not _is_junk_employer(entry['employer']):
                        career_history.append(entry)
                        continue

                # Determine if this is a company line or title line
                if _is_company_line(text):
                    # Reject junk ONLY when it would become an employer
                    if _is_junk_employer(text):
                        continue
                    if current_job:
                        career_history.append(current_job)
                    current_job = _parse_company_line(text)
                elif current_job is not None and not current_job['title']:
                    # First title for current company
                    title, start, end = _parse_title_dates(text)
                    current_job['title'] = title
                    if start:
                        current_job['start_date'] = start
                    if end:
                        current_job['end_date'] = None if end.lower() in ('present', 'current') else end
                        current_job['is_current'] = end.lower() in ('present', 'current')
                elif current_job is not None and current_job['title']:
                    # Decide: another role at same company, or a new company?
                    # If text has title words or date patterns → title at same company
                    # If neither → likely a new company name
                    has_title_words = bool(_TITLE_WORDS.search(text))
                    has_dates = _has_date_pattern(text)
                    if has_title_words or has_dates:
                        # Additional role at same company — split into separate entry
                        career_history.append(current_job)
                        title, start, end = _parse_title_dates(text)
                        current_job = {
                            'employer': current_job['employer'],
                            'title': title,
                            'start_date': start or current_job.get('start_date'),
                            'end_date': (None if end and end.lower() in ('present', 'current') else end) or current_job.get('end_date'),
                            'location': current_job.get('location'),
                            'industry': current_job.get('industry'),
                            'intro_text': '',
                            'is_current': bool(end and end.lower() in ('present', 'current')),
                            'bullets': [],
                            'notes': current_job.get('notes'),
                        }
                    else:
                        # No title words or dates — treat as new company
                        if _is_junk_employer(text):
                            continue
                        career_history.append(current_job)
                        current_job = _parse_company_line(text)
                else:
                    # No current job context — treat as company
                    if _is_junk_employer(text):
                        continue
                    if current_job:
                        career_history.append(current_job)
                    current_job = _parse_company_line(text)

            elif ptype == 'job_intro':
                if current_job:
                    if current_job['intro_text']:
                        current_job['intro_text'] += ' ' + text
                    else:
                        current_job['intro_text'] = text

            elif ptype == 'bullet':
                if current_job:
                    current_job['bullets'].append(text)

        # -- Header / headline / unknown — skip for KB purposes --

    # Flush last job
    if current_job:
        career_history.append(current_job)

    # Convert date strings to date objects for DB compatibility
    for job in career_history:
        job['start_date'] = _parse_date_str(job.get('start_date'))
        job['end_date'] = _parse_date_str(job.get('end_date'))

    # Deduplicate skills preserving order
    seen_skills = set()
    unique_skills = []
    for s in skills:
        key = s.lower()
        if key not in seen_skills:
            seen_skills.add(key)
            unique_skills.append(s)

    # Confidence score
    has_jobs = bool(career_history)
    has_bullets = any(j['bullets'] for j in career_history)
    has_skills = bool(unique_skills)
    has_edu = bool(education)
    signals = sum([has_jobs, has_bullets, has_skills, has_edu])

    logger.info(
        "KB parse: %d jobs, %d total bullets, %d skills, %d education, %d highlights",
        len(career_history),
        sum(len(j['bullets']) for j in career_history),
        len(unique_skills),
        len(education),
        len(highlights),
    )

    return {
        'career_history': career_history,
        'skills': unique_skills,
        'education': education,
        'certifications': certifications,
        'highlights': highlights,
        'summary': ' '.join(summary_parts),
        'confidence': round(signals / 4, 2),
    }
